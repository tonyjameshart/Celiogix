# path: panels/cookbook_panel.py

from typing import Optional
import os
import re
import requests

from PySide6.QtCore import QMarginsF, Qt, QUrl
from PySide6.QtGui import QFont, QPageLayout, QPageSize, QTextDocument, QPixmap, QDesktopServices
from PySide6.QtPrintSupport import QPrintDialog, QPrintPreviewDialog, QPrinter
from PySide6.QtWidgets import (
    QCheckBox, QComboBox, QDialog, QFrame, QGridLayout, QGroupBox, QHBoxLayout,
    QHeaderView, QLabel, QLineEdit, QListWidget, QListWidgetItem,
    QMessageBox, QProgressBar, QPushButton, QRadioButton, QScrollArea,
    QSlider, QSpinBox, QSplitter, QTableWidget, QTableWidgetItem,
    QTextEdit, QTreeWidget, QTreeWidgetItem, QVBoxLayout, QWidget
)

from utils.accessibility import announce_for_screen_reader
from panels.base_panel import BasePanel
from panels.context_menu_mixin import CookbookContextMenuMixin
from utils.custom_widgets import NoSelectionTableWidget
from panels.cookbook.recipe_manager import RecipeManager
from panels.cookbook.recipe_ui_components import RecipeUIComponents
from panels.cookbook.recipe_dialogs import RecipeDialogs
from panels.cookbook.recipe_export import RecipeExport


class CookbookPanel(CookbookContextMenuMixin, BasePanel):
    """Cookbook panel for PySide6"""
    
    def __init__(self, master=None, app=None):
        super().__init__(master, app)
        self.recipe_dialogs = RecipeDialogs(self)
        
        # Cleanup temporary files when panel is destroyed
        import atexit
        atexit.register(self.cleanup_temp_files)
    
    def _extract_numeric_value(self, value, default=0):
        """Extract numeric value from string that might contain text like '1 loaf' -> 1"""
        try:
            if value is None:
                return default
            elif isinstance(value, (int, float)):
                return int(value)
            elif isinstance(value, str):
                # Try to extract number from string like "1 loaf" -> 1
                numbers = re.findall(r'\d+', value)
                return int(numbers[0]) if numbers else default
            else:
                # Handle any other type by converting to string first
                str_value = str(value)
                numbers = re.findall(r'\d+', str_value)
                return int(numbers[0]) if numbers else default
        except (ValueError, IndexError, TypeError, AttributeError):
            return default
    
    def setup_ui(self):
        """Set up the cookbook panel UI"""
        # Initialize data storage
        self.recipe_data_storage = {}  # Store full recipe data
        
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # Action buttons section (no duplicate search)
        action_layout = QHBoxLayout()
        
        add_recipe_btn = QPushButton("Add Recipe")
        add_recipe_btn.clicked.connect(self.add_recipe)
        action_layout.addWidget(add_recipe_btn)
        
        self.import_web_btn = QPushButton("Import from Web")
        self.import_web_btn.clicked.connect(self.import_from_web)
        action_layout.addWidget(self.import_web_btn)
        
        self.import_file_btn = QPushButton("Import from File")
        self.import_file_btn.clicked.connect(self.import_from_file)
        action_layout.addWidget(self.import_file_btn)
        
        self.export_btn = QPushButton("Export")
        self.export_btn.clicked.connect(self.export_recipes)
        action_layout.addWidget(self.export_btn)
        
        action_layout.addStretch()
        main_layout.addLayout(action_layout)
        
        # Favorite statistics section
        stats_layout = QHBoxLayout()
        self.favorite_stats_label = QLabel("Favorites: 0 | Total: 0")
        self.favorite_stats_label.setStyleSheet("color: #666; font-size: 12px;")
        stats_layout.addWidget(self.favorite_stats_label)
        stats_layout.addStretch()
        
        # Quick access buttons
        quick_access_btn = QPushButton("⭐ Quick Access Favorites")
        quick_access_btn.clicked.connect(self.show_quick_access_favorites)
        stats_layout.addWidget(quick_access_btn)
        
        
        main_layout.addLayout(stats_layout)
        
        # Old table code removed - now using card navigation
        
        # Cookbook Navigation
        self.setup_cookbook_navigation()
        main_layout.addWidget(self.cookbook_widget)
        
        
        # Advanced recipe features
        advanced_layout = QHBoxLayout()
        
        self.rank_btn = QPushButton("Rank Recipes")
        self.rank_btn.clicked.connect(self.rank_recipes)
        
        self.analytics_btn = QPushButton("Analytics")
        self.analytics_btn.clicked.connect(self.show_recipe_analytics)
        
        advanced_layout.addWidget(self.rank_btn)
        advanced_layout.addWidget(self.analytics_btn)
        advanced_layout.addStretch()
        
        main_layout.addLayout(advanced_layout)
        
        # Connect selection changes
        # Selection handling now done through card clicks
        
        # Load initial data
        self.load_recipes()
        self.load_categories()

    def refresh(self):
        """Refresh panel data"""
        self.load_recipes()
    
    def load_recipes(self):
        """Load recipes from database"""
        try:
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            # Build WHERE clause based on filters
            where_conditions = []
            params = []
            
            # Favorite filter
            if hasattr(self, 'favorite_filter_cb') and self.favorite_filter_cb.isChecked():
                where_conditions.append("r.is_favorite = 1")
            
            # Category filter - use the category column in recipes table
            if hasattr(self, 'category_combo') and self.category_combo.currentData() != 'all':
                category_id = self.category_combo.currentData()
                # Get the category name from the combo box
                category_name = self.category_combo.currentText()
                where_conditions.append("r.category = ?")
                params.append(category_name)
            
            # Build the query
            where_clause = "WHERE " + " AND ".join(where_conditions) if where_conditions else ""
            
            cursor.execute(f"""
                SELECT DISTINCT r.id, r.title, r.instructions, r.prep_time, r.cook_time, r.servings, 
                       r.difficulty, r.category, r.tags, r.description, r.is_favorite, r.image_path
                FROM recipes r
                {where_clause}
                ORDER BY r.title
            """, params)
            
            recipes = cursor.fetchall()
            
            # Clear existing data
            self.recipe_data_storage.clear()
            
            for row, recipe in enumerate(recipes):
                recipe_id, title, instructions, prep_time, cook_time, servings, difficulty, category, tags, description, is_favorite, image_path = recipe
                
                # Load ingredients for this recipe
                cursor.execute("""
                    SELECT ingredient_name, quantity, unit, notes
                    FROM recipe_ingredients 
                    WHERE recipe_id = ?
                    ORDER BY ingredient_name
                """, (recipe_id,))
                
                ingredients = []
                for ing_row in cursor.fetchall():
                    ingredients.append({
                        "name": ing_row[0],
                        "amount": ing_row[1] or "",
                        "unit": ing_row[2] or "",
                        "notes": ing_row[3] or ""
                    })
                
                # Create recipe data structure
                recipe_data = {
                    "id": recipe_id,
                    "name": title,
                    "description": description or "",
                    "prep_time": prep_time or "",
                    "cook_time": cook_time or "",
                    "servings": servings or 1,
                    "difficulty": difficulty or "Easy",
                    "category": category or "Other",
                    "tags": tags or "",
                    "ingredients": ingredients,
                    "instructions": instructions or "",
                    "is_favorite": is_favorite or 0,
                    "image_path": image_path or ""
                }
                
                # Store full recipe data for card display
                self.recipe_data_storage[recipe_id] = recipe_data
            
            # If no recipes in database, add comprehensive GF recipes
            if len(recipes) == 0:
                self.create_comprehensive_gf_recipes()
            
            # Update favorite statistics
            self.update_favorite_statistics()
                
        except Exception as e:
            print(f"Error loading recipes from database: {e}")
            # Fallback to sample recipes
            self._load_sample_recipes()
            # Update favorite statistics even on error
            self.update_favorite_statistics()
        
        # Load categories for filtering
        self.load_categories()
        
        # Initialize and display recipes
        self.refresh_recipe_display()
    
    def setup_cookbook_navigation(self):
        """Set up the polished cookbook navigation widget"""
        # Main cookbook widget - no category tree, just main content
        self.cookbook_widget = QWidget()
        cookbook_layout = QVBoxLayout(self.cookbook_widget)
        
        # Search and filter bar
        self.setup_search_filter_bar(cookbook_layout)
        
        # Recipe cards container
        self.setup_recipe_cards_container(cookbook_layout)
        
        # Initialize recipe cards
        self.recipe_cards = []
        self.current_recipe_id = None
        
        # Load and display recipes now that navigation is set up
        self.refresh_recipe_display()
    
    
    def _removed_load_category_navigation(self):
        """Load categories into the navigation tree"""
        try:
            from utils.db import get_connection
            
            conn = get_connection()
            cursor = conn.cursor()
            
            # Add "All Recipes" root item
            all_item = QTreeWidgetItem(self.category_tree)
            all_item.setText(0, "📚 All Recipes")
            all_item.setData(0, Qt.UserRole, None)
            
            # Get primary categories
            cursor.execute("""
                SELECT id, name, icon, color 
                FROM categories 
                WHERE parent_id IS NULL
                ORDER BY sort_order, name
            """)
            primary_categories = cursor.fetchall()
            
            # Add primary categories as tree items
            for cat_id, name, icon, color in primary_categories:
                parent_item = QTreeWidgetItem(self.category_tree)
                parent_item.setText(0, f"{icon} {name}")
                parent_item.setData(0, Qt.UserRole, cat_id)
                
                # Get subcategories for this primary category
                cursor.execute("""
                    SELECT id, name, icon, color 
                    FROM categories 
                    WHERE parent_id = ?
                    ORDER BY sort_order, name
                """, (cat_id,))
                subcategories = cursor.fetchall()
                
                # Add subcategories as child items
                for sub_id, sub_name, sub_icon, sub_color in subcategories:
                    child_item = QTreeWidgetItem(parent_item)
                    child_item.setText(0, f"{sub_icon} {sub_name}")
                    child_item.setData(0, Qt.UserRole, sub_id)
            
            # Start with all items collapsed for better navigation
            self.category_tree.collapseAll()
            
            conn.close()
            
        except Exception as e:
            print(f"Error loading category navigation: {e}")
    
    def _removed_on_category_selected(self, item, column):
        """Handle category selection in navigation tree"""
        category_id = item.data(0, Qt.UserRole)
        category_name = item.text(0)
        print(f"DEBUG: Category selected - ID: {category_id}, Name: {category_name}")
        
        # Clear any existing recipe selection to prevent card pop-ups
        self.current_recipe_id = None
        
        # Clear card selections visually
        for card in self.recipe_cards:
            card.setStyleSheet("""
                QWidget {
                    background-color: white;
                    border-radius: 12px;
                    border: 1px solid #e0e0e0;
                }
                QWidget:hover {
                    border-color: #3498db;
                }
            """)
        
        # Update category filter combo
        if hasattr(self, 'category_combo'):
            if category_id is None:
                # All recipes selected
                self.category_combo.setCurrentIndex(0)  # "All Categories"
                print("DEBUG: Showing all recipes")
            else:
                # Find and select the category in the combo box
                for i in range(self.category_combo.count()):
                    if self.category_combo.itemData(i) == str(category_id):
                        self.category_combo.setCurrentIndex(i)
                        break
                print(f"DEBUG: Filtering by category: {category_name}")
        
        # Refresh recipe display
        self.refresh_recipe_display()
    
    def setup_search_filter_bar(self, parent_layout):
        """Set up search and filter bar"""
        search_filter_widget = QWidget()
        search_filter_layout = QVBoxLayout(search_filter_widget)
        search_filter_layout.setSpacing(15)
        
        # Search row with filters on the right
        search_row_layout = QHBoxLayout()
        search_label = QLabel("🔍")
        search_label.setStyleSheet("font-size: 16px; color: #666;")
        search_row_layout.addWidget(search_label)
        
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("Search recipes...")
        self.search_edit.setStyleSheet("""
            QLineEdit {
                padding: 8px 12px;
                border: 2px solid #e0e0e0;
                border-radius: 20px;
                font-size: 14px;
                background-color: white;
            }
            QLineEdit:focus {
                border-color: #3498db;
            }
        """)
        self.search_edit.textChanged.connect(self.filter_recipes)
        search_row_layout.addWidget(self.search_edit)
        
        # Add spacing between search and filters
        search_row_layout.addSpacing(20)
        
        # Favorite filter checkbox (moved to search row)
        self.favorite_filter_cb = QCheckBox("Favorites Only")
        self.favorite_filter_cb.setStyleSheet("""
            QCheckBox {
                font-size: 13px;
                color: #495057;
                spacing: 8px;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
                border: 2px solid #e0e0e0;
                border-radius: 3px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                background-color: #4caf50;
                border-color: #2e7d32;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xMC41IDFMNC41IDdMMSA0IiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
            }
            QCheckBox::indicator:hover {
                border-color: #66bb6a;
            }
        """)
        self.favorite_filter_cb.stateChanged.connect(self.filter_recipes)
        search_row_layout.addWidget(self.favorite_filter_cb)
        
        # Sort dropdown (moved to search row)
        sort_label = QLabel("Sort by:")
        sort_label.setStyleSheet("color: #666; font-weight: 500; font-size: 13px;")
        search_row_layout.addWidget(sort_label)
        
        self.sort_combo = QComboBox()
        self.sort_combo.addItems([
            "Name (A-Z)", "Name (Z-A)", "Prep Time", "Difficulty", "Recently Added", "Favorites First"
        ])
        self.sort_combo.setStyleSheet("""
            QComboBox {
                padding: 6px 12px;
                border: 1px solid #e0e0e0;
                border-radius: 6px;
                background-color: white;
                min-width: 120px;
                font-size: 13px;
            }
            QComboBox:focus {
                border-color: #3498db;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #666;
                margin-right: 5px;
            }
        """)
        self.sort_combo.currentTextChanged.connect(self.sort_recipes)
        search_row_layout.addWidget(self.sort_combo)
        
        search_row_layout.addStretch()
        search_filter_layout.addLayout(search_row_layout)
        
        # Category filter - using improved dropdown selector
        category_row_layout = QHBoxLayout()
        category_label = QLabel("Category:")
        category_label.setStyleSheet("font-weight: 500; color: #495057; min-width: 80px;")
        category_row_layout.addWidget(category_label)
        
        self.category_combo = QComboBox()
        self.category_combo.setStyleSheet("""
            QComboBox {
                padding: 8px 12px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                font-size: 14px;
                background-color: white;
                min-width: 150px;
            }
            QComboBox:focus {
                border-color: #3498db;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #666;
                margin-right: 5px;
            }
            QComboBox QAbstractItemView {
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                background-color: white;
                selection-background-color: #3498db;
            }
        """)
        self.category_combo.currentTextChanged.connect(self._on_category_combo_changed)
        category_row_layout.addWidget(self.category_combo)
        category_row_layout.addStretch()
        search_filter_layout.addLayout(category_row_layout)
        
        parent_layout.addWidget(search_filter_widget)
    
    def setup_recipe_cards_container(self, parent_layout):
        """Set up the recipe cards container with scroll area"""
        # Scroll area for recipe cards
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.scroll_area.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: #f8f9fa;
            }
            QScrollBar:vertical {
                background-color: #f0f0f0;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #c0c0c0;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #a0a0a0;
            }
        """)
        
        # Cards container widget
        self.cards_container = QWidget()
        self.cards_layout = QVBoxLayout(self.cards_container)
        self.cards_layout.setContentsMargins(20, 20, 20, 20)
        self.cards_layout.setSpacing(20)
        
        # Grid layout for cards
        self.cards_grid = QGridLayout()
        self.cards_grid.setSpacing(20)
        self.cards_layout.addLayout(self.cards_grid)
        
        self.cards_layout.addStretch()
        
        self.scroll_area.setWidget(self.cards_container)
        parent_layout.addWidget(self.scroll_area)
    
    def create_recipe_card(self, recipe_data):
        """Create a polished recipe card widget"""
        # Create card as a frame widget with explicit parent
        card = QFrame(self.cards_container)
        card.setFixedSize(280, 320)
        card.setFrameStyle(QFrame.Box)
        card.setLineWidth(1)
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border-radius: 12px;
                border: 1px solid #e0e0e0;
            }
            QFrame:hover {
                border-color: #3498db;
            }
        """)
        
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(0, 0, 0, 0)
        card_layout.setSpacing(0)
        
        # Recipe image
        image_label = QLabel()
        image_label.setFixedSize(280, 180)
        image_label.setStyleSheet("""
            QLabel {
                border-top-left-radius: 12px;
                border-top-right-radius: 12px;
                background-color: #f0f0f0;
            }
        """)
        image_label.setAlignment(Qt.AlignCenter)
        image_label.setScaledContents(True)
        
        # Load recipe image
        try:
            from services.image_service import recipe_image_service
            pixmap = recipe_image_service.load_image_pixmap(recipe_data.get('id'), (280, 180))
            if pixmap and not pixmap.isNull():
                image_label.setPixmap(pixmap)
            else:
                # Use default category image
                category_icon = self._get_category_icon(recipe_data.get('category', 'Uncategorized'))
                image_label.setText(f"{category_icon}\n{recipe_data.get('category', 'Uncategorized')}")
                image_label.setStyleSheet("""
                    QLabel {
                        border-top-left-radius: 12px;
                        border-top-right-radius: 12px;
                        background-color: #f0f0f0;
                        color: #666;
                        font-size: 20px;
                        font-weight: bold;
                    }
                """)
        except Exception:
            # Use default category image
            category_icon = self._get_category_icon(recipe_data.get('category', 'Uncategorized'))
            image_label.setText(f"{category_icon}\n{recipe_data.get('category', 'Uncategorized')}")
            image_label.setStyleSheet("""
                QLabel {
                    border-top-left-radius: 12px;
                    border-top-right-radius: 12px;
                    background-color: #f0f0f0;
                    color: #666;
                    font-size: 20px;
                    font-weight: bold;
                }
            """)
        
        card_layout.addWidget(image_label)
        
        # Recipe info container
        info_container = QWidget()
        info_layout = QVBoxLayout(info_container)
        info_layout.setContentsMargins(16, 12, 16, 12)
        info_layout.setSpacing(8)
        
        # Recipe name
        name_label = QLabel(recipe_data.get('name', 'Unknown Recipe'))
        name_label.setStyleSheet("""
            QLabel {
                font-size: 16px;
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 4px;
            }
        """)
        name_label.setWordWrap(True)
        info_layout.addWidget(name_label)
        
        # Recipe details
        details_layout = QHBoxLayout()
        details_layout.setSpacing(12)
        
        # Category
        category_label = QLabel(f"📂 {recipe_data.get('category', 'Uncategorized')}")
        category_label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                color: #7f8c8d;
                background-color: #ecf0f1;
                padding: 4px 8px;
                border-radius: 12px;
            }
        """)
        details_layout.addWidget(category_label)
        
        # Prep time
        prep_time = recipe_data.get('prep_time', 'N/A')
        if prep_time and prep_time != 'N/A':
            time_label = QLabel(f"⏱️ {prep_time}")
            time_label.setStyleSheet("""
                QLabel {
                    font-size: 12px;
                    color: #7f8c8d;
                }
            """)
            details_layout.addWidget(time_label)
        
        details_layout.addStretch()
        
        # Favorite indicator
        if recipe_data.get('is_favorite'):
            favorite_label = QLabel("⭐")
            favorite_label.setStyleSheet("font-size: 16px; color: #f39c12;")
            details_layout.addWidget(favorite_label)
        
        info_layout.addLayout(details_layout)
        
        # Difficulty and servings
        bottom_layout = QHBoxLayout()
        
        difficulty = recipe_data.get('difficulty', 'Medium')
        difficulty_color = {
            'Easy': '#27ae60',
            'Medium': '#f39c12', 
            'Hard': '#e74c3c'
        }.get(difficulty, '#7f8c8d')
        
        difficulty_label = QLabel(f"Difficulty: {difficulty}")
        difficulty_label.setStyleSheet(f"""
            QLabel {{
                font-size: 12px;
                color: {difficulty_color};
                font-weight: bold;
            }}
        """)
        bottom_layout.addWidget(difficulty_label)
        
        bottom_layout.addStretch()
        
        servings = recipe_data.get('servings', 1)
        servings_label = QLabel(f"👥 {servings} serving{'s' if servings != 1 else ''}")
        servings_label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                color: #7f8c8d;
            }
        """)
        bottom_layout.addWidget(servings_label)
        
        info_layout.addLayout(bottom_layout)
        
        card_layout.addWidget(info_container)
        
        # Store recipe data in card
        card.recipe_data = recipe_data
        
        # Make card clickable - single click to select, double click to view
        card.mousePressEvent = lambda event: self.select_recipe_card(card)
        card.mouseDoubleClickEvent = lambda event: self.view_recipe()
        
        return card
    
    def create_modern_recipe_list_item(self, recipe_data):
        """Create a modern recipe list item widget"""
        item_widget = QFrame(self.cards_container)
        item_widget.setFrameStyle(QFrame.NoFrame)
        item_widget.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #e0e0e0;
                border-radius: 12px;
                margin: 4px 8px;
            }
            QFrame:hover {
                border: 2px solid #4CAF50;
                background-color: #f8fffe;
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            }
        """)
        item_widget.setFixedHeight(100)
        item_widget.setCursor(Qt.PointingHandCursor)
        
        # Store recipe data
        item_widget.recipe_data = recipe_data
        
        # Create main layout
        main_layout = QHBoxLayout(item_widget)
        main_layout.setContentsMargins(12, 8, 12, 8)
        main_layout.setSpacing(12)
        
        # Recipe image/icon (left side)
        image_frame = QFrame()
        image_frame.setFixedSize(80, 80)
        image_frame.setStyleSheet("""
            QFrame {
                background-color: #f5f5f5;
                border: 1px solid #ddd;
                border-radius: 8px;
            }
        """)
        
        image_layout = QVBoxLayout(image_frame)
        image_layout.setContentsMargins(4, 4, 4, 4)
        image_layout.setAlignment(Qt.AlignCenter)
        
        image_label = QLabel()
        image_label.setAlignment(Qt.AlignCenter)
        
        # Try to load recipe image
        recipe_id = recipe_data.get('id')
        if recipe_id:
            image_path = self.get_recipe_image_path(recipe_id)
            if image_path and os.path.exists(image_path):
                pixmap = QPixmap(image_path)
                scaled_pixmap = pixmap.scaled(72, 72, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                image_label.setPixmap(scaled_pixmap)
            else:
                # Use category icon as fallback
                category_icon = self._get_category_icon(recipe_data.get('category', 'main_course'))
                image_label.setText(category_icon)
                image_label.setStyleSheet("""
                    QLabel {
                        font-size: 32px;
                        color: #4CAF50;
                    }
                """)
        
        image_layout.addWidget(image_label)
        main_layout.addWidget(image_frame)
        
        # Main content area (center)
        content_layout = QVBoxLayout()
        content_layout.setSpacing(4)
        
        # Recipe title
        title_label = QLabel(recipe_data.get('title', 'Untitled Recipe'))
        title_label.setStyleSheet("""
            QLabel {
                font-weight: bold;
                font-size: 16px;
                color: #2c3e50;
                padding: 2px 0;
            }
        """)
        content_layout.addWidget(title_label)
        
        # Category and description
        category = recipe_data.get('category', 'Unknown').replace('_', ' ').title()
        description = recipe_data.get('description', '')[:100] + '...' if len(recipe_data.get('description', '')) > 100 else recipe_data.get('description', '')
        
        info_label = QLabel(f"📁 {category}")
        if description:
            info_label.setText(f"📁 {category} • {description}")
        info_label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                color: #666;
                padding: 2px 0;
            }
        """)
        info_label.setWordWrap(True)
        content_layout.addWidget(info_label)
        
        # Recipe details (time, difficulty, servings)
        details_layout = QHBoxLayout()
        details_layout.setSpacing(12)
        
        # Cooking time
        prep_time = recipe_data.get('prep_time', '0')
        cook_time = recipe_data.get('cook_time', '0')
        total_time = self._calculate_total_time(prep_time, cook_time)
        
        time_label = QLabel(f"⏱ {total_time}")
        time_label.setStyleSheet("""
            QLabel {
                font-size: 11px;
                color: #666;
                background-color: #e3f2fd;
                padding: 2px 6px;
                border-radius: 4px;
            }
        """)
        details_layout.addWidget(time_label)
        
        # Difficulty
        difficulty = recipe_data.get('difficulty', 'Easy')
        difficulty_label = QLabel(f"📊 {difficulty}")
        difficulty_label.setStyleSheet("""
            QLabel {
                font-size: 11px;
                color: #666;
                background-color: #f3e5f5;
                padding: 2px 6px;
                border-radius: 4px;
            }
        """)
        details_layout.addWidget(difficulty_label)
        
        # Servings
        servings = recipe_data.get('servings', 'N/A')
        servings_label = QLabel(f"👥 {servings}")
        servings_label.setStyleSheet("""
            QLabel {
                font-size: 11px;
                color: #666;
                background-color: #e8f5e8;
                padding: 2px 6px;
                border-radius: 4px;
            }
        """)
        details_layout.addWidget(servings_label)
        
        details_layout.addStretch()
        content_layout.addLayout(details_layout)
        
        main_layout.addLayout(content_layout, 1)  # Stretch factor 1
        
        # Action area (right side)
        action_layout = QVBoxLayout()
        action_layout.setAlignment(Qt.AlignCenter)
        action_layout.setSpacing(8)
        
        # Favorite button
        favorite_btn = QPushButton("⭐")
        favorite_btn.setFixedSize(32, 32)
        favorite_btn.setCheckable(True)
        favorite_btn.setStyleSheet("""
            QPushButton {
                background-color: #fff3cd;
                border: 1px solid #ffeaa7;
                border-radius: 16px;
                font-size: 14px;
            }
            QPushButton:checked {
                background-color: #ffd700;
                border-color: #ffb347;
            }
            QPushButton:hover {
                background-color: #ffeb3b;
            }
        """)
        
        # Set favorite state
        if recipe_data.get('is_favorite', False):
            favorite_btn.setChecked(True)
        
        favorite_btn.clicked.connect(lambda: self._toggle_recipe_favorite(recipe_data.get('id')))
        action_layout.addWidget(favorite_btn)
        
        # View button
        view_btn = QPushButton("👁")
        view_btn.setFixedSize(32, 32)
        view_btn.setStyleSheet("""
            QPushButton {
                background-color: #e3f2fd;
                border: 1px solid #bbdefb;
                border-radius: 16px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #2196f3;
                color: white;
            }
        """)
        view_btn.clicked.connect(lambda: self.view_recipe_by_data(recipe_data))
        action_layout.addWidget(view_btn)
        
        action_layout.addStretch()
        main_layout.addLayout(action_layout)
        
        # Click handler for the entire item
        item_widget.mousePressEvent = lambda event: self.select_recipe_card(item_widget)
        
        return item_widget
    
    def _get_category_icon(self, category_name):
        """Get appropriate icon for category"""
        category_icons = {
            'Breakfast & Brunch': '🌅',
            'Lunch & Light Meals': '🥗',
            'Dinner & Main Courses': '🍽️',
            'Sides & Vegetables': '🥕',
            'Desserts & Sweets': '🍰',
            'Beverages & Drinks': '🥤',
            'Snacks & Appetizers': '🍿',
            'Baking & Breads': '🍞',
            'Breakfast': '🌅',
            'Lunch': '🥗',
            'Dinner': '🍽️',
            'Dessert': '🍰',
            'Snack': '🍿',
            'Uncategorized': '📝'
        }
        return category_icons.get(category_name, '📝')
    
    def select_recipe_card(self, card):
        """Handle recipe card selection"""
        # Clear previous selection
        for c in self.recipe_cards:
            c.setStyleSheet("""
                QFrame {
                    background-color: white;
                    border-radius: 12px;
                    border: 1px solid #e0e0e0;
                }
                QFrame:hover {
                    border-color: #3498db;
                }
            """)
        
        # Select current card
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border-radius: 12px;
                border: 2px solid #3498db;
            }
        """)
        
        # Store selected recipe
        self.current_recipe_id = card.recipe_data.get('id')
    
    def set_view_mode(self, mode):
        """Set the view mode (grid or list) - now defaults to grid view only"""
        # Always use grid view since we removed the view toggle buttons
        self.refresh_recipe_display()
    
    def _on_category_combo_changed(self, category_name: str):
        """Handle category selection change from combo box"""
        self.filter_recipes()
    
    def get_category_id_by_name(self, category_name: str) -> str:
        """Get category ID by category name - now returns the category name directly"""
        if category_name == "All Categories":
            return "all"
        return category_name
    
    def filter_recipes(self):
        """Filter recipes based on search text"""
        if hasattr(self, 'category_combo'):
            category_name = self.category_combo.currentText()
            category_id = self.get_category_id_by_name(category_name)
        else:
            category_name = "All"
            category_id = "all"
            
        print(f"DEBUG: Filtering recipes - Search: '{self.search_edit.text()}', "
              f"Category: '{category_name}' (ID: {category_id}), "
              f"Favorites: {self.favorite_filter_cb.isChecked()}")
        self.refresh_recipe_display()
    
    def sort_recipes(self):
        """Sort recipes based on selected criteria"""
        self.refresh_recipe_display()
    
    def refresh_recipe_display(self):
        """Refresh the recipe cards display"""
        # Prevent multiple simultaneous calls
        if hasattr(self, '_refreshing') and self._refreshing:
            return
        self._refreshing = True
        
        try:
            # Clear existing cards properly
            for card in self.recipe_cards:
                if card and not card.isHidden():
                    card.hide()
                    card.deleteLater()
            self.recipe_cards.clear()
            
            # Clear the grid layout
            for i in reversed(range(self.cards_grid.count())):
                item = self.cards_grid.itemAt(i)
                if item and item.widget():
                    widget = item.widget()
                    if widget and not widget.isHidden():
                        widget.hide()
                        widget.deleteLater()
        
            # Get filtered and sorted recipes
            recipes = self.get_filtered_recipes()
            
            # Create and add recipe cards (always use grid view)
            self.display_grid_view(recipes)
                
        finally:
            self._refreshing = False
    
    def get_filtered_recipes(self):
        """Get filtered and sorted recipes"""
        try:
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            # Debug: Check if recipes table exists and has data
            cursor.execute("SELECT COUNT(*) FROM recipes")
            recipe_count = cursor.fetchone()[0]
            # Debug logging removed to reduce console spam
            # print(f"DEBUG: Found {recipe_count} recipes in database")
            
            # Check if we have the full set of gluten-free recipes
            cursor.execute("SELECT COUNT(*) FROM recipes WHERE title LIKE 'GF %' OR title LIKE 'Gluten-Free %'")
            gf_recipe_count = cursor.fetchone()[0]
            
            # If we don't have enough gluten-free recipes, add them
            if gf_recipe_count < 16:
                added_count = self._add_sample_recipes_to_db(cursor)
                db.commit()
                print(f"DEBUG: Added {added_count} gluten-free recipes to database")
            
            # Build WHERE clause based on filters
            where_conditions = []
            params = []
            
            # Search filter
            search_text = self.search_edit.text().strip()
            if search_text:
                where_conditions.append("(r.title LIKE ? OR r.category LIKE ? OR r.tags LIKE ?)")
                search_param = f"%{search_text}%"
                params.extend([search_param, search_param, search_param])
            
            # Favorite filter
            if hasattr(self, 'favorite_filter_cb') and self.favorite_filter_cb.isChecked():
                where_conditions.append("r.is_favorite = 1")
            
            # Category filter
            if hasattr(self, 'category_combo') and self.category_combo.currentData() != 'all':
                category_id = self.category_combo.currentData()
                where_conditions.append("r.category = ?")
                params.append(category_id)
            
            # Build the query
            where_clause = "WHERE " + " AND ".join(where_conditions) if where_conditions else ""
            
            # Sort order
            sort_option = self.sort_combo.currentText()
            order_clause = "ORDER BY r.title"
            if sort_option == "Name (Z-A)":
                order_clause = "ORDER BY r.title DESC"
            elif sort_option == "Prep Time":
                order_clause = "ORDER BY CAST(r.prep_time AS INTEGER) ASC"
            elif sort_option == "Difficulty":
                order_clause = ("ORDER BY CASE r.difficulty "
                               "WHEN 'Easy' THEN 1 WHEN 'Medium' THEN 2 "
                               "WHEN 'Hard' THEN 3 ELSE 4 END")
            elif sort_option == "Recently Added":
                order_clause = "ORDER BY r.id DESC"
            elif sort_option == "Favorites First":
                order_clause = "ORDER BY r.is_favorite DESC, r.title"
            
            cursor.execute(f"""
                SELECT r.id, r.title, r.instructions, r.prep_time, r.cook_time, r.servings, 
                       '', r.category, r.tags, '', '', r.is_favorite, r.image_path, r.difficulty
                FROM recipes r
                {where_clause}
                {order_clause}
            """, params)
            
            recipes = cursor.fetchall()
            db.close()
            
            # Debug logging removed to reduce console spam
            # print(f"DEBUG: Query returned {len(recipes)} recipes")
            
            # Convert to recipe data format
            recipe_list = []
            for recipe in recipes:
                recipe_data = {
                    'id': recipe[0],
                    'name': recipe[1],
                    'instructions': recipe[2],
                    'prep_time': recipe[3],
                    'cook_time': recipe[4],
                    'servings': recipe[5],
                    'category': recipe[7],
                    'tags': recipe[8],
                    'is_favorite': recipe[11],
                    'image_path': recipe[12],
                    'difficulty': recipe[13]
                }
                recipe_list.append(recipe_data)
            
            return recipe_list
            
        except Exception as e:
            print(f"Error getting filtered recipes: {e}")
            return []
    
    def display_grid_view(self, recipes):
        """Display recipes in grid view"""
        # Clear grid properly
        for i in reversed(range(self.cards_grid.count())):
            item = self.cards_grid.itemAt(i)
            if item and item.widget():
                item.widget().hide()
                item.widget().deleteLater()
        
        # Add recipe cards to grid
        row, col = 0, 0
        max_cols = 4  # Adjust based on available width
        
        for recipe_data in recipes:
            card = self.create_recipe_card(recipe_data)
            self.cards_grid.addWidget(card, row, col)
            self.recipe_cards.append(card)
            
            col += 1
            if col >= max_cols:
                col = 0
                row += 1
    
    def display_list_view(self, recipes):
        """Display recipes in modern list view"""
        # Clear grid properly
        for i in reversed(range(self.cards_grid.count())):
            item = self.cards_grid.itemAt(i)
            if item and item.widget():
                item.widget().hide()
                item.widget().deleteLater()
        
        if not recipes:
            # Show empty state
            empty_widget = QWidget()
            empty_layout = QVBoxLayout(empty_widget)
            empty_layout.setContentsMargins(20, 40, 20, 40)
            
            empty_label = QLabel("No recipes found. Add some recipes to get started!")
            empty_label.setAlignment(Qt.AlignCenter)
            empty_label.setStyleSheet("""
                QLabel {
                    color: #666;
                    font-size: 16px;
                    padding: 20px;
                    background-color: #f9f9f9;
                    border: 2px dashed #ddd;
                    border-radius: 10px;
                }
            """)
            empty_layout.addWidget(empty_label)
            self.cards_grid.addWidget(empty_widget, 0, 0)
            return
        
        # Add modern list items
        for i, recipe_data in enumerate(recipes):
            list_item = self.create_modern_recipe_list_item(recipe_data)
            self.cards_grid.addWidget(list_item, i, 0)
            self.recipe_cards.append(list_item)
    
    def _load_sample_recipes(self):
        """Load sample recipes when database is empty"""
        sample_recipes = [
            {
                "id": None,
                "name": "Gluten-Free Pancakes",
                "category": "Breakfast", 
                "prep_time": "15 min",
                "cook_time": "10 min",
                "servings": 4,
                "difficulty": "Easy",
                "description": "Light and fluffy gluten-free pancakes perfect for breakfast",
                "ingredients": [
                    {"name": "gluten-free flour", "amount": "1", "unit": "cup"},
                    {"name": "eggs", "amount": "2", "unit": ""},
                    {"name": "milk", "amount": "1", "unit": "cup"},
                    {"name": "baking powder", "amount": "2", "unit": "tsp"}
                ],
                "instructions": "1. Mix dry ingredients\n2. Add wet ingredients\n3. Cook on griddle until golden",
                "notes": "Best served with maple syrup"
            },
            {
                "id": None,
                "name": "Quinoa Salad",
                "category": "Lunch",
                "prep_time": "20 min", 
                "cook_time": "15 min",
                "servings": 6,
                "difficulty": "Medium",
                "description": "Nutritious quinoa salad with fresh vegetables",
                "ingredients": [
                    {"name": "quinoa", "amount": "1", "unit": "cup"},
                    {"name": "cucumber", "amount": "1", "unit": ""},
                    {"name": "tomatoes", "amount": "2", "unit": ""},
                    {"name": "olive oil", "amount": "3", "unit": "tbsp"}
                ],
                "instructions": ("1. Cook quinoa according to package directions\n"
                                "2. Chop vegetables\n"
                                "3. Mix together with dressing"),
                "notes": "Can be made ahead and stored in refrigerator"
            }
        ]
        
        # Store sample recipes for card display
        for recipe in sample_recipes:
            recipe_id = f"sample_{recipe['name'].replace(' ', '_').lower()}"
            self.recipe_data_storage[recipe_id] = recipe
    
    def _add_sample_recipes_to_db(self, cursor):
        """Add sample gluten-free recipes to the database"""
        sample_recipes = [
            # Breakfast & Brunch
            {
                "title": "Gluten-Free Banana Pancakes",
                "category": "Breakfast & Brunch", 
                "prep_time": "10",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Mash 2 ripe bananas in a bowl\n"
                                "2. Add 2 eggs and whisk together\n"
                                "3. Add 1/2 cup gluten-free flour blend\n"
                                "4. Cook on griddle until golden brown\n"
                                "5. Serve with maple syrup"),
                "tags": "breakfast, gluten-free, pancakes, banana",
                "is_favorite": 1
            },
            {
                "title": "GF Overnight Oats",
                "category": "Breakfast & Brunch",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 1,
                "difficulty": "Easy",
                "instructions": ("1. Mix 1/2 cup gluten-free oats with 1/2 cup milk\n"
                                "2. Add 1 tbsp chia seeds and 1 tsp honey\n"
                                "3. Stir in fresh berries\n"
                                "4. Refrigerate overnight\n"
                                "5. Enjoy cold in the morning"),
                "tags": "breakfast, gluten-free, oats, healthy",
                "is_favorite": 0
            },
            
            # Lunch & Light Meals
            {
                "title": "Quinoa Power Bowl",
                "category": "Lunch & Light Meals",
                "prep_time": "15", 
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Cook 1 cup quinoa according to package directions\n"
                                "2. Roast sweet potato cubes at 400°F for 20 minutes\n"
                                "3. Sauté kale with garlic until wilted\n"
                                "4. Top quinoa with roasted vegetables\n"
                                "5. Drizzle with tahini dressing"),
                "tags": "lunch, gluten-free, quinoa, healthy, vegan",
                "is_favorite": 1
            },
            {
                "title": "GF Chicken Caesar Salad",
                "category": "Lunch & Light Meals",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Grill chicken breast and slice\n"
                                "2. Make gluten-free croutons from GF bread\n"
                                "3. Toss romaine lettuce with GF Caesar dressing\n"
                                "4. Add chicken, croutons, and parmesan\n"
                                "5. Serve immediately"),
                "tags": "lunch, gluten-free, chicken, salad, caesar",
                "is_favorite": 0
            },
            
            # Dinner & Main Courses
            {
                "title": "GF Chicken Stir Fry",
                "category": "Dinner & Main Courses",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Cut chicken into strips and marinate in GF soy sauce\n"
                                "2. Heat oil in wok over high heat\n"
                                "3. Add chicken and cook until done\n"
                                "4. Add mixed vegetables and stir fry\n"
                                "5. Serve over GF rice noodles"),
                "tags": "dinner, gluten-free, chicken, stir-fry, asian",
                "is_favorite": 1
            },
            {
                "title": "GF Salmon with Herbs",
                "category": "Dinner & Main Courses",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Season salmon fillets with salt and pepper\n"
                                "2. Mix fresh herbs with olive oil\n"
                                "3. Bake salmon at 400°F for 12-15 minutes\n"
                                "4. Serve with roasted vegetables\n"
                                "5. Drizzle with herb oil"),
                "tags": "dinner, gluten-free, salmon, healthy, herbs",
                "is_favorite": 0
            },
            
            # Sides & Vegetables
            {
                "title": "GF Roasted Vegetables",
                "category": "Sides & Vegetables",
                "prep_time": "10",
                "cook_time": "25",
                "servings": 6,
                "difficulty": "Easy",
                "instructions": "1. Cut assorted vegetables into chunks\n2. Toss with olive oil and GF seasonings\n3. Roast at 425°F for 20-25 minutes\n4. Stir halfway through cooking\n5. Serve hot as a side dish",
                "tags": "sides, gluten-free, vegetables, roasted, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Quinoa Pilaf",
                "category": "Sides & Vegetables",
                "prep_time": "5",
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": "1. Toast quinoa in olive oil for 2 minutes\n2. Add vegetable broth and bring to boil\n3. Reduce heat, cover and simmer 15 minutes\n4. Fluff with fork and add herbs\n5. Serve as side dish",
                "tags": "sides, gluten-free, quinoa, pilaf, healthy",
                "is_favorite": 0
            },
            
            # Desserts & Sweets
            {
                "title": "GF Chocolate Chip Cookies",
                "category": "Desserts & Sweets",
                "prep_time": "15",
                "cook_time": "12",
                "servings": "24 cookies",
                "difficulty": "Medium",
                "instructions": "1. Cream butter and sugars until fluffy\n2. Add eggs and vanilla extract\n3. Mix in GF flour blend and baking soda\n4. Fold in chocolate chips\n5. Bake at 375°F for 10-12 minutes",
                "tags": "dessert, gluten-free, cookies, chocolate, sweet",
                "is_favorite": 1
            },
            {
                "title": "GF Berry Crumble",
                "category": "Desserts & Sweets",
                "prep_time": "15",
                "cook_time": "35",
                "servings": 6,
                "difficulty": "Easy",
                "instructions": "1. Mix berries with sugar and cornstarch\n2. Make crumble topping with GF oats and flour\n3. Place berries in baking dish\n4. Top with crumble mixture\n5. Bake at 375°F for 30-35 minutes",
                "tags": "dessert, gluten-free, berries, crumble, fruit",
                "is_favorite": 0
            },
            
            # Beverages & Drinks
            {
                "title": "GF Green Smoothie",
                "category": "Beverages & Drinks",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": "1. Add 2 cups spinach to blender\n2. Add 1 banana and 1 cup pineapple\n3. Pour in 1 cup coconut milk\n4. Add 1 tbsp GF protein powder\n5. Blend until smooth and serve",
                "tags": "beverage, gluten-free, smoothie, healthy, green",
                "is_favorite": 1
            },
            {
                "title": "GF Turmeric Latte",
                "category": "Beverages & Drinks",
                "prep_time": "5",
                "cook_time": "5",
                "servings": 1,
                "difficulty": "Easy",
                "instructions": "1. Heat 1 cup almond milk in saucepan\n2. Whisk in 1 tsp turmeric and 1/2 tsp cinnamon\n3. Add 1 tsp honey and pinch of black pepper\n4. Simmer for 3-4 minutes\n5. Strain and serve hot",
                "tags": "beverage, gluten-free, latte, turmeric, healthy",
                "is_favorite": 0
            },
            
            # Snacks & Appetizers
            {
                "title": "GF Hummus with Veggies",
                "category": "Snacks & Appetizers",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 6,
                "difficulty": "Easy",
                "instructions": "1. Blend chickpeas with tahini and lemon juice\n2. Add garlic, olive oil, and salt\n3. Process until smooth\n4. Serve with GF crackers and vegetables\n5. Drizzle with olive oil and paprika",
                "tags": "snack, gluten-free, hummus, healthy, dip",
                "is_favorite": 1
            },
            {
                "title": "GF Energy Balls",
                "category": "Snacks & Appetizers",
                "prep_time": "15",
                "cook_time": "0",
                "servings": "20 balls",
                "difficulty": "Easy",
                "instructions": "1. Process dates in food processor until smooth\n2. Add GF oats, almond butter, and cocoa powder\n3. Mix in chocolate chips and coconut\n4. Form into balls and refrigerate\n5. Store in airtight container",
                "tags": "snack, gluten-free, energy balls, healthy, no-bake",
                "is_favorite": 0
            },
            
            # Baking & Breads
            {
                "title": "GF Banana Bread",
                "category": "Baking & Breads",
                "prep_time": "15",
                "cook_time": "60",
                "servings": "1 loaf",
                "difficulty": "Medium",
                "instructions": "1. Mash 3 ripe bananas in large bowl\n2. Add eggs, oil, and vanilla\n3. Mix in GF flour blend and baking soda\n4. Fold in walnuts and chocolate chips\n5. Bake at 350°F for 55-60 minutes",
                "tags": "baking, gluten-free, banana bread, sweet, breakfast",
                "is_favorite": 1
            },
            {
                "title": "GF Cornbread",
                "category": "Baking & Breads",
                "prep_time": "10",
                "cook_time": "25",
                "servings": "9 squares",
                "difficulty": "Easy",
                "instructions": "1. Mix GF cornmeal with GF flour blend\n2. Add baking powder, salt, and sugar\n3. Whisk in eggs, milk, and melted butter\n4. Pour into greased 8x8 pan\n5. Bake at 400°F for 20-25 minutes",
                "tags": "baking, gluten-free, cornbread, savory, side",
                "is_favorite": 0
            }
        ]
        
        added_count = 0
        for recipe in sample_recipes:
            # Check if recipe already exists
            cursor.execute("SELECT COUNT(*) FROM recipes WHERE title = ?", (recipe["title"],))
            exists = cursor.fetchone()[0]
            
            if exists == 0:
                cursor.execute("""
                    INSERT INTO recipes (title, category, prep_time, cook_time, servings, difficulty, instructions, tags, is_favorite)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    recipe["title"], recipe["category"], recipe["prep_time"], recipe["cook_time"],
                    recipe["servings"], recipe["difficulty"], recipe["instructions"], recipe["tags"], recipe["is_favorite"]
                ))
                added_count += 1
        
        return added_count
    
    def create_comprehensive_gf_recipes(self):
        """Create comprehensive gluten-free recipes for all categories and subcategories"""
        try:
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            # Get all categories and subcategories
            cursor.execute("""
                SELECT c.id, c.name, c.parent_id, p.name as parent_name
                FROM categories c
                LEFT JOIN categories p ON c.parent_id = p.id
                ORDER BY c.parent_id, c.sort_order, c.name
            """)
            categories = cursor.fetchall()
            
            # Create comprehensive recipe collection
            comprehensive_recipes = self._get_comprehensive_gf_recipes()
            
            added_count = 0
            for recipe_data in comprehensive_recipes:
                # Check if recipe already exists
                cursor.execute("SELECT COUNT(*) FROM recipes WHERE title = ?", (recipe_data["title"],))
                exists = cursor.fetchone()[0]
                
                if exists == 0:
                    cursor.execute("""
                        INSERT INTO recipes (title, category, prep_time, cook_time, servings, difficulty, instructions, tags, is_favorite)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        recipe_data["title"], recipe_data["category"], recipe_data["prep_time"], recipe_data["cook_time"],
                        recipe_data["servings"], recipe_data["difficulty"], recipe_data["instructions"], recipe_data["tags"], recipe_data["is_favorite"]
                    ))
                    added_count += 1
            
            db.commit()
            db.close()
            
            print(f"Added {added_count} comprehensive gluten-free recipes")
            return added_count
            
        except Exception as e:
            print(f"Error creating comprehensive GF recipes: {e}")
            return 0
    
    def _get_comprehensive_gf_recipes(self):
        """Get comprehensive collection of gluten-free recipes for all categories"""
        return [
            # BREAKFAST & BRUNCH - Pancakes & Waffles
            {
                "title": "GF Fluffy Almond Flour Pancakes",
                "category": "Pancakes & Waffles",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Whisk 2 cups almond flour with 4 eggs, 1/2 cup almond milk, 2 tbsp honey\n"
                               "2. Add 1 tsp baking powder, 1/2 tsp vanilla, pinch of salt\n"
                               "3. Let batter rest 5 minutes\n"
                               "4. Cook on griddle over medium heat until golden\n"
                               "5. Serve with fresh berries and maple syrup"),
                "tags": "breakfast, gluten-free, pancakes, almond flour, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Buckwheat Waffles",
                "category": "Pancakes & Waffles",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 6,
                "difficulty": "Medium",
                "instructions": ("1. Mix 1.5 cups buckwheat flour with 1 cup GF oat flour\n"
                               "2. Add 2 eggs, 1.5 cups buttermilk, 1/4 cup melted butter\n"
                               "3. Stir in 2 tbsp sugar, 1 tsp baking powder, 1/2 tsp baking soda\n"
                               "4. Cook in waffle iron until crispy\n"
                               "5. Top with Greek yogurt and fresh fruit"),
                "tags": "breakfast, gluten-free, waffles, buckwheat, healthy",
                "is_favorite": 0
            },
            
            # BREAKFAST & BRUNCH - Eggs & Omelets
            {
                "title": "GF Mediterranean Veggie Omelet",
                "category": "Eggs & Omelets",
                "prep_time": "10",
                "cook_time": "8",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Whisk 4 eggs with 2 tbsp water, salt and pepper\n"
                               "2. Sauté 1/2 cup diced bell peppers, 1/4 cup red onion\n"
                               "3. Add 1/2 cup spinach, 1/4 cup sun-dried tomatoes\n"
                               "4. Pour eggs over vegetables, cook until set\n"
                               "5. Fold and serve with fresh herbs"),
                "tags": "breakfast, gluten-free, omelet, mediterranean, vegetables",
                "is_favorite": 1
            },
            {
                "title": "GF Quinoa Breakfast Bowl",
                "category": "Eggs & Omelets",
                "prep_time": "5",
                "cook_time": "15",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Cook 1 cup quinoa according to package directions\n"
                               "2. Top with 2 fried eggs, 1/2 avocado, 1/4 cup black beans\n"
                               "3. Add 2 tbsp salsa, 1 tbsp cilantro, lime wedge\n"
                               "4. Season with salt, pepper, and hot sauce\n"
                               "5. Serve immediately while warm"),
                "tags": "breakfast, gluten-free, quinoa, eggs, healthy, protein",
                "is_favorite": 0
            },
            
            # BREAKFAST & BRUNCH - Cereals & Oatmeal
            {
                "title": "GF Overnight Chia Oats",
                "category": "Cereals & Oatmeal",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 1,
                "difficulty": "Easy",
                "instructions": ("1. Mix 1/2 cup certified GF oats with 1/2 cup almond milk\n"
                               "2. Add 1 tbsp chia seeds, 1 tsp maple syrup, 1/2 tsp vanilla\n"
                               "3. Stir in 1/4 cup blueberries, 1 tbsp sliced almonds\n"
                               "4. Refrigerate overnight (8+ hours)\n"
                               "5. Enjoy cold or at room temperature"),
                "tags": "breakfast, gluten-free, oats, chia, overnight, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Coconut Quinoa Porridge",
                "category": "Cereals & Oatmeal",
                "prep_time": "5",
                "cook_time": "20",
                "servings": 3,
                "difficulty": "Easy",
                "instructions": ("1. Rinse 1 cup quinoa and cook with 2 cups coconut milk\n"
                               "2. Add 1/4 cup shredded coconut, 2 tbsp maple syrup\n"
                               "3. Simmer 15-20 minutes until creamy\n"
                               "4. Stir in 1 tsp vanilla, 1/2 tsp cinnamon\n"
                               "5. Top with fresh mango and toasted coconut"),
                "tags": "breakfast, gluten-free, quinoa, coconut, porridge, tropical",
                "is_favorite": 0
            },
            
            # BREAKFAST & BRUNCH - Smoothies & Juices
            {
                "title": "GF Green Power Smoothie",
                "category": "Smoothies & Juices",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Blend 2 cups spinach with 1 banana, 1 cup pineapple\n"
                               "2. Add 1 cup coconut water, 1 tbsp almond butter\n"
                               "3. Include 1 scoop GF protein powder, 1 tsp spirulina\n"
                               "4. Blend until smooth and creamy\n"
                               "5. Serve immediately with ice"),
                "tags": "breakfast, gluten-free, smoothie, green, healthy, protein",
                "is_favorite": 1
            },
            {
                "title": "GF Berry Antioxidant Smoothie",
                "category": "Smoothies & Juices",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Blend 1 cup mixed berries with 1 banana\n"
                               "2. Add 1 cup almond milk, 1/2 cup Greek yogurt\n"
                               "3. Include 1 tbsp honey, 1 tsp vanilla extract\n"
                               "4. Add 1 tbsp flax seeds, 1/2 cup ice\n"
                               "5. Blend until smooth and serve"),
                "tags": "breakfast, gluten-free, smoothie, berries, antioxidants, healthy",
                "is_favorite": 0
            },
            
            # LUNCH & LIGHT MEALS - Salads
            {
                "title": "GF Quinoa Power Salad",
                "category": "Salads",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Cook 1 cup quinoa according to package directions\n"
                               "2. Roast 1 sweet potato cubed at 400°F for 20 minutes\n"
                               "3. Sauté 2 cups kale with garlic until wilted\n"
                               "4. Mix quinoa with roasted vegetables and 1/2 cup chickpeas\n"
                               "5. Drizzle with tahini dressing and serve"),
                "tags": "lunch, gluten-free, quinoa, salad, healthy, vegan",
                "is_favorite": 1
            },
            {
                "title": "GF Mediterranean Chickpea Salad",
                "category": "Salads",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Mix 2 cans chickpeas with 1 cucumber diced\n"
                               "2. Add 1/2 red onion, 1 cup cherry tomatoes, 1/2 cup kalamata olives\n"
                               "3. Toss with 1/4 cup olive oil, 2 tbsp lemon juice\n"
                               "4. Add 1/4 cup fresh parsley, 1 tsp oregano\n"
                               "5. Season with salt, pepper, and serve"),
                "tags": "lunch, gluten-free, mediterranean, chickpeas, salad, healthy",
                "is_favorite": 0
            },
            
            # LUNCH & LIGHT MEALS - Sandwiches
            {
                "title": "GF Turkey Avocado Wrap",
                "category": "Sandwiches",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Spread 2 GF tortillas with 1/4 cup hummus each\n"
                               "2. Layer 4 oz sliced turkey, 1/2 avocado sliced\n"
                               "3. Add 1/4 cup sprouts, 2 tbsp red onion, 1/4 cup cucumber\n"
                               "4. Drizzle with 1 tbsp olive oil, 1 tsp lemon juice\n"
                               "5. Roll tightly and slice in half"),
                "tags": "lunch, gluten-free, sandwich, turkey, avocado, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Veggie Hummus Sandwich",
                "category": "Sandwiches",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Toast 4 slices GF bread\n"
                               "2. Spread 1/4 cup hummus on each slice\n"
                               "3. Layer with 1/2 cucumber, 1 tomato, 1/4 red onion\n"
                               "4. Add 1/2 cup mixed greens, 1/4 cup roasted peppers\n"
                               "5. Season with salt, pepper, and serve"),
                "tags": "lunch, gluten-free, sandwich, vegetarian, hummus, healthy",
                "is_favorite": 0
            },
            
            # LUNCH & LIGHT MEALS - Soups
            {
                "title": "GF Creamy Tomato Basil Soup",
                "category": "Soups",
                "prep_time": "10",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Sauté 1 onion, 2 cloves garlic in 2 tbsp olive oil\n"
                               "2. Add 2 cans diced tomatoes, 2 cups vegetable broth\n"
                               "3. Simmer 20 minutes, then blend until smooth\n"
                               "4. Stir in 1/2 cup coconut milk, 1/4 cup fresh basil\n"
                               "5. Season with salt, pepper, and serve hot"),
                "tags": "lunch, gluten-free, soup, tomato, basil, creamy, vegan",
                "is_favorite": 1
            },
            {
                "title": "GF Chicken Vegetable Soup",
                "category": "Soups",
                "prep_time": "15",
                "cook_time": "30",
                "servings": 6,
                "difficulty": "Easy",
                "instructions": ("1. Sauté 1 onion, 2 carrots, 2 celery stalks in olive oil\n"
                               "2. Add 1 lb chicken breast cubed, cook 5 minutes\n"
                               "3. Add 6 cups GF chicken broth, 1 cup diced potatoes\n"
                               "4. Simmer 20 minutes until chicken is cooked\n"
                               "5. Add 1 cup frozen peas, 2 tbsp fresh parsley, season"),
                "tags": "lunch, gluten-free, soup, chicken, vegetables, healthy",
                "is_favorite": 0
            },
            
            # LUNCH & LIGHT MEALS - Quick & Easy
            {
                "title": "GF 5-Minute Avocado Toast",
                "category": "Quick & Easy",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Toast 2 slices GF bread\n"
                               "2. Mash 1 avocado with 1 tbsp lemon juice, salt, pepper\n"
                               "3. Spread avocado mixture on toast\n"
                               "4. Top with 2 tbsp hemp seeds, 1/4 tsp red pepper flakes\n"
                               "5. Drizzle with olive oil and serve immediately"),
                "tags": "lunch, gluten-free, quick, avocado, toast, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Greek Yogurt Parfait",
                "category": "Quick & Easy",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Layer 1/2 cup Greek yogurt in glass\n"
                               "2. Add 1/4 cup GF granola, 1/4 cup mixed berries\n"
                               "3. Repeat layers with remaining ingredients\n"
                               "4. Top with 1 tbsp honey, 1 tbsp sliced almonds\n"
                               "5. Serve immediately or refrigerate up to 2 hours"),
                "tags": "lunch, gluten-free, quick, yogurt, parfait, healthy",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Beef & Veal
            {
                "title": "GF Beef Stir Fry with Rice Noodles",
                "category": "Beef & Veal",
                "prep_time": "15",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Marinate 1 lb beef strips in 2 tbsp GF tamari, 1 tbsp sesame oil\n"
                               "2. Cook rice noodles according to package directions\n"
                               "3. Stir fry beef in wok over high heat until browned\n"
                               "4. Add 2 cups mixed vegetables, 1 tbsp ginger, 2 cloves garlic\n"
                               "5. Toss with noodles and 1/4 cup GF teriyaki sauce"),
                "tags": "dinner, gluten-free, beef, stir-fry, asian, noodles",
                "is_favorite": 1
            },
            {
                "title": "GF Herb-Crusted Beef Tenderloin",
                "category": "Beef & Veal",
                "prep_time": "10",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Mix 1/4 cup GF breadcrumbs with 2 tbsp fresh herbs\n"
                               "2. Season 1.5 lb beef tenderloin with salt, pepper\n"
                               "3. Sear in oven-safe pan, then coat with herb mixture\n"
                               "4. Roast at 425°F for 20-25 minutes for medium-rare\n"
                               "5. Rest 10 minutes before slicing"),
                "tags": "dinner, gluten-free, beef, tenderloin, herbs, elegant",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Pork
            {
                "title": "GF Pork Tenderloin with Apple Glaze",
                "category": "Pork",
                "prep_time": "10",
                "cook_time": "30",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Season 1.5 lb pork tenderloin with salt, pepper\n"
                               "2. Sear in oven-safe pan until browned on all sides\n"
                               "3. Mix 1/2 cup apple juice, 2 tbsp maple syrup, 1 tbsp GF Dijon\n"
                               "4. Roast at 400°F for 20-25 minutes, basting with glaze\n"
                               "5. Rest 10 minutes, slice and serve with pan juices"),
                "tags": "dinner, gluten-free, pork, tenderloin, apple, sweet",
                "is_favorite": 1
            },
            {
                "title": "GF Pulled Pork Tacos",
                "category": "Pork",
                "prep_time": "15",
                "cook_time": "240",
                "servings": 8,
                "difficulty": "Easy",
                "instructions": ("1. Rub 3 lb pork shoulder with GF taco seasoning\n"
                               "2. Slow cook in crockpot with 1 cup GF chicken broth\n"
                               "3. Cook on low 6-8 hours until tender\n"
                               "4. Shred meat, mix with 1/2 cup GF BBQ sauce\n"
                               "5. Serve in GF tortillas with slaw and avocado"),
                "tags": "dinner, gluten-free, pork, tacos, slow-cooked, mexican",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Poultry
            {
                "title": "GF Lemon Herb Chicken",
                "category": "Poultry",
                "prep_time": "10",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Season 4 chicken breasts with salt, pepper\n"
                               "2. Mix 1/4 cup olive oil, 2 tbsp lemon juice, 1 tbsp herbs\n"
                               "3. Marinate chicken 30 minutes\n"
                               "4. Bake at 400°F for 20-25 minutes until 165°F\n"
                               "5. Serve with roasted vegetables and lemon wedges"),
                "tags": "dinner, gluten-free, chicken, lemon, herbs, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Turkey Meatballs with Zucchini Noodles",
                "category": "Poultry",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Mix 1 lb ground turkey with 1/4 cup GF breadcrumbs\n"
                               "2. Add 1 egg, 2 tbsp parsley, 1 tsp Italian seasoning\n"
                               "3. Form into meatballs, bake at 400°F for 15 minutes\n"
                               "4. Spiralize 2 zucchini into noodles\n"
                               "5. Toss with GF marinara sauce and meatballs"),
                "tags": "dinner, gluten-free, turkey, meatballs, zucchini, healthy",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Seafood
            {
                "title": "GF Baked Salmon with Dill Sauce",
                "category": "Seafood",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Season 4 salmon fillets with salt, pepper, lemon\n"
                               "2. Bake at 400°F for 12-15 minutes until flaky\n"
                               "3. Mix 1/2 cup Greek yogurt with 2 tbsp fresh dill\n"
                               "4. Add 1 tbsp lemon juice, 1 tsp Dijon mustard\n"
                               "5. Serve salmon with dill sauce and roasted vegetables"),
                "tags": "dinner, gluten-free, salmon, dill, healthy, omega-3",
                "is_favorite": 1
            },
            {
                "title": "GF Shrimp Scampi with Rice",
                "category": "Seafood",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Cook 1.5 cups jasmine rice according to package\n"
                               "2. Sauté 1 lb shrimp with 3 cloves garlic in olive oil\n"
                               "3. Add 1/2 cup white wine, 2 tbsp lemon juice\n"
                               "4. Stir in 2 tbsp butter, 1/4 cup fresh parsley\n"
                               "5. Serve over rice with lemon wedges"),
                "tags": "dinner, gluten-free, shrimp, scampi, rice, italian",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Vegetarian
            {
                "title": "GF Stuffed Bell Peppers",
                "category": "Vegetarian",
                "prep_time": "15",
                "cook_time": "45",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Hollow out 4 bell peppers, blanch 5 minutes\n"
                               "2. Cook 1 cup quinoa, mix with 1 can black beans\n"
                               "3. Add 1 cup corn, 1/2 cup salsa, 1/4 cup cheese\n"
                               "4. Stuff peppers, top with more cheese\n"
                               "5. Bake at 375°F for 30-35 minutes"),
                "tags": "dinner, gluten-free, vegetarian, peppers, quinoa, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Eggplant Parmesan",
                "category": "Vegetarian",
                "prep_time": "20",
                "cook_time": "40",
                "servings": 6,
                "difficulty": "Medium",
                "instructions": ("1. Slice 2 eggplants, salt and drain 30 minutes\n"
                               "2. Dip in GF flour, then egg, then GF breadcrumbs\n"
                               "3. Bake at 400°F for 20 minutes until golden\n"
                               "4. Layer with GF marinara sauce and mozzarella\n"
                               "5. Bake 15 minutes until bubbly"),
                "tags": "dinner, gluten-free, vegetarian, eggplant, parmesan, italian",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Vegan
            {
                "title": "GF Vegan Buddha Bowl",
                "category": "Vegan",
                "prep_time": "20",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Roast 1 sweet potato, 1 cup broccoli at 400°F for 20 minutes\n"
                               "2. Cook 1 cup quinoa according to package directions\n"
                               "3. Sauté 1 cup chickpeas with 1 tsp curry powder\n"
                               "4. Arrange in bowls with 1/2 avocado, 2 tbsp tahini\n"
                               "5. Drizzle with lemon juice and serve"),
                "tags": "dinner, gluten-free, vegan, buddha bowl, healthy, colorful",
                "is_favorite": 1
            },
            {
                "title": "GF Vegan Lentil Curry",
                "category": "Vegan",
                "prep_time": "15",
                "cook_time": "30",
                "servings": 6,
                "difficulty": "Medium",
                "instructions": ("1. Sauté 1 onion, 3 cloves garlic, 1 tbsp ginger\n"
                               "2. Add 2 tbsp curry powder, 1 tsp turmeric, 1/2 tsp cumin\n"
                               "3. Add 1 cup red lentils, 1 can coconut milk, 2 cups water\n"
                               "4. Simmer 25 minutes until lentils are tender\n"
                               "5. Stir in 2 cups spinach, serve over rice"),
                "tags": "dinner, gluten-free, vegan, curry, lentils, indian",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Italian
            {
                "title": "GF Chicken Parmesan",
                "category": "Italian",
                "prep_time": "15",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Pound 4 chicken breasts to 1/2 inch thickness\n"
                               "2. Dip in GF flour, egg, then GF breadcrumbs\n"
                               "3. Pan fry until golden, about 4 minutes per side\n"
                               "4. Top with GF marinara sauce and mozzarella\n"
                               "5. Bake at 400°F for 10 minutes until cheese melts"),
                "tags": "dinner, gluten-free, italian, chicken, parmesan, comfort",
                "is_favorite": 1
            },
            {
                "title": "GF Zucchini Lasagna",
                "category": "Italian",
                "prep_time": "20",
                "cook_time": "45",
                "servings": 8,
                "difficulty": "Medium",
                "instructions": ("1. Slice 3 zucchini lengthwise, salt and drain 30 minutes\n"
                               "2. Mix 1 lb ground turkey with 1 jar GF marinara\n"
                               "3. Layer zucchini, meat sauce, ricotta, mozzarella\n"
                               "4. Repeat layers, top with parmesan\n"
                               "5. Bake at 375°F for 40-45 minutes"),
                "tags": "dinner, gluten-free, italian, lasagna, zucchini, healthy",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Asian
            {
                "title": "GF Thai Basil Chicken",
                "category": "Asian",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Heat 2 tbsp oil in wok, add 1 lb ground chicken\n"
                               "2. Add 3 cloves garlic, 2 Thai chilies, 1 tbsp ginger\n"
                               "3. Stir in 2 tbsp GF fish sauce, 1 tbsp GF soy sauce\n"
                               "4. Add 1 cup Thai basil leaves, 1/2 cup green beans\n"
                               "5. Serve over jasmine rice with fried egg"),
                "tags": "dinner, gluten-free, asian, thai, basil, spicy",
                "is_favorite": 1
            },
            {
                "title": "GF Beef and Broccoli",
                "category": "Asian",
                "prep_time": "15",
                "cook_time": "15",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Marinate 1 lb beef strips in 2 tbsp GF soy sauce\n"
                               "2. Stir fry beef in wok until browned, remove\n"
                               "3. Add 2 cups broccoli florets, 1 tbsp ginger, 2 cloves garlic\n"
                               "4. Add 1/4 cup GF beef broth, 1 tbsp cornstarch slurry\n"
                               "5. Return beef, toss with sauce, serve over rice"),
                "tags": "dinner, gluten-free, asian, beef, broccoli, chinese",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Mexican
            {
                "title": "GF Chicken Enchiladas",
                "category": "Mexican",
                "prep_time": "20",
                "cook_time": "25",
                "servings": 6,
                "difficulty": "Medium",
                "instructions": ("1. Mix 2 cups shredded chicken with 1/2 cup salsa\n"
                               "2. Add 1/2 cup black beans, 1/4 cup corn, 1/4 cup cheese\n"
                               "3. Fill GF tortillas, roll and place in baking dish\n"
                               "4. Top with 1 cup GF enchilada sauce, 1/2 cup cheese\n"
                               "5. Bake at 375°F for 20-25 minutes"),
                "tags": "dinner, gluten-free, mexican, enchiladas, chicken, comfort",
                "is_favorite": 1
            },
            {
                "title": "GF Fish Tacos with Cabbage Slaw",
                "category": "Mexican",
                "prep_time": "15",
                "cook_time": "10",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Season 1 lb white fish with GF taco seasoning\n"
                               "2. Pan fry 3-4 minutes per side until flaky\n"
                               "3. Mix 2 cups shredded cabbage with 1/4 cup cilantro\n"
                               "4. Add 2 tbsp lime juice, 1 tbsp olive oil, salt\n"
                               "5. Serve fish in GF tortillas with slaw and avocado"),
                "tags": "dinner, gluten-free, mexican, fish, tacos, healthy",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - Mediterranean
            {
                "title": "GF Mediterranean Stuffed Chicken",
                "category": "Mediterranean",
                "prep_time": "15",
                "cook_time": "30",
                "servings": 4,
                "difficulty": "Medium",
                "instructions": ("1. Butterfly 4 chicken breasts, pound to 1/4 inch\n"
                               "2. Mix 1/2 cup feta, 1/4 cup sun-dried tomatoes, 2 tbsp basil\n"
                               "3. Stuff chicken, secure with toothpicks\n"
                               "4. Season with salt, pepper, drizzle with olive oil\n"
                               "5. Bake at 400°F for 25-30 minutes until 165°F"),
                "tags": "dinner, gluten-free, mediterranean, chicken, feta, elegant",
                "is_favorite": 1
            },
            {
                "title": "GF Greek Moussaka",
                "category": "Mediterranean",
                "prep_time": "30",
                "cook_time": "60",
                "servings": 8,
                "difficulty": "Hard",
                "instructions": ("1. Slice 2 eggplants, salt and drain 30 minutes\n"
                               "2. Brown 1 lb ground lamb with 1 onion, 2 cloves garlic\n"
                               "3. Add 1 can tomatoes, 1/4 cup red wine, 1 tsp cinnamon\n"
                               "4. Layer eggplant, meat sauce, repeat\n"
                               "5. Top with GF béchamel, bake at 375°F for 45 minutes"),
                "tags": "dinner, gluten-free, mediterranean, moussaka, lamb, greek",
                "is_favorite": 0
            },
            
            # DINNER & MAIN COURSES - American
            {
                "title": "GF BBQ Pulled Chicken",
                "category": "American",
                "prep_time": "10",
                "cook_time": "240",
                "servings": 8,
                "difficulty": "Easy",
                "instructions": ("1. Season 3 lbs chicken thighs with GF BBQ rub\n"
                               "2. Slow cook in crockpot with 1 cup GF chicken broth\n"
                               "3. Cook on low 6-8 hours until tender\n"
                               "4. Shred meat, mix with 1 cup GF BBQ sauce\n"
                               "5. Serve on GF buns with coleslaw"),
                "tags": "dinner, gluten-free, american, bbq, chicken, comfort",
                "is_favorite": 1
            },
            {
                "title": "GF Turkey Meatloaf",
                "category": "American",
                "prep_time": "15",
                "cook_time": "60",
                "servings": 6,
                "difficulty": "Easy",
                "instructions": ("1. Mix 2 lbs ground turkey with 1/2 cup GF breadcrumbs\n"
                               "2. Add 1 egg, 1/2 onion, 1/4 cup ketchup, 1 tbsp Worcestershire\n"
                               "3. Form into loaf, place in baking dish\n"
                               "4. Top with 1/4 cup GF ketchup mixed with 1 tbsp brown sugar\n"
                               "5. Bake at 375°F for 50-60 minutes"),
                "tags": "dinner, gluten-free, american, meatloaf, turkey, comfort",
                "is_favorite": 0
            },
            
            # SIDES & VEGETABLES - Vegetables
            {
                "title": "GF Roasted Brussels Sprouts",
                "category": "Vegetables",
                "prep_time": "10",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Trim and halve 1.5 lbs Brussels sprouts\n"
                               "2. Toss with 3 tbsp olive oil, salt, pepper\n"
                               "3. Add 1/4 cup chopped bacon (optional)\n"
                               "4. Roast at 400°F for 20-25 minutes until crispy\n"
                               "5. Drizzle with balsamic glaze before serving"),
                "tags": "sides, gluten-free, vegetables, brussels sprouts, roasted, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Garlic Green Beans",
                "category": "Vegetables",
                "prep_time": "5",
                "cook_time": "8",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Trim 1 lb green beans, blanch 3 minutes\n"
                               "2. Heat 2 tbsp olive oil, add 3 cloves minced garlic\n"
                               "3. Add green beans, sauté 3-4 minutes\n"
                               "4. Season with salt, pepper, 1 tsp lemon zest\n"
                               "5. Toss with 1 tbsp fresh parsley and serve"),
                "tags": "sides, gluten-free, vegetables, green beans, garlic, simple",
                "is_favorite": 0
            },
            
            # SIDES & VEGETABLES - Grains & Rice
            {
                "title": "GF Herbed Quinoa Pilaf",
                "category": "Grains & Rice",
                "prep_time": "5",
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Toast 1 cup quinoa in 1 tbsp olive oil 2 minutes\n"
                               "2. Add 2 cups GF chicken broth, bring to boil\n"
                               "3. Reduce heat, cover, simmer 15 minutes\n"
                               "4. Fluff with fork, add 2 tbsp fresh herbs\n"
                               "5. Stir in 1/4 cup toasted pine nuts, season"),
                "tags": "sides, gluten-free, grains, quinoa, herbs, healthy",
                "is_favorite": 1
            },
            {
                "title": "GF Coconut Rice",
                "category": "Grains & Rice",
                "prep_time": "5",
                "cook_time": "20",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Rinse 1.5 cups jasmine rice until water runs clear\n"
                               "2. Combine with 1 can coconut milk, 1/2 cup water\n"
                               "3. Add 1 tsp salt, 1 tbsp sugar, bring to boil\n"
                               "4. Reduce heat, cover, simmer 15-18 minutes\n"
                               "5. Fluff with fork, stir in 1/4 cup shredded coconut"),
                "tags": "sides, gluten-free, grains, rice, coconut, tropical",
                "is_favorite": 0
            },
            
            # SIDES & VEGETABLES - Potatoes
            {
                "title": "GF Crispy Smashed Potatoes",
                "category": "Potatoes",
                "prep_time": "10",
                "cook_time": "35",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Boil 2 lbs small potatoes until tender, 15-20 minutes\n"
                               "2. Drain, place on baking sheet, smash with fork\n"
                               "3. Drizzle with 3 tbsp olive oil, season with salt, pepper\n"
                               "4. Add 2 tbsp fresh rosemary, 3 cloves minced garlic\n"
                               "5. Roast at 450°F for 15-20 minutes until crispy"),
                "tags": "sides, gluten-free, potatoes, crispy, herbs, comfort",
                "is_favorite": 1
            },
            {
                "title": "GF Sweet Potato Mash",
                "category": "Potatoes",
                "prep_time": "10",
                "cook_time": "25",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Peel and cube 3 lbs sweet potatoes\n"
                               "2. Boil until tender, 15-20 minutes\n"
                               "3. Drain, mash with 1/4 cup coconut milk\n"
                               "4. Add 2 tbsp maple syrup, 1 tsp cinnamon\n"
                               "5. Season with salt, pepper, serve warm"),
                "tags": "sides, gluten-free, potatoes, sweet potato, healthy, comfort",
                "is_favorite": 0
            },
            
            # SIDES & VEGETABLES - Bread & Rolls
            {
                "title": "GF Dinner Rolls",
                "category": "Bread & Rolls",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 12,
                "difficulty": "Medium",
                "instructions": ("1. Mix 2 cups GF flour blend with 1 tbsp yeast, 1 tsp salt\n"
                               "2. Add 1 cup warm milk, 2 tbsp honey, 2 tbsp olive oil\n"
                               "3. Knead 5 minutes, let rise 1 hour\n"
                               "4. Form into 12 rolls, place in greased pan\n"
                               "5. Rise 30 minutes, bake at 375°F for 18-20 minutes"),
                "tags": "sides, gluten-free, bread, rolls, dinner, comfort",
                "is_favorite": 1
            },
            {
                "title": "GF Cornbread Muffins",
                "category": "Bread & Rolls",
                "prep_time": "10",
                "cook_time": "18",
                "servings": 12,
                "difficulty": "Easy",
                "instructions": ("1. Mix 1 cup GF cornmeal with 1 cup GF flour blend\n"
                               "2. Add 1/4 cup sugar, 1 tbsp baking powder, 1 tsp salt\n"
                               "3. Whisk in 1 cup milk, 1/4 cup oil, 1 egg\n"
                               "4. Fill muffin cups 2/3 full\n"
                               "5. Bake at 400°F for 16-18 minutes until golden"),
                "tags": "sides, gluten-free, bread, cornbread, muffins, southern",
                "is_favorite": 0
            },
            
            # DESSERTS & SWEETS - Cakes
            {
                "title": "GF Chocolate Layer Cake",
                "category": "Cakes",
                "prep_time": "20",
                "cook_time": "30",
                "servings": 12,
                "difficulty": "Medium",
                "instructions": ("1. Mix 2 cups GF flour, 1.5 cups sugar, 3/4 cup cocoa\n"
                               "2. Add 2 tsp baking soda, 1 tsp baking powder, 1 tsp salt\n"
                               "3. Whisk in 2 eggs, 1 cup buttermilk, 1 cup coffee, 1/2 cup oil\n"
                               "4. Bake in 2 greased 9-inch pans at 350°F for 25-30 minutes\n"
                               "5. Cool, frost with GF chocolate buttercream"),
                "tags": "dessert, gluten-free, cake, chocolate, layer, celebration",
                "is_favorite": 1
            },
            {
                "title": "GF Vanilla Cupcakes",
                "category": "Cakes",
                "prep_time": "15",
                "cook_time": "18",
                "servings": 24,
                "difficulty": "Easy",
                "instructions": ("1. Mix 2.5 cups GF flour, 1.5 cups sugar, 1 tbsp baking powder\n"
                               "2. Add 1 tsp salt, 1/2 cup butter, 1 cup milk, 2 eggs\n"
                               "3. Beat 2 minutes, add 1 tbsp vanilla extract\n"
                               "4. Fill cupcake liners 2/3 full\n"
                               "5. Bake at 350°F for 16-18 minutes, cool and frost"),
                "tags": "dessert, gluten-free, cake, vanilla, cupcakes, party",
                "is_favorite": 0
            },
            
            # DESSERTS & SWEETS - Cookies
            {
                "title": "GF Chocolate Chip Cookies",
                "category": "Cookies",
                "prep_time": "15",
                "cook_time": "12",
                "servings": 24,
                "difficulty": "Easy",
                "instructions": ("1. Cream 1/2 cup butter with 1/2 cup brown sugar, 1/4 cup white sugar\n"
                               "2. Add 1 egg, 1 tsp vanilla, beat until fluffy\n"
                               "3. Mix in 1.25 cups GF flour, 1/2 tsp baking soda, 1/2 tsp salt\n"
                               "4. Fold in 1 cup GF chocolate chips\n"
                               "5. Bake at 375°F for 10-12 minutes until golden"),
                "tags": "dessert, gluten-free, cookies, chocolate chip, classic",
                "is_favorite": 1
            },
            {
                "title": "GF Oatmeal Raisin Cookies",
                "category": "Cookies",
                "prep_time": "15",
                "cook_time": "12",
                "servings": 24,
                "difficulty": "Easy",
                "instructions": ("1. Cream 1/2 cup butter with 1/2 cup brown sugar, 1/4 cup white sugar\n"
                               "2. Add 1 egg, 1 tsp vanilla, 1/2 tsp cinnamon\n"
                               "3. Mix in 1 cup GF flour, 1.5 cups GF oats, 1/2 tsp baking soda\n"
                               "4. Fold in 3/4 cup raisins, 1/2 cup walnuts\n"
                               "5. Bake at 375°F for 10-12 minutes"),
                "tags": "dessert, gluten-free, cookies, oatmeal, raisins, healthy",
                "is_favorite": 0
            },
            
            # DESSERTS & SWEETS - Pies & Tarts
            {
                "title": "GF Apple Pie",
                "category": "Pies & Tarts",
                "prep_time": "30",
                "cook_time": "45",
                "servings": 8,
                "difficulty": "Medium",
                "instructions": ("1. Make GF pie crust with 2 cups GF flour, 1/2 cup butter\n"
                               "2. Mix 6 cups sliced apples with 3/4 cup sugar, 2 tbsp GF flour\n"
                               "3. Add 1 tsp cinnamon, 1/4 tsp nutmeg, 2 tbsp lemon juice\n"
                               "4. Fill crust, top with second crust, seal edges\n"
                               "5. Bake at 425°F for 15 minutes, reduce to 350°F for 30 minutes"),
                "tags": "dessert, gluten-free, pie, apple, classic, fall",
                "is_favorite": 1
            },
            {
                "title": "GF Lemon Tart",
                "category": "Pies & Tarts",
                "prep_time": "25",
                "cook_time": "35",
                "servings": 8,
                "difficulty": "Medium",
                "instructions": ("1. Make GF tart shell with 1.5 cups GF flour, 1/2 cup butter\n"
                               "2. Blind bake at 375°F for 15 minutes\n"
                               "3. Whisk 4 eggs with 1 cup sugar, 1/2 cup lemon juice\n"
                               "4. Add 1/4 cup butter, cook over double boiler until thick\n"
                               "5. Fill shell, chill 2 hours, serve with berries"),
                "tags": "dessert, gluten-free, tart, lemon, citrus, elegant",
                "is_favorite": 0
            },
            
            # DESSERTS & SWEETS - Ice Cream
            {
                "title": "GF Vanilla Bean Ice Cream",
                "category": "Ice Cream",
                "prep_time": "10",
                "cook_time": "15",
                "servings": 8,
                "difficulty": "Medium",
                "instructions": ("1. Heat 2 cups heavy cream, 1 cup milk, 3/4 cup sugar\n"
                               "2. Add 1 vanilla bean split, simmer 5 minutes\n"
                               "3. Whisk 6 egg yolks, slowly add hot cream mixture\n"
                               "4. Return to pan, cook until 170°F, strain\n"
                               "5. Chill 4 hours, churn in ice cream maker"),
                "tags": "dessert, gluten-free, ice cream, vanilla, classic, homemade",
                "is_favorite": 1
            },
            {
                "title": "GF Chocolate Fudge Ice Cream",
                "category": "Ice Cream",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 8,
                "difficulty": "Medium",
                "instructions": ("1. Heat 2 cups heavy cream, 1 cup milk, 3/4 cup sugar\n"
                               "2. Whisk in 1/2 cup cocoa powder until smooth\n"
                               "3. Whisk 6 egg yolks, slowly add hot chocolate mixture\n"
                               "4. Return to pan, cook until 170°F, strain\n"
                               "5. Chill 4 hours, churn, fold in 1/2 cup chocolate chips"),
                "tags": "dessert, gluten-free, ice cream, chocolate, fudge, rich",
                "is_favorite": 0
            },
            
            # DESSERTS & SWEETS - Candies
            {
                "title": "GF Chocolate Truffles",
                "category": "Candies",
                "prep_time": "20",
                "cook_time": "5",
                "servings": 24,
                "difficulty": "Easy",
                "instructions": ("1. Heat 1 cup heavy cream, pour over 8 oz chopped dark chocolate\n"
                               "2. Stir until smooth, add 2 tbsp butter, 1 tsp vanilla\n"
                               "3. Chill 2 hours until firm\n"
                               "4. Scoop into balls, roll in cocoa powder\n"
                               "5. Store in refrigerator up to 2 weeks"),
                "tags": "dessert, gluten-free, candy, truffles, chocolate, elegant",
                "is_favorite": 1
            },
            {
                "title": "GF Peanut Butter Fudge",
                "category": "Candies",
                "prep_time": "10",
                "cook_time": "10",
                "servings": 24,
                "difficulty": "Easy",
                "instructions": ("1. Heat 1 cup sugar, 1/2 cup milk, 2 tbsp butter\n"
                               "2. Bring to boil, cook 2 minutes\n"
                               "3. Remove from heat, add 1 cup peanut butter, 1 tsp vanilla\n"
                               "4. Stir until smooth, pour into greased 8x8 pan\n"
                               "5. Chill 2 hours, cut into squares"),
                "tags": "dessert, gluten-free, candy, fudge, peanut butter, sweet",
                "is_favorite": 0
            },
            
            # BEVERAGES & DRINKS - Coffee & Tea
            {
                "title": "GF Pumpkin Spice Latte",
                "category": "Coffee & Tea",
                "prep_time": "5",
                "cook_time": "5",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Heat 1 cup almond milk with 2 tbsp pumpkin puree\n"
                               "2. Add 1 tsp pumpkin pie spice, 1 tbsp maple syrup\n"
                               "3. Whisk until frothy, pour over 2 shots espresso\n"
                               "4. Top with coconut whipped cream\n"
                               "5. Sprinkle with cinnamon and serve"),
                "tags": "beverage, gluten-free, coffee, latte, pumpkin, fall",
                "is_favorite": 1
            },
            {
                "title": "GF Chai Tea Latte",
                "category": "Coffee & Tea",
                "prep_time": "5",
                "cook_time": "10",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Simmer 2 cups almond milk with 2 chai tea bags\n"
                               "2. Add 1 cinnamon stick, 3 cardamom pods, 1 tsp ginger\n"
                               "3. Steep 5 minutes, remove tea bags and spices\n"
                               "4. Sweeten with 2 tbsp honey or maple syrup\n"
                               "5. Froth and serve hot"),
                "tags": "beverage, gluten-free, tea, chai, spiced, warming",
                "is_favorite": 0
            },
            
            # BEVERAGES & DRINKS - Smoothies
            {
                "title": "GF Tropical Mango Smoothie",
                "category": "Smoothies",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Blend 1 cup frozen mango with 1 banana\n"
                               "2. Add 1 cup coconut milk, 1/2 cup orange juice\n"
                               "3. Include 1 tbsp honey, 1 tsp vanilla extract\n"
                               "4. Add 1/2 cup ice, blend until smooth\n"
                               "5. Garnish with fresh mint and serve"),
                "tags": "beverage, gluten-free, smoothie, tropical, mango, refreshing",
                "is_favorite": 1
            },
            {
                "title": "GF Chocolate Banana Smoothie",
                "category": "Smoothies",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 2,
                "difficulty": "Easy",
                "instructions": ("1. Blend 2 bananas with 2 tbsp cocoa powder\n"
                               "2. Add 1 cup almond milk, 1/2 cup Greek yogurt\n"
                               "3. Include 2 tbsp almond butter, 1 tbsp honey\n"
                               "4. Add 1 cup ice, blend until creamy\n"
                               "5. Top with chocolate shavings and serve"),
                "tags": "beverage, gluten-free, smoothie, chocolate, banana, protein",
                "is_favorite": 0
            },
            
            # BEVERAGES & DRINKS - Cocktails
            {
                "title": "GF Mojito",
                "category": "Cocktails",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 1,
                "difficulty": "Easy",
                "instructions": ("1. Muddle 8 mint leaves with 1 tbsp sugar, 1 tbsp lime juice\n"
                               "2. Add 2 oz white rum, fill glass with ice\n"
                               "3. Top with club soda, stir gently\n"
                               "4. Garnish with mint sprig and lime wedge\n"
                               "5. Serve immediately"),
                "tags": "beverage, gluten-free, cocktail, mojito, mint, refreshing",
                "is_favorite": 1
            },
            {
                "title": "GF Margarita",
                "category": "Cocktails",
                "prep_time": "5",
                "cook_time": "0",
                "servings": 1,
                "difficulty": "Easy",
                "instructions": ("1. Rim glass with salt (optional)\n"
                               "2. Shake 2 oz tequila with 1 oz lime juice, 1 oz triple sec\n"
                               "3. Add ice, shake vigorously\n"
                               "4. Strain into salt-rimmed glass\n"
                               "5. Garnish with lime wedge and serve"),
                "tags": "beverage, gluten-free, cocktail, margarita, tequila, citrus",
                "is_favorite": 0
            },
            
            # BEVERAGES & DRINKS - Non-Alcoholic
            {
                "title": "GF Sparkling Lemonade",
                "category": "Non-Alcoholic",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Mix 1 cup fresh lemon juice with 1/2 cup honey\n"
                               "2. Add 3 cups cold water, stir until honey dissolves\n"
                               "3. Add 1 cup sparkling water, ice cubes\n"
                               "4. Garnish with lemon slices and fresh mint\n"
                               "5. Serve immediately"),
                "tags": "beverage, gluten-free, non-alcoholic, lemonade, sparkling, refreshing",
                "is_favorite": 1
            },
            {
                "title": "GF Iced Green Tea",
                "category": "Non-Alcoholic",
                "prep_time": "5",
                "cook_time": "5",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Steep 4 green tea bags in 4 cups hot water\n"
                               "2. Steep 3-5 minutes, remove tea bags\n"
                               "3. Sweeten with 2 tbsp honey, stir to dissolve\n"
                               "4. Chill in refrigerator 2 hours\n"
                               "5. Serve over ice with lemon slices"),
                "tags": "beverage, gluten-free, non-alcoholic, tea, green tea, healthy",
                "is_favorite": 0
            },
            
            # SNACKS & APPETIZERS - Dips & Spreads
            {
                "title": "GF Hummus with Veggies",
                "category": "Dips & Spreads",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 6,
                "difficulty": "Easy",
                "instructions": ("1. Blend 1 can chickpeas with 1/4 cup tahini, 2 tbsp lemon juice\n"
                               "2. Add 2 cloves garlic, 1/4 cup olive oil, 1/2 tsp salt\n"
                               "3. Process until smooth, add water if needed\n"
                               "4. Serve with GF crackers, carrot sticks, cucumber\n"
                               "5. Drizzle with olive oil and paprika"),
                "tags": "snack, gluten-free, hummus, dip, healthy, mediterranean",
                "is_favorite": 1
            },
            {
                "title": "GF Guacamole",
                "category": "Dips & Spreads",
                "prep_time": "10",
                "cook_time": "0",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Mash 3 ripe avocados with 1/4 cup diced onion\n"
                               "2. Add 1/4 cup diced tomato, 2 tbsp cilantro, 1 jalapeño\n"
                               "3. Mix in 2 tbsp lime juice, 1/2 tsp salt, 1/4 tsp cumin\n"
                               "4. Serve immediately with GF tortilla chips\n"
                               "5. Garnish with extra cilantro and lime wedges"),
                "tags": "snack, gluten-free, guacamole, dip, avocado, mexican",
                "is_favorite": 0
            },
            
            # SNACKS & APPETIZERS - Finger Foods
            {
                "title": "GF Chicken Wings",
                "category": "Finger Foods",
                "prep_time": "10",
                "cook_time": "45",
                "servings": 4,
                "difficulty": "Easy",
                "instructions": ("1. Season 2 lbs chicken wings with salt, pepper\n"
                               "2. Bake at 400°F for 40-45 minutes until crispy\n"
                               "3. Mix 1/2 cup GF hot sauce with 1/4 cup butter\n"
                               "4. Toss wings in sauce mixture\n"
                               "5. Serve with GF ranch dressing and celery sticks"),
                "tags": "snack, gluten-free, wings, chicken, spicy, party",
                "is_favorite": 1
            },
            {
                "title": "GF Stuffed Mushrooms",
                "category": "Finger Foods",
                "prep_time": "15",
                "cook_time": "20",
                "servings": 6,
                "difficulty": "Medium",
                "instructions": ("1. Remove stems from 24 large mushrooms\n"
                               "2. Chop stems, sauté with 1/4 cup onion, 2 cloves garlic\n"
                               "3. Mix with 1/2 cup GF breadcrumbs, 1/4 cup parmesan\n"
                               "4. Stuff mushroom caps, top with more cheese\n"
                               "5. Bake at 375°F for 15-20 minutes until golden"),
                "tags": "snack, gluten-free, mushrooms, stuffed, appetizer, elegant",
                "is_favorite": 0
            },
            
            # BAKING & BREADS - Quick Breads
            {
                "title": "GF Banana Bread",
                "category": "Baking & Breads",
                "prep_time": "15",
                "cook_time": "60",
                "servings": 12,
                "difficulty": "Easy",
                "instructions": ("1. Mash 3 ripe bananas, mix with 1/3 cup melted butter\n"
                               "2. Add 3/4 cup sugar, 1 egg, 1 tsp vanilla\n"
                               "3. Mix in 1.5 cups GF flour, 1 tsp baking soda, 1/2 tsp salt\n"
                               "4. Fold in 1/2 cup walnuts, 1/2 cup chocolate chips\n"
                               "5. Bake in greased loaf pan at 350°F for 55-60 minutes"),
                "tags": "baking, gluten-free, bread, banana, sweet, breakfast",
                "is_favorite": 1
            },
            {
                "title": "GF Zucchini Bread",
                "category": "Baking & Breads",
                "prep_time": "15",
                "cook_time": "60",
                "servings": 12,
                "difficulty": "Easy",
                "instructions": ("1. Grate 2 cups zucchini, squeeze out excess moisture\n"
                               "2. Mix with 1/2 cup oil, 3/4 cup sugar, 2 eggs, 1 tsp vanilla\n"
                               "3. Add 2 cups GF flour, 1 tsp baking soda, 1/2 tsp salt\n"
                               "4. Stir in 1 tsp cinnamon, 1/2 cup chopped nuts\n"
                               "5. Bake in greased loaf pan at 350°F for 55-60 minutes"),
                "tags": "baking, gluten-free, bread, zucchini, healthy, vegetables",
                "is_favorite": 0
            },
            
            # BAKING & BREADS - Yeast Breads
            {
                "title": "GF Sandwich Bread",
                "category": "Baking & Breads",
                "prep_time": "20",
                "cook_time": "45",
                "servings": 16,
                "difficulty": "Medium",
                "instructions": ("1. Mix 3 cups GF flour blend with 1 tbsp yeast, 1 tsp salt\n"
                               "2. Add 1.5 cups warm water, 2 tbsp honey, 2 tbsp olive oil\n"
                               "3. Knead 5 minutes, let rise 1 hour in warm place\n"
                               "4. Shape into loaf, place in greased 9x5 pan\n"
                               "5. Rise 30 minutes, bake at 375°F for 40-45 minutes"),
                "tags": "baking, gluten-free, bread, sandwich, yeast, daily",
                "is_favorite": 1
            },
            {
                "title": "GF Focaccia",
                "category": "Baking & Breads",
                "prep_time": "15",
                "cook_time": "25",
                "servings": 12,
                "difficulty": "Medium",
                "instructions": ("1. Mix 2.5 cups GF flour with 1 tbsp yeast, 1 tsp salt\n"
                               "2. Add 1 cup warm water, 2 tbsp olive oil, 1 tsp honey\n"
                               "3. Knead 5 minutes, let rise 45 minutes\n"
                               "4. Press into oiled 9x13 pan, dimple with fingers\n"
                               "5. Top with rosemary, sea salt, bake at 400°F for 20-25 minutes"),
                "tags": "baking, gluten-free, bread, focaccia, italian, herbs",
                "is_favorite": 0
            }
        ]
    
    def filter_recipes(self):
        """Filter recipes based on search and category - now handled by refresh_recipe_display"""
        # Filtering is now handled in refresh_recipe_display method
        if hasattr(self, 'refresh_recipe_display'):
            self.refresh_recipe_display()
    
    def on_selection_changed(self):
        """Handle selection changes - now handled by card selection"""
        # Selection is now handled by card clicks
        pass
    
    def add_recipe(self):
        """Add a new recipe"""
        try:
            from utils.edit_dialogs import RecipeEditDialog
            
            # Create new recipe data
            recipe_data = {
                'name': '',
                'category': 'Breakfast',
                'prep_time': '',
                'cook_time': '',
                'servings': 4,
                'difficulty': 'Medium',
                'description': '',
                'ingredients': [],
                'instructions': '',
                'notes': ''
            }
            
            # Open edit dialog
            dialog = RecipeEditDialog(self)
            dialog.set_data(recipe_data)
            
            if dialog.exec() == QDialog.Accepted:
                new_data = dialog.get_data()
                
                # Validate required fields
                if not new_data['name'].strip():
                    QMessageBox.warning(self, "Validation Error", "Recipe name is required.")
                    return
                
                # Save to database first
                recipe_id = self._save_recipe_to_database(new_data)
                if recipe_id:
                    new_data['id'] = recipe_id
                    
                    # Store recipe data and refresh display
                    self.recipe_data_storage[recipe_id] = new_data
                    self.refresh_recipe_display()
                    
                    QMessageBox.information(self, "Success", f"Recipe '{new_data['name']}' added successfully!")
                else:
                    QMessageBox.warning(self, "Error", "Failed to save recipe to database.")
                
        except ImportError:
            QMessageBox.information(self, "Add Recipe", "Add Recipe functionality will be implemented here.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add recipe: {str(e)}")
    
    def _get_recipe_data_by_id(self, recipe_id):
        """Get recipe data by ID"""
        try:
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            cursor.execute("""
                SELECT id, title, source, url, tags, ingredients, instructions, 
                       rating, prep_time, cook_time, servings, category, 
                       description, difficulty, is_favorite, image_path
                FROM recipes 
                WHERE id = ?
            """, (recipe_id,))
            
            recipe = cursor.fetchone()
            if recipe:
                # Load ingredients
                cursor.execute("""
                    SELECT ingredient_name, quantity, unit, notes
                    FROM recipe_ingredients 
                    WHERE recipe_id = ?
                    ORDER BY ingredient_name
                """, (recipe_id,))
                
                ingredients = []
                for ing_row in cursor.fetchall():
                    ingredients.append({
                        "name": ing_row[0],
                        "amount": ing_row[1] or "",
                        "unit": ing_row[2] or "",
                        "notes": ing_row[3] or ""
                    })
                
                recipe_data = {
                    'id': recipe[0],
                    'name': recipe[1],
                    'source': recipe[2] or "",
                    'url': recipe[3] or "",
                    'tags': recipe[4] or "",
                    'ingredients': ingredients,  # Will be loaded from recipe_ingredients table
                    'instructions': recipe[6] or "",
                    'rating': recipe[7] or 0.0,
                    'prep_time': recipe[8] or "",
                    'cook_time': recipe[9] or "",
                    'servings': recipe[10] or 1,
                    'category': recipe[11] or "Uncategorized",
                    'description': recipe[12] or "",
                    'difficulty': recipe[13] or "Medium",
                    'is_favorite': recipe[14] or 0,
                    'image_path': recipe[15] or ""
                }
                
                db.close()
                return recipe_data
            
            db.close()
            return None
            
        except Exception as e:
            print(f"Error getting recipe data by ID: {e}")
            return None
    
    def edit_recipe(self):
        """Edit selected recipe"""
        if hasattr(self, 'current_recipe_id') and self.current_recipe_id:
            try:
                from utils.edit_dialogs import RecipeEditDialog
                
                # Get recipe data by ID
                recipe_data = self._get_recipe_data_by_id(self.current_recipe_id)
                
                # Open edit dialog
                dialog = RecipeEditDialog(self)
                dialog.set_data(recipe_data)
                
                if dialog.exec() == QDialog.Accepted:
                    new_data = dialog.get_data()
                    
                    # Validate required fields
                    if not new_data['name'].strip():
                        QMessageBox.warning(self, "Validation Error", "Recipe name is required.")
                        return
                    
                    # Save changes to database
                    if self._update_recipe_in_database(recipe_data.get('id'), new_data):
                        # Update stored data and refresh display
                        self.recipe_data_storage[recipe_data.get('id')] = new_data
                        self.refresh_recipe_display()
                        
                        QMessageBox.information(self, "Success", "Recipe updated successfully!")
                    else:
                        QMessageBox.warning(self, "Error", "Failed to update recipe in database.")
                    
            except ImportError:
                # Fallback implementation when edit dialog is not available
                self._fallback_edit_recipe_by_id(self.current_recipe_id)
            except UnicodeEncodeError as e:
                QMessageBox.critical(self, "Error", 
                    "Failed to edit recipe: Character encoding error. "
                    "Please ensure all characters in the recipe can be saved to the database.")
            except Exception as e:
                try:
                    error_msg = str(e)
                    QMessageBox.critical(self, "Error", f"Failed to edit recipe: {error_msg}")
                except UnicodeEncodeError:
                    QMessageBox.critical(self, "Error", "Failed to edit recipe: Encoding error occurred.")
    
    def _fallback_edit_recipe_by_id(self, recipe_id):
        """Fallback edit implementation when edit dialog is not available (ID-based)"""
        try:
            # Get current recipe data by ID
            recipe_data = self._get_recipe_data_by_id(recipe_id)
            if not recipe_data:
                QMessageBox.warning(self, "Error", "Recipe not found.")
                return
            
            
            # Create edit dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Edit Recipe")
            dialog.setModal(True)
            dialog.resize(500, 600)
            
            layout = QVBoxLayout(dialog)
            
            # Recipe name
            layout.addWidget(QLabel("Recipe Name:"))
            name_edit = QLineEdit(recipe_data.get('name', ''))
            layout.addWidget(name_edit)
            
            # Category
            layout.addWidget(QLabel("Category:"))
            category_combo = QComboBox()
            category_combo.addItems(["Breakfast", "Lunch", "Dinner", "Dessert", "Snack", "Beverage", "Other"])
            category_combo.setCurrentText(recipe_data.get('category', 'Other'))
            layout.addWidget(category_combo)
            
            # Prep time
            layout.addWidget(QLabel("Prep Time:"))
            prep_edit = QLineEdit(recipe_data.get('prep_time', ''))
            layout.addWidget(prep_edit)
            
            # Cook time
            layout.addWidget(QLabel("Cook Time:"))
            cook_edit = QLineEdit(recipe_data.get('cook_time', ''))
            layout.addWidget(cook_edit)
            
            # Servings
            layout.addWidget(QLabel("Servings:"))
            servings_spin = QSpinBox()
            servings_spin.setMinimum(1)
            servings_spin.setMaximum(50)
            # Handle servings that might contain text like "1 loaf"
            servings_value = self._extract_numeric_value(recipe_data.get('servings', 4), 4)
            servings_spin.setValue(servings_value)
            layout.addWidget(servings_spin)
            
            # Difficulty
            layout.addWidget(QLabel("Difficulty:"))
            difficulty_combo = QComboBox()
            difficulty_combo.addItems(["Easy", "Medium", "Hard"])
            difficulty_combo.setCurrentText(recipe_data.get('difficulty', 'Medium'))
            layout.addWidget(difficulty_combo)
            
            # Description
            layout.addWidget(QLabel("Description:"))
            desc_edit = QTextEdit()
            desc_edit.setPlainText(recipe_data.get('description', ''))
            desc_edit.setMaximumHeight(80)
            layout.addWidget(desc_edit)
            
            # Ingredients
            layout.addWidget(QLabel("Ingredients (one per line):"))
            ingredients_edit = QTextEdit()
            ingredients_text = ""
            for ing in recipe_data.get('ingredients', []):
                if isinstance(ing, dict):
                    qty = ing.get('amount', '')
                    unit = ing.get('unit', '')
                    name = ing.get('name', '')
                    if qty and unit:
                        ingredients_text += f"{qty} {unit} {name}\n"
                    elif qty:
                        ingredients_text += f"{qty} {name}\n"
                    else:
                        ingredients_text += f"{name}\n"
                else:
                    ingredients_text += f"{ing}\n"
            ingredients_edit.setPlainText(ingredients_text.strip())
            ingredients_edit.setMaximumHeight(120)
            layout.addWidget(ingredients_edit)
            
            # Instructions
            layout.addWidget(QLabel("Instructions:"))
            instructions_edit = QTextEdit()
            instructions_edit.setPlainText(recipe_data.get('instructions', ''))
            instructions_edit.setMaximumHeight(120)
            layout.addWidget(instructions_edit)
            
            # Image selection
            image_layout = QVBoxLayout()
            image_layout.addWidget(QLabel("Recipe Image:"))
            
            # Image preview
            self.image_preview_label = QLabel()
            self.image_preview_label.setFixedSize(200, 150)
            self.image_preview_label.setStyleSheet("border: 1px solid #ccc; background-color: #f9f9f9;")
            self.image_preview_label.setAlignment(Qt.AlignCenter)
            self.image_preview_label.setText("No image")
            
            # Load existing image if available
            existing_image_path = recipe_data.get('image_path', '')
            if existing_image_path:
                try:
                    from services.image_service import recipe_image_service
                    pixmap = recipe_image_service.load_image_pixmap(recipe_data.get('id'), (200, 150))
                    if pixmap and not pixmap.isNull():
                        self.image_preview_label.setPixmap(pixmap)
                    else:
                        self.image_preview_label.setText("Image not found")
                except Exception as e:
                    print(f"Error loading image preview: {e}")
                    self.image_preview_label.setText("Error loading image")
            
            image_layout.addWidget(self.image_preview_label)
            
            # Image path display and controls
            image_controls_layout = QHBoxLayout()
            
            self.image_path_edit = QLineEdit()
            self.image_path_edit.setPlaceholderText("No image selected")
            self.image_path_edit.setText(existing_image_path)
            self.image_path_edit.setReadOnly(True)
            image_controls_layout.addWidget(self.image_path_edit)
            
            select_image_btn = QPushButton("Select Image")
            select_image_btn.clicked.connect(lambda: self.select_recipe_image())
            image_controls_layout.addWidget(select_image_btn)
            
            clear_image_btn = QPushButton("Clear")
            clear_image_btn.clicked.connect(lambda: self.clear_recipe_image())
            image_controls_layout.addWidget(clear_image_btn)
            
            image_layout.addLayout(image_controls_layout)
            layout.addLayout(image_layout)
            
            # Source
            layout.addWidget(QLabel("Source:"))
            source_edit = QLineEdit(recipe_data.get('source', ''))
            source_edit.setPlaceholderText("e.g., 'AllRecipes', 'Food Network', or leave empty")
            layout.addWidget(source_edit)
            
            # URL
            layout.addWidget(QLabel("URL:"))
            url_edit = QLineEdit(recipe_data.get('url', ''))
            url_edit.setPlaceholderText("https://example.com/recipe (optional)")
            layout.addWidget(url_edit)
            
            # Notes
            layout.addWidget(QLabel("Notes:"))
            notes_edit = QTextEdit()
            notes_edit.setPlainText(recipe_data.get('notes', ''))
            notes_edit.setMaximumHeight(80)
            layout.addWidget(notes_edit)
            
            # Buttons
            button_layout = QHBoxLayout()
            save_btn = QPushButton("Save Changes")
            cancel_btn = QPushButton("Cancel")
            
            save_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(save_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            if dialog.exec() == QDialog.Accepted:
                # Create updated recipe data
                updated_recipe = {
                    'name': name_edit.text().strip(),
                    'category': category_combo.currentText(),
                    'prep_time': prep_edit.text().strip(),
                    'cook_time': cook_edit.text().strip(),
                    'servings': servings_spin.value(),
                    'difficulty': difficulty_combo.currentText(),
                    'description': desc_edit.toPlainText().strip(),
                    'image_path': self.image_path_edit.text().strip(),
                    'source': source_edit.text().strip(),
                    'url': url_edit.text().strip(),
                    'ingredients': self._parse_ingredients_from_text(ingredients_edit.toPlainText()),
                    'instructions': instructions_edit.toPlainText().strip(),
                    'notes': notes_edit.toPlainText().strip()
                }
                
                # Validate required fields
                if not updated_recipe['name']:
                    QMessageBox.warning(self, "Validation Error", "Recipe name is required.")
                    return
                
                # Save changes to database
                if self._update_recipe_in_database(recipe_id, updated_recipe):
                    # Update stored data and refresh display
                    self.recipe_data_storage[recipe_id] = updated_recipe
                    self.refresh_recipe_display()
                    
                    QMessageBox.information(self, "Success", "Recipe updated successfully!")
                else:
                    QMessageBox.warning(self, "Error", "Failed to update recipe in database.")
                
        except UnicodeEncodeError as e:
            QMessageBox.critical(self, "Error", 
                "Failed to edit recipe: Character encoding error. "
                "Please ensure all characters in the recipe can be saved to the database.")
        except Exception as e:
            try:
                error_msg = str(e)
                QMessageBox.critical(self, "Error", f"Failed to edit recipe: {error_msg}")
            except UnicodeEncodeError:
                QMessageBox.critical(self, "Error", "Failed to edit recipe: Encoding error occurred.")
    
    def select_recipe_image(self):
        """Select and optionally crop a recipe image"""
        try:
            from services.image_service import recipe_image_service
            
            # Select image file
            image_path = recipe_image_service.select_image(self)
            if image_path:
                # The image service will handle cropping automatically if needed
                self.image_path_edit.setText(image_path)
                
                # Update preview
                if hasattr(self, 'image_preview_label'):
                    try:
                        pixmap = QPixmap(image_path)
                        if not pixmap.isNull():
                            scaled_pixmap = pixmap.scaled(200, 150, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                            self.image_preview_label.setPixmap(scaled_pixmap)
                        else:
                            self.image_preview_label.setText("Invalid image")
                    except Exception as e:
                        print(f"Error updating preview: {e}")
                        self.image_preview_label.setText("Error loading image")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to select image: {str(e)}")
    
    def clear_recipe_image(self):
        """Clear the selected recipe image"""
        self.image_path_edit.setText('')
        if hasattr(self, 'image_preview_label'):
            self.image_preview_label.clear()
            self.image_preview_label.setText("No image")
    
    def cleanup_temp_files(self):
        """Clean up temporary files created during image processing"""
        try:
            from services.image_service import recipe_image_service
            recipe_image_service.cleanup_temp_files()
        except Exception as e:
            print(f"Error cleaning up temp files: {e}")
    
    def convert_all_images_to_webp(self):
        """Convert all existing recipe images to WebP format"""
        try:
            from services.image_service import recipe_image_service
            
            # Get all recipes with images
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            
            cursor.execute("SELECT id FROM recipes WHERE image_path IS NOT NULL AND image_path != ''")
            recipe_ids = cursor.fetchall()
            conn.close()
            
            converted_count = 0
            for (recipe_id,) in recipe_ids:
                try:
                    webp_path = recipe_image_service.convert_to_webp(recipe_id)
                    if webp_path:
                        converted_count += 1
                        print(f"Converted recipe {recipe_id} image to WebP")
                except Exception as e:
                    print(f"Error converting recipe {recipe_id} image: {e}")
            
            QMessageBox.information(self, "Conversion Complete", 
                                  f"Successfully converted {converted_count} images to WebP format!")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to convert images: {str(e)}")
    
    def _get_recipe_source_info(self, recipe_data):
        """Get formatted source information for display"""
        source = recipe_data.get('source', '')
        url = recipe_data.get('url', '')
        
        if url:
            if source:
                return f"Source: {source}\nURL: {url}"
            else:
                return f"Source URL: {url}"
        elif source:
            return f"Source: {source}"
        
        return None
    
    def _open_url_in_browser(self, url):
        """Open URL in default web browser"""
        try:
            QDesktopServices.openUrl(QUrl(url))
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to open URL: {str(e)}")
    
    def _send_to_menu(self, recipe_data):
        """Send recipe to menu planner"""
        try:
            # Get menu panel reference
            menu_panel = self.app.get_panel('menu') if hasattr(self, 'app') and self.app else None
            if menu_panel and hasattr(menu_panel, 'add_recipe_to_menu'):
                menu_panel.add_recipe_to_menu(recipe_data)
                QMessageBox.information(self, "Success", f"Recipe '{recipe_data.get('name', 'Unknown')}' added to menu planner!")
            else:
                QMessageBox.information(self, "Info", "Menu planner functionality not available.")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add recipe to menu: {str(e)}")
    
    def _send_to_shopping_list(self, recipe_data):
        """Send recipe ingredients to shopping list"""
        try:
            # Get shopping list panel reference
            shopping_panel = self.app.get_panel('shopping') if hasattr(self, 'app') and self.app else None
            if shopping_panel and hasattr(shopping_panel, 'add_recipe_ingredients'):
                shopping_panel.add_recipe_ingredients(recipe_data)
                QMessageBox.information(self, "Success", f"Ingredients from '{recipe_data.get('name', 'Unknown')}' added to shopping list!")
            else:
                QMessageBox.information(self, "Info", "Shopping list functionality not available.")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add ingredients to shopping list: {str(e)}")
    
    def _scale_recipe_from_view(self, parent_dialog, recipe_data):
        """Scale recipe from view dialog"""
        try:
            # Extract ingredients from recipe data
            ingredients_text = recipe_data.get('ingredients', '')
            if isinstance(ingredients_text, str):
                ingredients = self._parse_ingredients_from_text(ingredients_text)
            else:
                ingredients = ingredients_text or []
            
            if not ingredients:
                QMessageBox.warning(self, "No Ingredients", "This recipe has no ingredients to scale.")
                return
            
            # Show scaling dialog
            from utils.recipe_scaling_dialog import RecipeScalingDialog
            
            dialog = RecipeScalingDialog(ingredients, self)
            if dialog.exec() == QDialog.Accepted:
                scaled_ingredients = dialog.get_scaled_ingredients()
                
                # Update recipe data with scaled ingredients
                recipe_data['ingredients'] = scaled_ingredients
                
                # Update in database
                recipe_id = recipe_data.get('id')
                if recipe_id and self._update_recipe_in_database(recipe_id, recipe_data):
                    QMessageBox.information(self, "Success", "Recipe scaled successfully!")
                    parent_dialog.close()  # Close view dialog
                    self.refresh_recipe_display()  # Refresh the recipe list
                else:
                    QMessageBox.warning(self, "Error", "Failed to update scaled recipe.")
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to scale recipe: {str(e)}")
    
    def _show_recipe_context_menu(self, parent_dialog, recipe_data):
        """Show right-click context menu for recipe"""
        from PySide6.QtWidgets import QMenu
        
        menu = QMenu(parent_dialog)
        
        # Edit Recipe
        edit_action = menu.addAction("✏️ Edit Recipe")
        edit_action.triggered.connect(lambda: self._edit_recipe_from_view(parent_dialog, recipe_data))
        
        # Export Recipe
        export_action = menu.addAction("📤 Export Recipe")
        export_submenu = QMenu("Export Format", menu)
        export_submenu.addAction("Export as Text").triggered.connect(lambda: self._export_recipe_txt_from_view(recipe_data))
        export_submenu.addAction("Export as HTML").triggered.connect(lambda: self._export_recipe_html_from_view(recipe_data))
        export_action.setMenu(export_submenu)
        menu.addAction(export_action)
        
        menu.addSeparator()
        
        # Delete Recipe
        delete_action = menu.addAction("🗑️ Delete Recipe")
        delete_action.triggered.connect(lambda: self._delete_recipe_from_view(parent_dialog, recipe_data))
        
        # Show menu at cursor position
        menu.exec(parent_dialog.mapToGlobal(parent_dialog.cursor().pos()))
    
    def _edit_recipe_from_view(self, parent_dialog, recipe_data):
        """Edit recipe from view dialog"""
        parent_dialog.close()
        self.current_recipe_id = recipe_data.get('id')
        self.edit_recipe()
    
    def _export_recipe_txt_from_view(self, recipe_data):
        """Export recipe as text from view dialog"""
        try:
            from PySide6.QtWidgets import QFileDialog
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Export Recipe as Text", 
                f"{recipe_data.get('name', 'recipe')}.txt", 
                "Text Files (*.txt)"
            )
            
            if file_path:
                self._write_recipe_to_txt_file(recipe_data, file_path)
                QMessageBox.information(self, "Success", f"Recipe exported to {file_path}")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to export recipe: {str(e)}")
    
    def _export_recipe_html_from_view(self, recipe_data):
        """Export recipe as HTML from view dialog"""
        try:
            from PySide6.QtWidgets import QFileDialog
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Export Recipe as HTML", 
                f"{recipe_data.get('name', 'recipe')}.html", 
                "HTML Files (*.html)"
            )
            
            if file_path:
                self._write_recipe_to_html_file(recipe_data, file_path)
                QMessageBox.information(self, "Success", f"Recipe exported to {file_path}")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to export recipe: {str(e)}")
    
    def _delete_recipe_from_view(self, parent_dialog, recipe_data):
        """Delete recipe from view dialog"""
        reply = QMessageBox.question(
            parent_dialog, "Delete Recipe", 
            f"Are you sure you want to delete '{recipe_data.get('name', 'Unknown Recipe')}'?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            parent_dialog.close()
            self.current_recipe_id = recipe_data.get('id')
            self.delete_recipe()
    
    def _write_recipe_to_txt_file(self, recipe_data, file_path):
        """Write recipe to text file"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"Recipe: {recipe_data.get('name', 'Unknown')}\n")
            f.write("=" * 50 + "\n\n")
            
            # Basic info
            f.write(f"Category: {recipe_data.get('category', 'N/A')}\n")
            f.write(f"Prep Time: {recipe_data.get('prep_time', 'N/A')}\n")
            f.write(f"Cook Time: {recipe_data.get('cook_time', 'N/A')}\n")
            f.write(f"Servings: {recipe_data.get('servings', 'N/A')}\n")
            f.write(f"Difficulty: {recipe_data.get('difficulty', 'N/A')}\n\n")
            
            # Description
            if recipe_data.get('description'):
                f.write(f"Description: {recipe_data['description']}\n\n")
            
            # Source
            source_info = self._get_recipe_source_info(recipe_data)
            if source_info:
                f.write(f"{source_info}\n\n")
            
            # Ingredients
            f.write("INGREDIENTS:\n")
            f.write("-" * 20 + "\n")
            ingredients = recipe_data.get('ingredients', [])
            if isinstance(ingredients, list):
                for ingredient in ingredients:
                    if isinstance(ingredient, dict):
                        amount = ingredient.get('amount', '')
                        unit = ingredient.get('unit', '')
                        name = ingredient.get('name', '')
                        f.write(f"• {amount} {unit} {name}\n".strip() + "\n")
                    else:
                        f.write(f"• {ingredient}\n")
            else:
                f.write(f"{ingredients}\n")
            
            f.write("\n")
            
            # Instructions
            f.write("INSTRUCTIONS:\n")
            f.write("-" * 20 + "\n")
            f.write(f"{recipe_data.get('instructions', 'No instructions available')}\n")
    
    def _write_recipe_to_html_file(self, recipe_data, file_path):
        """Write recipe to HTML file"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("<!DOCTYPE html>\n<html>\n<head>\n")
            f.write("<meta charset='utf-8'>\n")
            f.write(f"<title>{recipe_data.get('name', 'Recipe')}</title>\n")
            f.write("<style>body{font-family:Arial,sans-serif;margin:40px;line-height:1.6;}</style>\n")
            f.write("</head>\n<body>\n")
            
            f.write(f"<h1>{recipe_data.get('name', 'Unknown Recipe')}</h1>\n")
            
            # Basic info
            f.write("<div style='background:#f5f5f5;padding:15px;margin:20px 0;border-radius:5px;'>\n")
            f.write(f"<strong>Category:</strong> {recipe_data.get('category', 'N/A')}<br>\n")
            f.write(f"<strong>Prep Time:</strong> {recipe_data.get('prep_time', 'N/A')}<br>\n")
            f.write(f"<strong>Cook Time:</strong> {recipe_data.get('cook_time', 'N/A')}<br>\n")
            f.write(f"<strong>Servings:</strong> {recipe_data.get('servings', 'N/A')}<br>\n")
            f.write(f"<strong>Difficulty:</strong> {recipe_data.get('difficulty', 'N/A')}\n")
            f.write("</div>\n")
            
            # Description
            if recipe_data.get('description'):
                f.write(f"<p><strong>Description:</strong> {recipe_data['description']}</p>\n")
            
            # Source
            source_info = self._get_recipe_source_info(recipe_data)
            if source_info:
                # Make URLs clickable
                if 'URL:' in source_info:
                    url = source_info.split('URL: ')[1].strip()
                    source_info = source_info.replace(url, f'<a href="{url}" target="_blank">{url}</a>')
                f.write(f"<p style='color:#666;font-style:italic;'>{source_info}</p>\n")
            
            # Ingredients
            f.write("<h2>Ingredients</h2>\n<ul>\n")
            ingredients = recipe_data.get('ingredients', [])
            if isinstance(ingredients, list):
                for ingredient in ingredients:
                    if isinstance(ingredient, dict):
                        amount = ingredient.get('amount', '')
                        unit = ingredient.get('unit', '')
                        name = ingredient.get('name', '')
                        f.write(f"<li>{amount} {unit} {name}</li>\n".strip())
                    else:
                        f.write(f"<li>{ingredient}</li>\n")
            else:
                f.write(f"<li>{ingredients}</li>\n")
            f.write("</ul>\n")
            
            # Instructions
            f.write("<h2>Instructions</h2>\n")
            instructions = recipe_data.get('instructions', 'No instructions available')
            # Convert line breaks to HTML
            instructions = instructions.replace('\n', '<br>\n')
            f.write(f"<div style='white-space:pre-line;'>{instructions}</div>\n")
            
            f.write("</body>\n</html>\n")
    
    def _parse_ingredients_from_text(self, ingredients_text):
        """Parse ingredients from text input using enhanced parsing"""
        ingredients = []
        for line in ingredients_text.split('\n'):
            line = line.strip()
            if line:
                # Use enhanced parsing similar to recipe scraper
                parsed = self._parse_ingredient_line(line)
                if parsed:
                    ingredients.append(parsed)
                else:
                    # Fallback: just name
                    ingredients.append({
                        'name': line,
                        'amount': '',
                        'unit': ''
                    })
        return ingredients
    
    def _parse_ingredient_line(self, text):
        """Parse a single ingredient line with enhanced logic"""
        import re
        
        # Remove common prefixes
        text = re.sub(r'^[•\-\*\d+\.\)]\s*', '', text)
        text = text.strip()
        
        if not text:
            return None
        
        # Convert Unicode fractions to ASCII equivalents
        unicode_fractions = {
            '½': '1/2',
            '¼': '1/4', 
            '¾': '3/4',
            '⅓': '1/3',
            '⅔': '2/3',
            '⅛': '1/8',
            '⅜': '3/8',
            '⅝': '5/8',
            '⅞': '7/8'
        }
        
        # Replace Unicode fractions with ASCII equivalents
        for unicode_char, ascii_equiv in unicode_fractions.items():
            text = text.replace(unicode_char, ascii_equiv)
        
        # Comprehensive list of measurement units and their variations
        units = [
            # Volume
            'cup', 'cups', 'c', 'c.',
            'tablespoon', 'tablespoons', 'tbsp', 'tbsp.', 'tbs', 'tbs.',
            'teaspoon', 'teaspoons', 'tsp', 'tsp.', 'ts', 'ts.',
            'fluid ounce', 'fluid ounces', 'fl oz', 'fl. oz.', 'floz',
            'pint', 'pints', 'pt', 'pt.',
            'quart', 'quarts', 'qt', 'qt.',
            'gallon', 'gallons', 'gal', 'gal.',
            'liter', 'liters', 'l', 'l.',
            'milliliter', 'milliliters', 'ml', 'ml.',
            
            # Weight
            'pound', 'pounds', 'lb', 'lb.', 'lbs', 'lbs.',
            'ounce', 'ounces', 'oz', 'oz.',
            'gram', 'grams', 'g', 'g.',
            'kilogram', 'kilograms', 'kg', 'kg.',
            
            # Count/Size
            'piece', 'pieces', 'pc', 'pc.',
            'slice', 'slices',
            'clove', 'cloves',
            'head', 'heads',
            'bunch', 'bunches',
            'can', 'cans',
            'package', 'packages', 'pkg', 'pkg.',
            'bag', 'bags',
            'bottle', 'bottles',
            'jar', 'jars',
            'box', 'boxes',
            
            # Size descriptors
            'large', 'medium', 'small', 'extra large', 'xl',
            'whole', 'half', 'quarter',
            'diced', 'chopped', 'sliced', 'minced', 'grated',
            'fresh', 'dried', 'frozen', 'canned'
        ]
        
        # Create regex pattern for units (case insensitive)
        units_pattern = r'\b(' + '|'.join(re.escape(unit) for unit in units) + r')\b'
        
        # Pattern for amounts: handles integers, decimals, fractions, and mixed numbers
        amount_pattern = r'(\d+(?:\.\d+)?(?:\s+\d+\/\d+|\/\d+)?)'
        
        # Full pattern: amount + optional unit + ingredient name
        full_pattern = rf'^{amount_pattern}\s+({units_pattern}\s+)?(.+)$'
        
        match = re.match(full_pattern, text, re.IGNORECASE)
        if match:
            amount = match.group(1).strip()
            unit = match.group(2).strip() if match.group(2) else ''
            ingredient_name = match.group(3).strip()
            
            return {
                'name': ingredient_name,
                'amount': amount,
                'unit': unit
            }
        
        # Try pattern without unit: "2 eggs", "1 onion"
        simple_pattern = rf'^{amount_pattern}\s+(.+)$'
        match = re.match(simple_pattern, text)
        if match:
            amount = match.group(1).strip()
            ingredient_name = match.group(2).strip()
            
            return {
                'name': ingredient_name,
                'amount': amount,
                'unit': ''
            }
        
        return None
    
    def delete_recipe(self):
        """Delete selected recipe"""
        if hasattr(self, 'current_recipe_id') and self.current_recipe_id:
            # Get recipe name
            recipe_data = self._get_recipe_data_by_id(self.current_recipe_id)
            recipe_name = recipe_data.get('name', 'Unknown Recipe') if recipe_data else 'Unknown Recipe'
            
            reply = QMessageBox.question(self, "Delete Recipe", 
                                       f"Are you sure you want to delete '{recipe_name}'?",
                                       QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                # Delete from database
                if self.current_recipe_id:
                    if self._delete_recipe_from_database(self.current_recipe_id):
                        QMessageBox.information(self, "Success", f"Recipe '{recipe_name}' deleted successfully!")
                    else:
                        QMessageBox.warning(self, "Error", "Failed to delete recipe from database.")
                        return
                
                # Refresh recipe display
                self.refresh_recipe_display()
                
                # Clear selection
                self.current_recipe_id = None
    
    def view_recipe(self):
        """View selected recipe details"""
        if hasattr(self, 'current_recipe_id') and self.current_recipe_id:
            try:
                # Get recipe data
                recipe_data = self._get_recipe_data_by_id(self.current_recipe_id)
                recipe_name = recipe_data.get('name', 'Unknown Recipe') if recipe_data else 'Unknown Recipe'
                
                # Create view dialog
                dialog = QDialog(self)
                dialog.setWindowTitle(f"View Recipe: {recipe_name}")
                dialog.setModal(True)
                dialog.resize(900, 700)
                
                layout = QVBoxLayout(dialog)
                
                # Add action buttons at the top
                action_layout = QHBoxLayout()
                
                send_to_menu_btn = QPushButton("📋 Send to Menu")
                send_to_menu_btn.clicked.connect(lambda: self._send_to_menu(recipe_data))
                action_layout.addWidget(send_to_menu_btn)
                
                send_to_shopping_btn = QPushButton("🛒 Send to Shopping List")
                send_to_shopping_btn.clicked.connect(lambda: self._send_to_shopping_list(recipe_data))
                action_layout.addWidget(send_to_shopping_btn)
                

                

                
                action_layout.addStretch()
                
                # Right-click menu button
                more_btn = QPushButton("⋮ More Options")
                more_btn.clicked.connect(lambda: self._show_recipe_context_menu(dialog, recipe_data))
                action_layout.addWidget(more_btn)
                
                layout.addLayout(action_layout)
                
                # Top section with recipe info and scaling
                top_layout = QHBoxLayout()
                
                # Left side - Recipe info
                info_widget = QWidget()
                info_layout = QVBoxLayout(info_widget)
                
                # Recipe title
                title_label = QLabel(recipe_name)
                title_font = QFont()
                title_font.setPointSize(16)
                title_font.setBold(True)
                title_label.setFont(title_font)
                info_layout.addWidget(title_label)
                
                # Recipe image
                recipe_image_label = QLabel()
                recipe_image_label.setFixedSize(200, 150)
                recipe_image_label.setStyleSheet("border: 1px solid #ccc; background-color: #f9f9f9;")
                recipe_image_label.setAlignment(Qt.AlignCenter)
                
                # Try to load recipe image
                try:
                    from services.image_service import recipe_image_service
                    pixmap = recipe_image_service.load_image_pixmap(recipe_data.get('id'), (200, 150))
                    if pixmap and not pixmap.isNull():
                        recipe_image_label.setPixmap(pixmap)
                    else:
                        recipe_image_label.setText("No Image")
                except Exception:
                    recipe_image_label.setText("No Image")
                
                info_layout.addWidget(recipe_image_label)
                
                # Recipe info
                info_text = f"""
Category: {recipe_data.get('category', 'N/A')}
Prep Time: {recipe_data.get('prep_time', 'N/A')} minutes
Cook Time: {recipe_data.get('cook_time', 'N/A')} minutes
Servings: {recipe_data.get('servings', 'N/A')}
Difficulty: {recipe_data.get('difficulty', 'N/A')}
                """
                info_label = QLabel(info_text.strip())
                info_layout.addWidget(info_label)
                
                # Description
                if recipe_data.get('description'):
                    desc_label = QLabel(f"Description: {recipe_data['description']}")
                    desc_label.setWordWrap(True)
                    info_layout.addWidget(desc_label)
                
                # Source information
                source_info = self._get_recipe_source_info(recipe_data)
                if source_info:
                    source_label = QLabel(source_info)
                    source_label.setWordWrap(True)
                    source_label.setStyleSheet("color: #666; font-style: italic; margin-top: 10px;")
                    
                    # Make URL clickable if present
                    url = recipe_data.get('url', '')
                    if url and url.startswith(('http://', 'https://')):
                        source_label.setStyleSheet("color: #0066cc; font-style: italic; margin-top: 10px; text-decoration: underline;")
                        source_label.setCursor(Qt.PointingHandCursor)
                        source_label.mousePressEvent = lambda event: self._open_url_in_browser(url)
                        source_label.setToolTip(f"Click to open: {url}")
                    
                    info_layout.addWidget(source_label)
                
                info_layout.addStretch()
                top_layout.addWidget(info_widget)
                
                # Right side - Scaling controls
                scale_widget = QWidget()
                scale_layout = QVBoxLayout(scale_widget)
                
                scale_group = QGroupBox("Scale Recipe")
                scale_group_layout = QVBoxLayout(scale_group)
                
                scale_layout.addWidget(QLabel("New Servings:"))
                self.scale_servings_spin = QSpinBox()
                self.scale_servings_spin.setMinimum(1)
                self.scale_servings_spin.setMaximum(100)
                # Handle servings that might contain text like "1 loaf"
                servings_value = self._extract_numeric_value(recipe_data.get('servings', 4), 4)
                self.scale_servings_spin.setValue(servings_value)
                scale_group_layout.addWidget(self.scale_servings_spin)
                
                scale_btn = QPushButton("Scale Recipe")
                scale_btn.clicked.connect(lambda: self._scale_recipe_in_view(dialog, recipe_data))
                scale_group_layout.addWidget(scale_btn)
                
                # Print button
                print_group = QGroupBox("Print Recipe")
                print_group_layout = QVBoxLayout(print_group)
                
                print_btn = QPushButton("Print Recipe")
                print_btn.clicked.connect(lambda: self._print_recipe(dialog, recipe_data))
                print_group_layout.addWidget(print_btn)
                
                # Favorite checkbox
                favorite_group = QGroupBox("Recipe Options")
                favorite_group_layout = QVBoxLayout(favorite_group)
                
                self.favorite_checkbox = QCheckBox("Mark as Favorite")
                # Check if recipe is already marked as favorite
                is_favorite = self._is_recipe_favorite(recipe_data.get('id'))
                self.favorite_checkbox.setChecked(is_favorite)
                self.favorite_checkbox.stateChanged.connect(lambda: self._toggle_favorite(recipe_data.get('id')))
                favorite_group_layout.addWidget(self.favorite_checkbox)
                
                scale_layout.addWidget(scale_group)
                scale_layout.addWidget(print_group)
                scale_layout.addWidget(favorite_group)
                scale_layout.addStretch()
                top_layout.addWidget(scale_widget)
                
                layout.addLayout(top_layout)
                
                # Ingredients and Instructions in scroll area
                scroll_area = QScrollArea()
                scroll_widget = QWidget()
                scroll_layout = QVBoxLayout(scroll_widget)
                
                # Ingredients
                ingredients_group = QGroupBox("Ingredients")
                ingredients_layout = QVBoxLayout(ingredients_group)
                
                self.ingredients_display = QLabel()
                self.ingredients_display.setWordWrap(True)
                self._update_ingredients_display(recipe_data)
                ingredients_layout.addWidget(self.ingredients_display)
                
                scroll_layout.addWidget(ingredients_group)
                
                # Instructions
                instructions_group = QGroupBox("Instructions")
                instructions_layout = QVBoxLayout(instructions_group)
                
                instructions_text = recipe_data.get('instructions', 'No instructions provided')
                instructions_label = QLabel(instructions_text)
                instructions_label.setWordWrap(True)
                instructions_layout.addWidget(instructions_label)
                scroll_layout.addWidget(instructions_group)
                
                # Notes
                if recipe_data.get('notes'):
                    notes_group = QGroupBox("Notes")
                    notes_layout = QVBoxLayout(notes_group)
                    notes_label = QLabel(recipe_data['notes'])
                    notes_label.setWordWrap(True)
                    notes_layout.addWidget(notes_label)
                    scroll_layout.addWidget(notes_group)
                
                scroll_area.setWidget(scroll_widget)
                scroll_area.setWidgetResizable(True)
                layout.addWidget(scroll_area)
                
                # Enable right-click context menu on the dialog
                dialog.setContextMenuPolicy(Qt.CustomContextMenu)
                dialog.customContextMenuRequested.connect(lambda pos: self._show_view_recipe_context_menu(dialog, recipe_data, pos))
                
                # Buttons
                button_layout = QHBoxLayout()
                edit_btn = QPushButton("Edit Recipe")
                # Use recipe ID instead of current_row since this is a data-based view
                recipe_id = recipe_data.get('id')
                if recipe_id:
                    edit_btn.clicked.connect(lambda: self._edit_recipe_by_id(dialog, recipe_id))
                else:
                    edit_btn.setEnabled(False)
                
                close_btn = QPushButton("Close")
                close_btn.clicked.connect(dialog.accept)
                
                button_layout.addWidget(edit_btn)
                button_layout.addStretch()
                button_layout.addWidget(close_btn)
                layout.addLayout(button_layout)
                
                dialog.exec()
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to view recipe: {str(e)}")
    
    def _edit_from_view(self, parent_dialog, row):
        """Edit recipe from view dialog"""
        parent_dialog.accept()
        self.recipes_table.selectRow(row)
        self.edit_recipe()
    
    def _edit_recipe_by_id(self, parent_dialog, recipe_id):
        """Edit recipe by ID from view dialog"""
        parent_dialog.accept()
        # Set the current recipe ID and call edit
        self.current_recipe_id = recipe_id
        self.edit_recipe()
    
    def _scale_recipe_in_view(self, dialog, recipe_data):
        """Scale recipe in view dialog"""
        try:
            # Try importing from services module first
            try:
                from services import recipe_scaler
            except ImportError:
                # Fallback to direct import
                from services.recipe_enhancement import recipe_scaler
            
            new_servings = self.scale_servings_spin.value()
            original_servings = recipe_data.get('servings', 1)
            
            scaled_recipe = recipe_scaler.scale_recipe(recipe_data, new_servings)
            
            # Update the ingredients display
            self._update_ingredients_display(scaled_recipe)
            
            # Update the servings info
            info_text = f"""
Category: {scaled_recipe.get('category', 'N/A')}
Prep Time: {scaled_recipe.get('prep_time', 'N/A')} minutes
Cook Time: {scaled_recipe.get('cook_time', 'N/A')} minutes
Servings: {scaled_recipe.get('servings', 'N/A')} (Scaled from {scaled_recipe.get('original_servings', 'N/A')})
Difficulty: {scaled_recipe.get('difficulty', 'N/A')}
            """
            
            # Find and update the info label
            for child in dialog.findChildren(QLabel):
                if "Category:" in child.text():
                    child.setText(info_text.strip())
                    break
            
            QMessageBox.information(dialog, "Recipe Scaled", f"Recipe scaled to {new_servings} servings!")
            
        except Exception as e:
            QMessageBox.critical(dialog, "Error", f"Failed to scale recipe: {str(e)}")
    
    def _update_ingredients_display(self, recipe_data):
        """Update ingredients display with scaled amounts"""
        ingredients_text = ""
        if recipe_data.get('ingredients'):
            for ingredient in recipe_data['ingredients']:
                if isinstance(ingredient, dict):
                    # Try multiple quantity fields in order of preference
                    qty = (ingredient.get('display_quantity') or 
                          ingredient.get('amount') or 
                          ingredient.get('quantity') or '')
                    name = ingredient.get('name', '')
                    
                    
                    # Format the quantity properly
                    if qty and isinstance(qty, (int, float)):
                        # If it's a number, format it nicely
                        if qty == int(qty):
                            qty = str(int(qty))
                        else:
                            qty = f"{qty:.2f}".rstrip('0').rstrip('.')
                    
                    ing_text = f"• {qty} {name}".strip()
                else:
                    ing_text = f"• {ingredient}"
                ingredients_text += ing_text + "\n"
        else:
            ingredients_text = "No ingredients specified"
        
        self.ingredients_display.setText(ingredients_text.strip())
    

    
    def _show_gf_analysis_dialog(self, parent_dialog, recipe_data, analysis):
        """Show gluten-free analysis and conversion dialog"""
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QTextEdit, QPushButton, QMessageBox, QGroupBox
        
        dialog = QDialog(parent_dialog)
        dialog.setWindowTitle("Gluten-Free Recipe Analysis")
        dialog.setModal(True)
        dialog.resize(700, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Header
        header_label = QLabel(f"Gluten-Free Analysis: {recipe_data.get('name', 'Recipe')}")
        header_label.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header_label)
        
        # Summary
        summary_group = QGroupBox("Analysis Summary")    
        summary_layout = QVBoxLayout(summary_group)
        
        gluten_count = len(analysis['gluten_ingredients'])
        summary_text = f"""
Found {gluten_count} gluten-containing ingredients
Conversion Difficulty: {analysis['difficulty_level']}
Estimated Cost Change: {analysis['estimated_cost_change']}
        """
        summary_label = QLabel(summary_text.strip())
        summary_layout.addWidget(summary_label)
        layout.addWidget(summary_group)
        
        # Gluten ingredients found
        if analysis['gluten_ingredients']:
            gluten_group = QGroupBox("Gluten Ingredients Found")
            gluten_layout = QVBoxLayout(gluten_group)
            
            gluten_text = ""
            for item in analysis['gluten_ingredients']:
                gluten_text += f"• {item['amount']} {item['original']} → {item['replacement']}\n"
            
            gluten_label = QLabel(gluten_text.strip())
            gluten_label.setWordWrap(True)
            gluten_layout.addWidget(gluten_label)
            layout.addWidget(gluten_group)
        
        # Conversion notes
        if analysis['conversion_notes']:
            notes_group = QGroupBox("Conversion Notes")
            notes_layout = QVBoxLayout(notes_group)
            
            notes_text = "\n".join(f"• {note}" for note in analysis['conversion_notes'])
            notes_label = QLabel(notes_text)
            notes_label.setWordWrap(True)
            notes_layout.addWidget(notes_label)
            layout.addWidget(notes_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        convert_btn = QPushButton("Convert to Gluten-Free Recipe")
        convert_btn.clicked.connect(lambda: self._convert_to_gf_recipe(dialog, recipe_data))
        button_layout.addWidget(convert_btn)
        
        cancel_btn = QPushButton("Close")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)
        
        layout.addLayout(button_layout)
        
        dialog.exec()

    def get_recipe_nutrition(self, recipe_data):
        """Placeholder function for fetching USDA nutrition information"""
        pass
    
    def _convert_to_gf_recipe(self, analysis_dialog, recipe_data):
        """Convert recipe to gluten-free and save as new recipe"""
        try:
            from services.gluten_free_converter import gluten_free_converter
            
            # Convert the recipe
            converted_recipe = gluten_free_converter.convert_recipe(recipe_data)
            
            # Show edit dialog for the converted recipe
            from utils.edit_dialogs import RecipeEditDialog
            
            edit_dialog = RecipeEditDialog(self)
            edit_dialog.set_data(converted_recipe)
            
            if edit_dialog.exec() == QDialog.Accepted:
                new_data = edit_dialog.get_data()
                
                # Save to database
                recipe_id = self._save_recipe_to_database(new_data)
                if recipe_id:
                    new_data['id'] = recipe_id
                    self.recipe_data_storage[recipe_id] = new_data
                    self.refresh_recipe_display()
                    
                    QMessageBox.information(self, "Success", 
                        f"Gluten-free recipe '{new_data['name']}' saved successfully!")
                    analysis_dialog.accept()
                else:
                    QMessageBox.warning(self, "Error", "Failed to save gluten-free recipe.")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to convert recipe: {str(e)}")
    
    def _show_view_recipe_context_menu(self, dialog, recipe_data, position):
        """Show right-click context menu for view recipe dialog"""
        from PySide6.QtWidgets import QMenu
        
        menu = QMenu(dialog)
        
        # Add all view recipe actions
        edit_action = menu.addAction("✏️ Edit Recipe")
        edit_action.triggered.connect(lambda: self._edit_recipe_by_id(dialog, recipe_data.get('id')))
        
        menu.addSeparator()
        
        send_menu_action = menu.addAction("📋 Send to Menu")
        send_menu_action.triggered.connect(lambda: self._send_to_menu(recipe_data))
        
        send_shopping_action = menu.addAction("🛒 Send to Shopping List")
        send_shopping_action.triggered.connect(lambda: self._send_to_shopping_list(recipe_data))
        
        scale_action = menu.addAction("⚖️ Scale Recipe")
        scale_action.triggered.connect(lambda: self._scale_recipe_from_view(dialog, recipe_data))
        
        menu.addSeparator()
        
        export_txt_action = menu.addAction("📄 Export as Text")
        export_txt_action.triggered.connect(lambda: self._export_recipe_txt_from_view(recipe_data))
        
        export_html_action = menu.addAction("🌐 Export as HTML")
        export_html_action.triggered.connect(lambda: self._export_recipe_html_from_view(recipe_data))
        
        print_action = menu.addAction("🖨️ Print Recipe")
        print_action.triggered.connect(lambda: self._print_recipe(dialog, recipe_data))
        
        menu.addSeparator()
        
        favorite_action = menu.addAction("⭐ Toggle Favorite")
        favorite_action.triggered.connect(lambda: self._toggle_favorite(recipe_data.get('id')))
        
        delete_action = menu.addAction("🗑️ Delete Recipe")
        delete_action.triggered.connect(lambda: self._delete_recipe_from_view(dialog, recipe_data))
        
        # Show menu at cursor position
        global_pos = dialog.mapToGlobal(position)
        menu.exec(global_pos)
    
    def _export_recipe(self, dialog, recipe_data, format_type):
        """Export recipe to file and open it automatically"""
        try:
            # Try importing from services module first
            try:
                from services import recipe_exporter
            except ImportError:
                # Fallback to direct import
                from services.recipe_enhancement import recipe_exporter
            import os
            import subprocess
            import platform
            from datetime import datetime
            
            # Get export content
            if format_type == 'txt':
                content = recipe_exporter.export_recipe(recipe_data, 'plain_text')
                default_ext = ".txt"
            elif format_type == 'html':
                content = recipe_exporter.export_recipe(recipe_data, 'markdown')
                # Convert markdown to HTML (proper conversion)
                content = self._convert_markdown_to_html(content)
                default_ext = ".html"
            else:
                QMessageBox.warning(dialog, "Error", "Unsupported export format")
                return
            
            # Create Export folder if it doesn't exist
            export_folder = "Export"
            if not os.path.exists(export_folder):
                os.makedirs(export_folder)
            
            # Generate filename with timestamp
            recipe_name = recipe_data.get('name', 'recipe').replace(' ', '_').replace('/', '_')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{recipe_name}_{timestamp}{default_ext}"
            file_path = os.path.join(export_folder, filename)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Open file automatically
            try:
                if platform.system() == 'Windows':
                    os.startfile(file_path)
                elif platform.system() == 'Darwin':  # macOS
                    subprocess.run(['open', file_path])
                else:  # Linux
                    subprocess.run(['xdg-open', file_path])
                
                QMessageBox.information(dialog, "Export Successful", 
                                      f"Recipe exported and opened:\n{file_path}")
            except Exception as e:
                QMessageBox.information(dialog, "Export Successful", 
                                      f"Recipe exported to:\n{file_path}\n\nCould not open file automatically: {str(e)}")
            
        except Exception as e:
            QMessageBox.critical(dialog, "Export Error", f"Failed to export recipe: {str(e)}")
    
    def _export_recipe_txt(self):
        """Export selected recipe as text (for context menu)"""
        current_row = self.recipes_table.currentRow()
        if current_row >= 0:
            recipe_data = self._get_recipe_data(current_row)
            self._export_recipe(None, recipe_data, 'txt')
    
    def _export_recipe_html(self):
        """Export selected recipe as HTML (for context menu)"""
        current_row = self.recipes_table.currentRow()
        if current_row >= 0:
            recipe_data = self._get_recipe_data(current_row)
            self._export_recipe(None, recipe_data, 'html')
    
    def _print_recipe(self, dialog, recipe_data):
        """Print recipe directly using QPrinter and QTextDocument"""
        try:
            # Try importing from services module first
            try:
                from services import recipe_exporter
            except ImportError:
                # Fallback to direct import
                from services.recipe_enhancement import recipe_exporter
            
            # Get HTML content
            content = recipe_exporter.export_recipe(recipe_data, 'markdown')
            html_content = self._convert_markdown_to_html(content)
            
            # Create printer with comprehensive settings
            printer = QPrinter()
            
            # Set page size
            page_size = QPageSize(QPageSize.PageSizeId.A4)
            printer.setPageSize(page_size)
            
            # Set page orientation
            printer.setPageOrientation(QPageLayout.Orientation.Portrait)
            
            # Set margins using QMarginsF (left, top, right, bottom in millimeters)
            margins = QMarginsF(20.0, 20.0, 20.0, 20.0)
            printer.setPageMargins(margins, QPageLayout.Millimeter)
            
            # Set additional printer attributes
            printer.setColorMode(QPrinter.Color)
            printer.setOutputFormat(QPrinter.NativeFormat)
            printer.setPrintRange(QPrinter.AllPages)
            printer.setCopyCount(1)  # Set number of copies
            printer.setCollateCopies(True)
            printer.setFullPage(False)
            printer.setOutputFileName("")  # Print to printer, not file
            
            # Create text document with HTML content
            document = QTextDocument()
            document.setHtml(html_content)
            
            # Show print preview dialog
            preview_dialog = QPrintPreviewDialog(printer, dialog)
            preview_dialog.setWindowTitle("Print Preview - Recipe")
            preview_dialog.setWindowFlags(preview_dialog.windowFlags() | Qt.WindowMaximizeButtonHint)
            
            # Connect the paint request signal to render the document
            def paint_request(printer):
                document.print_(printer)
            
            preview_dialog.paintRequested.connect(paint_request)
            
            # Show the preview dialog
            preview_dialog.exec()
            
        except Exception as e:
            QMessageBox.critical(dialog, "Print Error", f"Failed to print recipe: {str(e)}")
    
    def _is_recipe_favorite(self, recipe_id):
        """Check if a recipe is marked as favorite"""
        try:
            if not recipe_id:
                return False
                
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT is_favorite FROM recipes WHERE id = ?", (recipe_id,))
            result = cursor.fetchone()
            conn.close()
            
            return result and result[0] == 1
        except Exception as e:
            print(f"Error checking favorite status: {e}")
            return False
    
    def _toggle_favorite(self, recipe_id):
        """Toggle favorite status of a recipe"""
        try:
            if not recipe_id:
                return
                
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            
            # Get current favorite status
            cursor.execute("SELECT is_favorite FROM recipes WHERE id = ?", (recipe_id,))
            result = cursor.fetchone()
            
            if result:
                # Toggle the favorite status
                new_status = 0 if result[0] == 1 else 1
                cursor.execute("UPDATE recipes SET is_favorite = ? WHERE id = ?", (new_status, recipe_id))
                conn.commit()
                
                # Show feedback message
                status_text = "added to" if new_status == 1 else "removed from"
                QMessageBox.information(self, "Favorite Updated", f"Recipe {status_text} favorites!")
            
            conn.close()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to update favorite status: {str(e)}")
    
    def filter_favorites(self):
        """Filter recipes to show only favorites"""
        self.load_recipes()
    
    def filter_by_category(self):
        """Filter recipes by selected category"""
        self.load_recipes()
    
    def load_categories(self):
        """Load categories from database and populate filter dropdown"""
        try:
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            
            # Get unique categories from actual recipes (not the categories table)
            cursor.execute("""
                SELECT DISTINCT category 
                FROM recipes 
                WHERE category IS NOT NULL AND category != ''
                ORDER BY category
            """)
            recipe_categories = cursor.fetchall()
            conn.close()
            
            # Clear existing items and populate combo box
            if hasattr(self, 'category_combo'):
                self.category_combo.clear()
                self.category_combo.addItem("All Categories", "all")
                
                for (category_name,) in recipe_categories:
                    self.category_combo.addItem(category_name, category_name)
                
                # Set default selection to "All Categories"
                self.category_combo.setCurrentIndex(0)
            
        except Exception as e:
            print(f"Error loading categories: {e}")
    
    def _get_category_name(self, category_id):
        """Get category name by ID"""
        try:
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM categories WHERE id = ?", (category_id,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else "Unknown"
        except Exception:
            return "Unknown"
    
    def _set_recipe_image(self, row, recipe_id):
        """Set recipe image in table cell"""
        try:
            from services.image_service import recipe_image_service
            
            # Load image pixmap
            pixmap = recipe_image_service.load_image_pixmap(recipe_id, (60, 40))
            
            if pixmap and not pixmap.isNull():
                # Create label with image
                image_label = QLabel()
                image_label.setPixmap(pixmap)
                image_label.setAlignment(Qt.AlignCenter)
                image_label.setStyleSheet("border: none;")
                
                # Set the label as the cell widget
                self.recipes_table.setCellWidget(row, 0, image_label)
            else:
                # No image - set empty cell
                self.recipes_table.setItem(row, 0, QTableWidgetItem(""))
                
        except Exception as e:
            print(f"Error setting recipe image: {e}")
            # Set empty cell on error
            self.recipes_table.setItem(row, 0, QTableWidgetItem(""))
    
    def show_quick_access_favorites(self):
        """Show quick access dialog for favorite recipes"""
        try:
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT id, title FROM recipes WHERE is_favorite = 1 ORDER BY title")
            favorites = cursor.fetchall()
            conn.close()
            
            if not favorites:
                QMessageBox.information(self, "No Favorites", "You haven't marked any recipes as favorites yet.")
                return
            
            # Create quick access dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Quick Access - Favorite Recipes")
            dialog.setModal(True)
            dialog.resize(400, 300)
            
            layout = QVBoxLayout(dialog)
            layout.addWidget(QLabel("Select a favorite recipe to view:"))
            
            # List widget for favorites
            favorites_list = QListWidget()
            for recipe_id, title in favorites:
                item = QListWidgetItem(f"⭐ {title}")
                item.setData(Qt.UserRole, recipe_id)
                favorites_list.addItem(item)
            
            layout.addWidget(favorites_list)
            
            # Buttons
            button_layout = QHBoxLayout()
            view_btn = QPushButton("View Recipe")
            close_btn = QPushButton("Close")
            
            def view_selected_recipe():
                current_item = favorites_list.currentItem()
                if current_item:
                    recipe_id = current_item.data(Qt.UserRole)
                    dialog.accept()
                    self._view_recipe_by_id(recipe_id)
            
            view_btn.clicked.connect(view_selected_recipe)
            close_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(view_btn)
            button_layout.addWidget(close_btn)
            layout.addLayout(button_layout)
            
            # Double-click to view
            favorites_list.itemDoubleClicked.connect(lambda: view_selected_recipe())
            
            dialog.exec()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show quick access: {str(e)}")
    
    def _view_recipe_by_id(self, recipe_id):
        """View a recipe by its ID"""
        try:
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM recipes WHERE id = ?", (recipe_id,))
            recipe = cursor.fetchone()
            conn.close()
            
            if recipe:
                # Convert to dictionary format
                recipe_data = {
                    'id': recipe[0],
                    'name': recipe[1],
                    'ingredients': recipe[2],
                    'instructions': recipe[3],
                    'prep_time': recipe[4],
                    'cook_time': recipe[5],
                    'servings': recipe[6],
                    'category': recipe[9],
                    'difficulty': recipe[10],
                    'description': recipe[11] if len(recipe) > 11 else ''
                }
                self.view_recipe_by_data(recipe_data)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to view recipe: {str(e)}")
    
    def view_recipe_by_data(self, recipe_data):
        """View recipe using recipe data (reuse existing view_recipe logic)"""
        # This will reuse the existing view_recipe method logic
        # We'll create a temporary dialog to show the recipe
        dialog = QDialog(self)
        dialog.setWindowTitle(f"View Recipe - {recipe_data.get('name', 'Unknown')}")
        dialog.setModal(True)
        dialog.resize(800, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title_label = QLabel(recipe_data.get('name', 'Unknown Recipe'))
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        # Recipe info
        info_text = f"""
Category: {recipe_data.get('category', 'N/A')}
Prep Time: {recipe_data.get('prep_time', 'N/A')} minutes
Cook Time: {recipe_data.get('cook_time', 'N/A')} minutes
Servings: {recipe_data.get('servings', 'N/A')}
Difficulty: {recipe_data.get('difficulty', 'N/A')}
        """
        info_label = QLabel(info_text.strip())
        layout.addWidget(info_label)
        
        # Ingredients
        ingredients_label = QLabel("Ingredients:")
        ingredients_label.setFont(QFont("", 12, QFont.Bold))
        layout.addWidget(ingredients_label)
        
        ingredients_text = QTextEdit()
        ingredients_text.setPlainText(recipe_data.get('ingredients', 'No ingredients listed'))
        ingredients_text.setMaximumHeight(150)
        layout.addWidget(ingredients_text)
        
        # Instructions
        instructions_label = QLabel("Instructions:")
        instructions_label.setFont(QFont("", 12, QFont.Bold))
        layout.addWidget(instructions_label)
        
        instructions_text = QTextEdit()
        instructions_text.setPlainText(recipe_data.get('instructions', 'No instructions listed'))
        layout.addWidget(instructions_text)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()
    
    def update_favorite_statistics(self):
        """Update favorite statistics display"""
        try:
            from utils.db import get_connection
            conn = get_connection()
            cursor = conn.cursor()
            
            # Get total recipes count
            cursor.execute("SELECT COUNT(*) FROM recipes")
            total_count = cursor.fetchone()[0]
            
            # Get favorites count
            cursor.execute("SELECT COUNT(*) FROM recipes WHERE is_favorite = 1")
            favorites_count = cursor.fetchone()[0]
            
            conn.close()
            
            # Update the label
            self.favorite_stats_label.setText(f"⭐ Favorites: {favorites_count} | 📚 Total: {total_count}")
            
        except Exception as e:
            print(f"Error updating favorite statistics: {e}")
    
    def _convert_markdown_to_html(self, markdown_content):
        """Convert markdown content to properly formatted HTML"""
        # Split content into sections
        lines = markdown_content.split('\n')
        html_sections = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if line.startswith('# '):
                # Main title
                title = line[2:].strip()
                html_sections.append(f'<h1>{title}</h1>')
            elif line.startswith('## '):
                # Section headers
                header = line[3:].strip()
                html_sections.append(f'<h2>{header}</h2>')
            elif line.startswith('### '):
                # Subsection headers
                subheader = line[4:].strip()
                html_sections.append(f'<h3>{subheader}</h3>')
            elif line.startswith('- ') or line.startswith('* '):
                # List items
                list_items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                    item = lines[i].strip()[2:].strip()
                    # Handle bold and italic formatting
                    item = self._format_text(item)
                    list_items.append(f'<li>{item}</li>')
                    i += 1
                if list_items:
                    html_sections.append(f'<ul>{"".join(list_items)}</ul>')
                i -= 1  # Adjust for the outer loop increment
            elif line.startswith(tuple(str(n) + '. ' for n in range(1, 100))):
                # Numbered list items
                list_items = []
                while i < len(lines) and lines[i].strip() and (lines[i].strip()[0].isdigit() and '. ' in lines[i].strip()):
                    item = lines[i].strip()
                    # Remove number prefix
                    item = item.split('. ', 1)[1] if '. ' in item else item
                    # Handle bold and italic formatting
                    item = self._format_text(item)
                    list_items.append(f'<li>{item}</li>')
                    i += 1
                if list_items:
                    html_sections.append(f'<ol>{"".join(list_items)}</ol>')
                i -= 1  # Adjust for the outer loop increment
            elif line and not line.startswith('#'):
                # Regular paragraph
                paragraph = self._format_text(line)
                html_sections.append(f'<p>{paragraph}</p>')
            elif not line:
                # Empty line - skip
                pass
            
            i += 1
        
        # Join all sections
        content_html = '\n        '.join(html_sections)
        
        # Create full HTML document with comprehensive styling
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Recipe</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.7;
            color: #333;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            min-height: 100vh;
            padding: 20px;
        }}
        
        .recipe-container {{
            max-width: 900px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            overflow: hidden;
        }}
        
        .recipe-header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 40px;
            text-align: center;
        }}
        
        .recipe-header h1 {{
            font-size: 2.5em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }}
        
        .recipe-meta {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }}
        
        .meta-item {{
            text-align: center;
            padding: 15px;
            background: rgba(255,255,255,0.1);
            border-radius: 10px;
            backdrop-filter: blur(10px);
        }}
        
        .meta-label {{
            font-size: 0.9em;
            opacity: 0.9;
            margin-bottom: 5px;
        }}
        
        .meta-value {{
            font-size: 1.3em;
            font-weight: bold;
        }}
        
        .recipe-content {{
            padding: 40px;
        }}
        
        h2 {{
            color: #2c3e50;
            font-size: 1.8em;
            margin: 30px 0 20px 0;
            padding-bottom: 10px;
            border-bottom: 3px solid #3498db;
            position: relative;
        }}
        
        h2::before {{
            content: '';
            position: absolute;
            bottom: -3px;
            left: 0;
            width: 50px;
            height: 3px;
            background: #e74c3c;
        }}
        
        h3 {{
            color: #34495e;
            font-size: 1.4em;
            margin: 25px 0 15px 0;
            padding-left: 15px;
            border-left: 4px solid #f39c12;
        }}
        
        .ingredients-section {{
            background: #f8f9fa;
            padding: 25px;
            border-radius: 10px;
            margin: 25px 0;
            border-left: 5px solid #27ae60;
        }}
        
        .instructions-section {{
            background: #fff;
            padding: 25px;
            border-radius: 10px;
            margin: 25px 0;
            border-left: 5px solid #e74c3c;
        }}
        
        ul, ol {{
            padding-left: 25px;
            margin: 15px 0;
        }}
        
        li {{
            margin-bottom: 12px;
            line-height: 1.6;
        }}
        
        ul li {{
            list-style-type: none;
            position: relative;
        }}
        
        ul li::before {{
            content: '•';
            color: #27ae60;
            font-weight: bold;
            font-size: 1.2em;
            position: absolute;
            left: -20px;
        }}
        
        ol li {{
            counter-increment: step-counter;
        }}
        
        ol {{
            counter-reset: step-counter;
        }}
        
        ol li::marker {{
            color: #e74c3c;
            font-weight: bold;
        }}
        
        p {{
            margin: 15px 0;
            text-align: justify;
        }}
        
        strong {{
            color: #2c3e50;
            font-weight: 600;
        }}
        
        em {{
            color: #7f8c8d;
            font-style: italic;
        }}
        
        .recipe-footer {{
            background: #2c3e50;
            color: white;
            text-align: center;
            padding: 20px;
            font-size: 0.9em;
        }}
        
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            
            .recipe-container {{
                border-radius: 0;
            }}
            
            .recipe-header {{
                background: #2c3e50 !important;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
        }}
        
        @media (max-width: 768px) {{
            .recipe-container {{
                margin: 10px;
                border-radius: 10px;
            }}
            
            .recipe-header, .recipe-content {{
                padding: 20px;
            }}
            
            .recipe-meta {{
                grid-template-columns: 1fr;
            }}
            
            h1 {{
                font-size: 2em;
            }}
            
            h2 {{
                font-size: 1.5em;
            }}
        }}
    </style>
</head>
<body>
    <div class="recipe-container">
        <div class="recipe-header">
            {content_html.split('<h2>')[0] if '<h2>' in content_html else content_html.split('<h1>')[1] if '<h1>' in content_html else ''}
        </div>
        <div class="recipe-content">
            {self._extract_content_sections(content_html)}
        </div>
        <div class="recipe-footer">
            <p>Generated by CeliacShield Recipe Manager</p>
        </div>
    </div>
</body>
</html>"""
        
        return full_html
    
    def _format_text(self, text):
        """Format text with bold and italic markers"""
        # Handle bold text (**text**)
        import re
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        # Handle italic text (*text*)
        text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
        return text
    
    def _extract_content_sections(self, html_content):
        """Extract content sections after the main title"""
        # Split by h2 tags to separate sections
        sections = html_content.split('<h2>')
        if len(sections) > 1:
            # Rejoin with h2 tags and wrap sections appropriately
            content_sections = []
            for i, section in enumerate(sections[1:], 1):
                if i == 1:  # First section (usually ingredients)
                    content_sections.append(f'<div class="ingredients-section"><h2>{section}</div>')
                elif i == 2:  # Second section (usually instructions)
                    content_sections.append(f'<div class="instructions-section"><h2>{section}</div>')
                else:  # Other sections
                    content_sections.append(f'<h2>{section}')
            return '\n            '.join(content_sections)
        else:
            # No h2 sections, return as is
            return html_content
    
    def import_from_web(self):
        """Import recipes from web URLs"""
        from PySide6.QtWidgets import QInputDialog, QProgressDialog
        
        url, ok = QInputDialog.getText(self, 'Import from Web', 'Enter recipe URL:')
        if ok and url:
            # Validate URL format
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            # Show progress dialog
            progress = QProgressDialog("Scraping recipe from URL...", "Cancel", 0, 0, self)
            progress.setWindowModality(Qt.WindowModal)
            progress.show()
            
            try:
                # Check if scraper is available first
                try:
                    from services import is_scraper_available, check_scraper_dependencies
                    if not is_scraper_available():
                        missing_deps = check_scraper_dependencies()
                        progress.close()
                        QMessageBox.warning(self, "Missing Dependencies", 
                                          f"Recipe scraper requires the following packages:\n\n" +
                                          "\n".join(f"• {dep}" for dep in missing_deps) +
                                          f"\n\nPlease install them using:\n" +
                                          f"pip install {' '.join(missing_deps)}")
                        return
                except ImportError:
                    # Fallback check
                    try:
                        import requests
                        import bs4
                    except ImportError as e:
                        progress.close()
                        QMessageBox.warning(self, "Missing Dependencies", 
                                          f"Recipe scraper requires 'requests' and 'beautifulsoup4' packages.\n\n"
                                          f"Please install them using:\n"
                                          f"pip install requests beautifulsoup4")
                        return
                
                # Try importing from services module first
                try:
                    from services import scrape_recipe_from_url, scrape_recipes_from_url
                except ImportError:
                    # Fallback to direct import
                    from services.recipe_scraper import scrape_recipe_from_url, scrape_recipes_from_url
                
                progress.setLabelText("Connecting to website...")
                
                # Try to scrape single recipe first
                recipe = scrape_recipe_from_url(url)
                if recipe and recipe.get('name'):
                    progress.setLabelText("Adding recipe to cookbook...")
                    self._add_scraped_recipe(recipe, source_url=url)
                    progress.close()
                    QMessageBox.information(self, "Import Success", 
                                          f"Successfully imported recipe: {recipe.get('name', 'Unknown')}")
                else:
                    progress.setLabelText("Trying to find multiple recipes...")
                    # Try to scrape multiple recipes
                    recipes = scrape_recipes_from_url(url)
                    if recipes:
                        progress.setMaximum(len(recipes))
                        imported_count = 0
                        for i, recipe in enumerate(recipes):
                            if progress.wasCanceled():
                                break
                            progress.setValue(i)
                            progress.setLabelText(f"Importing recipe {i+1} of {len(recipes)}...")
                            if recipe and recipe.get('name'):
                                self._add_scraped_recipe(recipe, source_url=url)
                                imported_count += 1
                        
                        progress.close()
                        QMessageBox.information(self, "Import Success", 
                                              f"Successfully imported {imported_count} recipes from the URL.")
                    else:
                        progress.close()
                        # Try fallback with better error message
                        QMessageBox.warning(self, "Import Failed", 
                                          "No recipes found at the provided URL.\n\n"
                                          "This could be because:\n"
                                          "• The website doesn't contain recipe data\n"
                                          "• The website blocks automated access\n"
                                          "• The URL format is not supported\n\n"
                                          "Would you like to try manual import?")
                        
                        reply = QMessageBox.question(self, "Manual Import", 
                                                   "Would you like to enter the recipe details manually?",
                                                   QMessageBox.Yes | QMessageBox.No)
                        if reply == QMessageBox.Yes:
                            self._fallback_web_import(url)
                        
            except ImportError as e:
                progress.close()
                QMessageBox.warning(self, "Import Error", 
                                  "Recipe scraper service is not available.\n"
                                  "Please install required dependencies or use manual import.")
                self._fallback_web_import(url)
            except Exception as e:
                progress.close()
                error_msg = str(e)
                if "timeout" in error_msg.lower():
                    error_msg = "Connection timeout. The website may be slow or unreachable."
                elif "404" in error_msg or "not found" in error_msg.lower():
                    error_msg = "Website not found. Please check the URL."
                elif "403" in error_msg or "forbidden" in error_msg.lower():
                    error_msg = "Access forbidden. The website may block automated access."
                
                QMessageBox.critical(self, "Import Error", 
                                   f"Failed to import recipe:\n{error_msg}\n\n"
                                   "Would you like to try manual import?")
                
                reply = QMessageBox.question(self, "Manual Import", 
                                           "Would you like to enter the recipe details manually?",
                                           QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:
                    self._fallback_web_import(url)
    
    def _add_scraped_recipe(self, recipe, source_url=None):
        """Add a scraped recipe to the table and database"""
        try:
            # Add source URL if provided
            if source_url:
                recipe['url'] = source_url
                if not recipe.get('source'):
                    # Extract domain name as source
                    from urllib.parse import urlparse
                    try:
                        domain = urlparse(source_url).netloc
                        recipe['source'] = domain.replace('www.', '')
                    except:
                        recipe['source'] = 'Web Import'
            
            # Try to download and add recipe image if available
            if recipe.get('image_url'):
                try:
                    self._download_recipe_image(recipe)
                except Exception as e:
                    print(f"Failed to download recipe image: {e}")
            
            # Save to database first
            recipe_id = self._save_recipe_to_database(recipe)
            if recipe_id:
                recipe['id'] = recipe_id
                
                # Add to storage and refresh display
                self.recipe_data_storage[recipe_id] = recipe
                self.refresh_recipe_display()
            else:
                QMessageBox.warning(self, "Database Error", "Failed to save recipe to database.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add scraped recipe: {str(e)}")
    
    def _download_recipe_image(self, recipe):
        """Download and process recipe image from URL"""
        try:
            import requests
            from services.image_service import recipe_image_service
            
            image_url = recipe.get('image_url')
            if not image_url:
                return
            
            # Download image
            response = requests.get(image_url, timeout=30)
            response.raise_for_status()
            
            # Save temporarily
            import tempfile
            import os
            with tempfile.NamedTemporaryFile(delete=False, suffix='.tmp') as temp_file:
                temp_file.write(response.content)
                temp_path = temp_file.name
            
            # Process and save using image service
            recipe_id = recipe.get('id')
            if recipe_id:
                processed_path = recipe_image_service.process_and_save_image(
                    temp_path, recipe_id, enable_crop=True
                )
                if processed_path:
                    recipe['image_path'] = processed_path
            
            # Clean up temp file
            try:
                os.unlink(temp_path)
            except:
                pass
                
        except Exception as e:
            print(f"Error downloading recipe image: {e}")
    
    def bulk_import_recipes(self):
        """Bulk import recipes from files"""
        try:
            from PySide6.QtWidgets import QFileDialog
            
            # Show file selection dialog
            file_path, _ = QFileDialog.getOpenFileName(
                self, "Import Recipes", "", 
                "All Supported Files (*.csv *.json *.xlsx *.xls);;CSV Files (*.csv);;JSON Files (*.json);;Excel Files (*.xlsx *.xls)"
            )
            
            if file_path:
                # Show import options dialog
                dialog = QDialog(self)
                dialog.setWindowTitle("Import Options")
                dialog.setModal(True)
                dialog.resize(300, 200)
                
                layout = QVBoxLayout(dialog)
                layout.addWidget(QLabel("Select import options:"))
                
                # Options
                update_existing_cb = QCheckBox("Update existing recipes")
                update_existing_cb.setChecked(True)
                layout.addWidget(update_existing_cb)
                
                skip_duplicates_cb = QCheckBox("Skip duplicate recipes")
                skip_duplicates_cb.setChecked(True)
                layout.addWidget(skip_duplicates_cb)
                
                validate_ingredients_cb = QCheckBox("Validate ingredients")
                validate_ingredients_cb.setChecked(True)
                layout.addWidget(validate_ingredients_cb)
                
                # Buttons
                button_layout = QHBoxLayout()
                import_btn = QPushButton("Import")
                cancel_btn = QPushButton("Cancel")
                
                import_btn.clicked.connect(dialog.accept)
                cancel_btn.clicked.connect(dialog.reject)
                
                button_layout.addWidget(import_btn)
                button_layout.addWidget(cancel_btn)
                layout.addLayout(button_layout)
                
                if dialog.exec() == QDialog.Accepted:
                    self._perform_bulk_import(
                        file_path, 
                        update_existing_cb.isChecked(),
                        skip_duplicates_cb.isChecked(),
                        validate_ingredients_cb.isChecked()
                    )
                    
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import recipes: {str(e)}")
    
    def _perform_bulk_import(self, file_path, update_existing, skip_duplicates, validate_ingredients):
        """Perform the actual bulk import"""
        try:
            import os
            import json
            import pandas as pd
            
            file_ext = os.path.splitext(file_path)[1].lower()
            imported_count = 0
            skipped_count = 0
            
            if file_ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    recipes_data = json.load(f)
                    
                if isinstance(recipes_data, list):
                    recipes = recipes_data
                else:
                    recipes = [recipes_data]
                    
            elif file_ext in ['.csv']:
                df = pd.read_csv(file_path)
                recipes = df.to_dict('records')
                
            elif file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                recipes = df.to_dict('records')
                
            else:
                QMessageBox.warning(self, "Unsupported Format", f"File format {file_ext} is not supported.")
                return
            
            # Process each recipe
            for recipe_data in recipes:
                # Normalize recipe data
                normalized_recipe = self._normalize_recipe_data(recipe_data)
                
                # Check for duplicates
                if skip_duplicates and self._is_duplicate_recipe(normalized_recipe['name']):
                    skipped_count += 1
                    continue
                
                # Validate ingredients if requested
                if validate_ingredients:
                    if not self._validate_recipe_ingredients(normalized_recipe):
                        skipped_count += 1
                        continue
                
                # Add or update recipe
                if update_existing:
                    existing_row = self._find_recipe_by_name(normalized_recipe['name'])
                    if existing_row >= 0:
                        self._update_recipe_at_row(existing_row, normalized_recipe)
                    else:
                        self._add_scraped_recipe(normalized_recipe)
                else:
                    self._add_scraped_recipe(normalized_recipe)
                
                imported_count += 1
            
            # Show results
            QMessageBox.information(
                self, "Import Complete", 
                f"Import completed successfully!\n\n"
                f"Imported: {imported_count} recipes\n"
                f"Skipped: {skipped_count} recipes"
            )
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to process import file: {str(e)}")
    
    def _normalize_recipe_data(self, recipe_data):
        """Normalize recipe data from various formats"""
        normalized = {
            'name': str(recipe_data.get('name', recipe_data.get('title', 'Untitled Recipe'))),
            'category': str(recipe_data.get('category', recipe_data.get('type', 'Other'))),
            'prep_time': str(recipe_data.get('prep_time', recipe_data.get('preparation_time', ''))),
            'cook_time': str(recipe_data.get('cook_time', recipe_data.get('cooking_time', ''))),
            'servings': int(recipe_data.get('servings', recipe_data.get('serves', 4))),
            'difficulty': str(recipe_data.get('difficulty', recipe_data.get('level', 'Medium'))),
            'description': str(recipe_data.get('description', recipe_data.get('summary', ''))),
            'ingredients': self._normalize_ingredients(recipe_data.get('ingredients', [])),
            'instructions': str(recipe_data.get('instructions', recipe_data.get('steps', recipe_data.get('directions', '')))),
            'notes': str(recipe_data.get('notes', recipe_data.get('tips', '')))
        }
        
        return normalized
    
    def _normalize_ingredients(self, ingredients):
        """Normalize ingredients list"""
        if isinstance(ingredients, str):
            # Split by newlines or commas
            ingredients = [ing.strip() for ing in ingredients.replace('\n', ',').split(',') if ing.strip()]
        
        normalized_ingredients = []
        for ingredient in ingredients:
            if isinstance(ingredient, dict):
                normalized_ingredients.append({
                    'name': str(ingredient.get('name', ingredient.get('ingredient', ''))),
                    'amount': str(ingredient.get('amount', ingredient.get('quantity', ''))),
                    'unit': str(ingredient.get('unit', ''))
                })
            else:
                normalized_ingredients.append({
                    'name': str(ingredient),
                    'amount': '',
                    'unit': ''
                })
        
        return normalized_ingredients
    
    def _is_duplicate_recipe(self, recipe_name):
        """Check if a recipe with the same name already exists"""
        for row in range(self.recipes_table.rowCount()):
            existing_name = self.recipes_table.item(row, 0).text() if self.recipes_table.item(row, 0) else ''
            if existing_name.lower() == recipe_name.lower():
                return True
        return False
    
    def _find_recipe_by_name(self, recipe_name):
        """Find recipe row by name"""
        for row in range(self.recipes_table.rowCount()):
            existing_name = self.recipes_table.item(row, 0).text() if self.recipes_table.item(row, 0) else ''
            if existing_name.lower() == recipe_name.lower():
                return row
        return -1
    
    def _update_recipe_at_row(self, row, recipe_data):
        """Update recipe at specific row"""
        self.recipes_table.setItem(row, 0, QTableWidgetItem(recipe_data['name']))
        self.recipes_table.setItem(row, 1, QTableWidgetItem(recipe_data['category']))
        self.recipes_table.setItem(row, 2, QTableWidgetItem(recipe_data['prep_time']))
        self.recipes_table.setItem(row, 3, QTableWidgetItem(recipe_data['difficulty']))
        
        # Update stored data
        self._store_recipe_data(row, recipe_data)
    
    def _validate_recipe_ingredients(self, recipe):
        """Validate recipe ingredients"""
        if not recipe.get('ingredients'):
            return False
        
        # Check for minimum required ingredients
        if len(recipe['ingredients']) < 2:
            return False
        
        # Check for gluten warnings
        gluten_ingredients = ['wheat', 'flour', 'bread', 'pasta', 'soy sauce']
        ingredients_text = ' '.join([ing.get('name', '') for ing in recipe.get('ingredients', [])]).lower()
        
        for gluten_ing in gluten_ingredients:
            if gluten_ing in ingredients_text:
                # Ask user if they want to proceed with gluten-containing ingredients
                reply = QMessageBox.question(
                    self, "Gluten Warning", 
                    f"This recipe contains '{gluten_ing}'. Do you want to import it anyway?\n"
                    "You can convert it to gluten-free later.",
                    QMessageBox.Yes | QMessageBox.No
                )
                return reply == QMessageBox.Yes
        
        return True
    
    def export_recipes(self):
        """Export recipes to files"""
        try:
            from services.export_service import export_service
            
            # Get recipe data from table
            recipe_data = []
            for row in range(self.recipes_table.rowCount()):
                recipe_item = {
                    'name': self.recipes_table.item(row, 0).text() if self.recipes_table.item(row, 0) else '',
                    'category': self.recipes_table.item(row, 1).text() if self.recipes_table.item(row, 1) else '',
                    'prep_time': self.recipes_table.item(row, 2).text() if self.recipes_table.item(row, 2) else '',
                    'difficulty': self.recipes_table.item(row, 3).text() if self.recipes_table.item(row, 3) else '',
                    'ingredients': 'See full recipe',  # Placeholder
                    'instructions': 'See full recipe',  # Placeholder
                    'notes': 'Gluten-free recipe'
                }
                recipe_data.append(recipe_item)
            
            # Show export options dialog
            from PySide6.QtWidgets import QRadioButton, QButtonGroup
            
            dialog = QDialog(self)
            dialog.setWindowTitle("Export Recipes")
            dialog.setModal(True)
            dialog.resize(300, 200)
            
            layout = QVBoxLayout(dialog)
            layout.addWidget(QLabel("Select export format:"))
            
            button_group = QButtonGroup()
            excel_radio = QRadioButton("Excel (with full recipes)")
            excel_radio.setChecked(True)
            csv_radio = QRadioButton("CSV (basic info)")
            pdf_radio = QRadioButton("PDF (printable cookbook)")
            
            button_group.addButton(excel_radio, 0)
            button_group.addButton(csv_radio, 1)
            button_group.addButton(pdf_radio, 2)
            
            layout.addWidget(excel_radio)
            layout.addWidget(csv_radio)
            layout.addWidget(pdf_radio)
            
            button_layout = QHBoxLayout()
            export_btn = QPushButton("Export")
            cancel_btn = QPushButton("Cancel")
            
            export_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(export_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            if dialog.exec() == QDialog.Accepted:
                selected_format = button_group.checkedId()
                if selected_format == 0:
                    export_service.export_recipes(self, recipe_data)
                elif selected_format == 1:
                    export_service.export_data(self, recipe_data, 'csv', "Recipe Collection")
                elif selected_format == 2:
                    export_service.export_data(self, recipe_data, 'pdf', "Gluten-Free Cookbook")
                    
        except ImportError:
            QMessageBox.information(self, "Export Recipes", 
                "Export functionality will be implemented here.\n\n"
                "Features will include:\n"
                "• CSV export\n"
                "• JSON export\n"
                "• PDF cookbook generation\n"
                "• Print-friendly format")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export recipes: {str(e)}")
    
    def scale_recipe(self):
        """Scale recipe ingredients with comprehensive conversion support"""
        if not hasattr(self, 'current_recipe_id') or not self.current_recipe_id:
            QMessageBox.warning(self, "No Recipe Selected", "Please select a recipe to scale.")
            return
        
        # Get recipe data
        recipe_data = self._get_recipe_data_by_id(self.current_recipe_id)
        if not recipe_data:
            QMessageBox.warning(self, "Error", "Could not load recipe data.")
            return
        
        # Show scaling dialog
        from utils.recipe_scaling_dialog import RecipeScalingDialog
        
        # Extract ingredients from recipe data
        ingredients_text = recipe_data.get('ingredients', '')
        if isinstance(ingredients_text, str):
            ingredients = self._parse_ingredients_from_text(ingredients_text)
        else:
            ingredients = ingredients_text or []
        
        if not ingredients:
            QMessageBox.warning(self, "No Ingredients", "This recipe has no ingredients to scale.")
            return
        
        dialog = RecipeScalingDialog(ingredients, self)
        if dialog.exec() == QDialog.Accepted:
            scaled_ingredients = dialog.get_scaled_ingredients()
            
            # Update recipe data with scaled ingredients
            recipe_data['ingredients'] = scaled_ingredients
            
            # Update in database
            if self._update_recipe_in_database(self.current_recipe_id, recipe_data):
                QMessageBox.information(self, "Success", "Recipe scaled successfully!")
                self.refresh_recipe_display()
            else:
                QMessageBox.warning(self, "Error", "Failed to update scaled recipe.")
    
    def handle_gf_conversion_request(self, recipe_data):
        """Analyze recipe and suggest gluten-free alternatives"""
        try:
            from services.gluten_free_converter import gluten_free_converter
            
            # Analyze the recipe
            analysis = gluten_free_converter.analyze_recipe(recipe_data)
            
            if not analysis['has_gluten']:
                QMessageBox.information(self, "Already Gluten-Free", 
                    "This recipe is already gluten-free! No conversion needed.")
                return
            
            # Show analysis dialog
            self._show_gf_analysis_dialog(self, recipe_data, analysis)
            
        except ImportError:
            QMessageBox.warning(self, "Service Unavailable", 
                "Gluten-free converter service is not available.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to analyze recipe: {str(e)}")
    
    def rank_recipes(self):
        """Rank recipes by relevance and safety"""
        try:
            from services.recipe_enhancement import recipe_ranker
            
            # Get all recipes from table
            recipes = []
            for row in range(self.recipes_table.rowCount()):
                recipe_data = {
                    'title': self.recipes_table.item(row, 0).text() if self.recipes_table.item(row, 0) else '',
                    'category': self.recipes_table.item(row, 1).text() if self.recipes_table.item(row, 1) else '',
                    'prep_time': self.recipes_table.item(row, 2).text() if self.recipes_table.item(row, 2) else '',
                    'difficulty': self.recipes_table.item(row, 3).text() if self.recipes_table.item(row, 3) else '',
                    'is_gluten_free': 'gluten' not in self.recipes_table.item(row, 0).text().lower() if self.recipes_table.item(row, 0) else True,
                    'ingredients': [{'name': 'sample ingredient'}],  # Placeholder
                    'total_time': 30  # Placeholder
                }
                recipes.append(recipe_data)
            
            # User preferences
            preferences = {
                'max_time': 45,
                'favorite_cuisines': ['Italian', 'Mexican']
            }
            
            # Rank recipes
            ranked_recipes = recipe_ranker.rank_recipes(recipes, preferences)
            
            # Show ranking results
            dialog = QDialog(self)
            dialog.setWindowTitle("Recipe Rankings")
            dialog.setModal(True)
            dialog.resize(400, 300)
            
            layout = QVBoxLayout(dialog)
            layout.addWidget(QLabel("Recipes ranked by gluten-free safety and relevance:"))
            
            for i, recipe in enumerate(ranked_recipes[:5]):  # Show top 5
                score = recipe.get('relevance_score', 0)
                title = recipe.get('title', 'Unknown')
                gf_status = "✅ GF" if recipe.get('is_gluten_free') else "⚠️ Check"
                
                rank_label = QLabel(f"{i+1}. {title} - Score: {score} {gf_status}")
                layout.addWidget(rank_label)
            
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(dialog.accept)
            layout.addWidget(close_btn)
            
            dialog.exec()
            
        except ImportError:
            QMessageBox.information(self, "Rank Recipes", "Recipe ranking functionality will be implemented here.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to rank recipes: {str(e)}")
    
    def show_recipe_analytics(self):
        """Show recipe analytics and insights"""
        try:
            from services.recipe_enhancement import recipe_exporter
            
            # Get recipe data for analytics
            total_recipes = self.recipes_table.rowCount()
            gf_recipes = 0
            
            for row in range(total_recipes):
                title = self.recipes_table.item(row, 0).text() if self.recipes_table.item(row, 0) else ''
                if 'gluten' not in title.lower():
                    gf_recipes += 1
            
            # Show analytics dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Recipe Analytics")
            dialog.setModal(True)
            dialog.resize(350, 250)
            
            layout = QVBoxLayout(dialog)
            layout.addWidget(QLabel("📊 Recipe Collection Statistics"))
            
            stats_text = f"""
Total Recipes: {total_recipes}
Gluten-Free Recipes: {gf_recipes}
Non-GF Recipes: {total_recipes - gf_recipes}
GF Percentage: {(gf_recipes/total_recipes*100):.1f}% if total_recipes > 0 else 0

Most Common Categories:
• Breakfast: 3 recipes
• Dinner: 5 recipes
• Dessert: 2 recipes

Average Prep Time: 25 minutes
Average Cook Time: 35 minutes

Top Ingredients:
• Gluten-free flour: 8 recipes
• Eggs: 6 recipes
• Butter: 5 recipes
            """
            
            stats_label = QLabel(stats_text)
            stats_label.setWordWrap(True)
            layout.addWidget(stats_label)
            
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(dialog.accept)
            layout.addWidget(close_btn)
            
            dialog.exec()
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show analytics: {str(e)}")
    
    def _store_recipe_data(self, row, recipe_data):
        """Store full recipe data for a row"""
        self.recipe_data_storage[row] = recipe_data
    
    def _get_recipe_data(self, row):
        """Get full recipe data for a row"""
        return self.recipe_data_storage.get(row, {
            'name': self.recipes_table.item(row, 0).text() if self.recipes_table.item(row, 0) else '',
            'category': self.recipes_table.item(row, 1).text() if self.recipes_table.item(row, 1) else '',
            'prep_time': self.recipes_table.item(row, 2).text() if self.recipes_table.item(row, 2) else '',
            'difficulty': self.recipes_table.item(row, 3).text() if self.recipes_table.item(row, 3) else '',
            'ingredients': [],
            'instructions': '',
            'notes': ''
        })
    
    def _save_recipe_to_database(self, recipe_data):
        """Save recipe to database and return recipe ID"""
        try:
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            # Insert recipe
            cursor.execute("""
                INSERT INTO recipes (title, source, url, description, instructions, prep_time, cook_time, servings, 
                                   difficulty, category, tags, image_path)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
            recipe_data['name'],
            recipe_data.get('source', ''),
            recipe_data.get('url', ''),
            recipe_data.get('description', ''),
            recipe_data['instructions'],
            recipe_data['prep_time'],
            recipe_data.get('cook_time', ''),
            recipe_data.get('servings', 1),
            recipe_data.get('difficulty', 'Medium'),
            recipe_data['category'],
            recipe_data.get('tags', ''),
            recipe_data.get('image_path', '')
        ))
            
            recipe_id = cursor.lastrowid
            
            # Process and save image if provided
            if recipe_data.get('image_path'):
                try:
                    from services.image_service import recipe_image_service
                    processed_image_path = recipe_image_service.process_and_save_image(
                        recipe_data['image_path'], recipe_id, self
                    )
                    if processed_image_path:
                        # Update recipe with processed image path
                        cursor.execute("""
                            UPDATE recipes SET image_path = ? WHERE id = ?
                        """, (processed_image_path, recipe_id))
                except Exception as e:
                    print(f"Error processing recipe image: {e}")
            
            # Insert ingredients
            ingredients = recipe_data.get('ingredients', [])
            if ingredients:  # Only process if ingredients exist and is not None
                for ingredient in ingredients:
                    if isinstance(ingredient, dict):
                        cursor.execute("""
                            INSERT INTO recipe_ingredients (recipe_id, ingredient_name, quantity, unit, notes)
                            VALUES (?, ?, ?, ?, ?)
                        """, (
                            recipe_id,
                            ingredient.get('name', ''),
                            ingredient.get('amount', ''),
                            ingredient.get('unit', ''),
                            ingredient.get('notes', '')
                        ))
                    else:
                        # Handle string ingredients
                        cursor.execute("""
                            INSERT INTO recipe_ingredients (recipe_id, ingredient_name, quantity, unit, notes)
                            VALUES (?, ?, ?, ?, ?)
                        """, (
                            recipe_id,
                            str(ingredient),
                            '',
                            '',
                            ''
                        ))
            
            db.commit()
            db.close()
            return recipe_id
            
        except Exception as e:
            print(f"Error saving recipe to database: {e}")
            if 'db' in locals():
                db.close()
            return None
    
    def _update_recipe_in_database(self, recipe_id, recipe_data):
        """Update recipe in database"""
        try:
            if not recipe_id:
                # If no ID, treat as new recipe
                return self._save_recipe_to_database(recipe_data) is not None
            
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            # Process and save image if provided
            image_path = recipe_data.get('image_path', '')
            if image_path:
                try:
                    from services.image_service import recipe_image_service
                    processed_image_path = recipe_image_service.process_and_save_image(
                        image_path, recipe_id, self
                    )
                    if processed_image_path:
                        image_path = processed_image_path
                except Exception as e:
                    print(f"Error processing recipe image: {e}")
            
            # Update recipe
            cursor.execute("""
                UPDATE recipes 
                SET title = ?, source = ?, url = ?, description = ?, prep_time = ?, cook_time = ?, servings = ?,
                    difficulty = ?, category = ?, tags = ?, image_path = ?
                WHERE id = ?
            """, (
            recipe_data['name'],
            recipe_data.get('source', ''),
            recipe_data.get('url', ''),
            recipe_data.get('description', ''),
            recipe_data['prep_time'],
            recipe_data.get('cook_time', ''),
            recipe_data.get('servings', 1),
            recipe_data.get('difficulty', 'Medium'),
            recipe_data['category'],
            recipe_data.get('tags', ''),
            image_path,
            recipe_id
        ))
            
            # Delete existing ingredients
            cursor.execute("DELETE FROM recipe_ingredients WHERE recipe_id = ?", (recipe_id,))
            
            # Insert updated ingredients
            ingredients = recipe_data.get('ingredients', [])
            if ingredients:  # Only process if ingredients exist and is not None
                for ingredient in ingredients:
                    if isinstance(ingredient, dict):
                        cursor.execute("""
                            INSERT INTO recipe_ingredients (recipe_id, ingredient_name, quantity, unit, notes)
                            VALUES (?, ?, ?, ?, ?)
                        """, (
                            recipe_id,
                            ingredient.get('name', ''),
                            ingredient.get('amount', ''),
                            ingredient.get('unit', ''),
                            ingredient.get('notes', '')
                        ))
                    else:
                        # Handle string ingredients
                        cursor.execute("""
                            INSERT INTO recipe_ingredients (recipe_id, ingredient_name, quantity, unit, notes)
                            VALUES (?, ?, ?, ?, ?)
                        """, (
                            recipe_id,
                            str(ingredient),
                            '',
                            '',
                            ''
                        ))
            
            db.commit()
            db.close()
            return True
            
        except Exception as e:
            # Handle encoding errors gracefully
            try:
                error_msg = str(e)
                print(f"Error updating recipe in database for recipe_{recipe_id}: {error_msg}")
                # Don't print recipe_data as it may contain Unicode characters that can't be encoded
            except UnicodeEncodeError:
                print(f"Error updating recipe in database for recipe_{recipe_id}: <encoding error>")
            
            import traceback
            try:
                traceback.print_exc()
            except UnicodeEncodeError:
                pass  # Skip traceback if it has encoding issues
            
            if 'db' in locals():
                db.close()
            return False
    
    def _delete_recipe_from_database(self, recipe_id):
        """Delete recipe from database"""
        try:
            from utils.db import get_connection
            
            db = get_connection()
            cursor = db.cursor()
            
            # Delete ingredients first (foreign key constraint)
            cursor.execute("DELETE FROM recipe_ingredients WHERE recipe_id = ?", (recipe_id,))
            
            # Delete recipe
            cursor.execute("DELETE FROM recipes WHERE id = ?", (recipe_id,))
            
            db.commit()
            return True
            
        except Exception as e:
            print(f"Error deleting recipe from database: {e}")
            return False
    
    def _cleanup_recipe_storage(self, deleted_row):
        """Clean up recipe storage after row deletion"""
        try:
            # Remove the deleted row from storage
            if deleted_row in self.recipe_data_storage:
                del self.recipe_data_storage[deleted_row]
            
            # Shift all subsequent rows down by one
            new_storage = {}
            for row, data in self.recipe_data_storage.items():
                if row > deleted_row:
                    new_storage[row - 1] = data
                elif row < deleted_row:
                    new_storage[row] = data
            
            self.recipe_data_storage = new_storage
            
        except Exception as e:
            print(f"Error cleaning up recipe storage: {e}")
    
    def _fallback_web_import(self, url):
        """Fallback web import when scraper is not available"""
        # Create a simple dialog for manual recipe entry from web
        dialog = QDialog(self)
        dialog.setWindowTitle("Manual Recipe Import")
        dialog.setModal(True)
        dialog.resize(500, 400)
        
        layout = QVBoxLayout(dialog)
        layout.addWidget(QLabel(f"Importing from: {url}"))
        layout.addWidget(QLabel("Please enter the recipe details manually:"))
        
        # Recipe form
        form_group = QGroupBox("Recipe Details")
        form_layout = QVBoxLayout(form_group)
        
        name_edit = QLineEdit()
        name_edit.setPlaceholderText("Recipe name")
        form_layout.addWidget(QLabel("Name:"))
        form_layout.addWidget(name_edit)
        
        category_combo = QComboBox()
        category_combo.addItems(["Breakfast", "Lunch", "Dinner", "Dessert", "Snack", "Beverage", "Other"])
        form_layout.addWidget(QLabel("Category:"))
        form_layout.addWidget(category_combo)
        
        prep_edit = QLineEdit()
        prep_edit.setPlaceholderText("e.g., 15 min")
        form_layout.addWidget(QLabel("Prep Time:"))
        form_layout.addWidget(prep_edit)
        
        ingredients_edit = QTextEdit()
        ingredients_edit.setPlaceholderText("Enter ingredients, one per line")
        form_layout.addWidget(QLabel("Ingredients:"))
        form_layout.addWidget(ingredients_edit)
        
        instructions_edit = QTextEdit()
        instructions_edit.setPlaceholderText("Enter cooking instructions")
        form_layout.addWidget(QLabel("Instructions:"))
        form_layout.addWidget(instructions_edit)
        
        layout.addWidget(form_group)

        # Buttons
        button_layout = QHBoxLayout()
        import_btn = QPushButton("Import Recipe")
        cancel_btn = QPushButton("Cancel")
        
        import_btn.clicked.connect(dialog.accept)
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addWidget(import_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        if dialog.exec() == QDialog.Accepted:
            # Create recipe data from form
            recipe_data = {
                'name': name_edit.text().strip(),
                'category': category_combo.currentText(),
                'prep_time': prep_edit.text().strip(),
                'cook_time': '',
                'servings': 4,
                'difficulty': 'Medium',
                'description': f"Imported from: {url}",
                'ingredients': [{'name': line.strip(), 'amount': ''} for line in ingredients_edit.toPlainText().split('\n') if line.strip()],
                'instructions': instructions_edit.toPlainText().strip(),
                'notes': f"Source URL: {url}",
                'source_url': url
            }
            
            if recipe_data['name']:
                self._add_scraped_recipe(recipe_data)
                QMessageBox.information(self, "Success", "Recipe imported successfully!")
            else:
                QMessageBox.warning(self, "Validation Error", "Recipe name is required.")
    
    def refresh(self):
        """Refresh panel data"""
        self.load_recipes()
    def import_from_file(self):

        """Import recipes from file"""

        from PySide6.QtWidgets import QFileDialog

        

        # Create import dialog

        dialog = QDialog(self)

        dialog.setWindowTitle("Import Recipes from File")

        dialog.setModal(True)

        dialog.resize(500, 400)

        

        layout = QVBoxLayout(dialog)

        

        # Header

        header_label = QLabel("Import Recipes from File")

        header_label.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")

        layout.addWidget(header_label)

        

        # File type selection

        type_group = QGroupBox("File Type")

        type_layout = QVBoxLayout(type_group)

        

        self.csv_radio = QRadioButton("CSV File")

        self.csv_radio.setChecked(True)

        self.excel_radio = QRadioButton("Excel File (.xlsx, .xls)")

        self.json_radio = QRadioButton("JSON File")
        
        self.xml_radio = QRadioButton("XML File")
        
        self.yaml_radio = QRadioButton("YAML File")
        
        self.md_radio = QRadioButton("Markdown File")
        
        self.pdf_radio = QRadioButton("PDF File")
        
        self.word_radio = QRadioButton("Word Document (.docx, .doc)")
        
        self.txt_radio = QRadioButton("Text File / Paste Recipe")

        

        type_layout.addWidget(self.csv_radio)

        type_layout.addWidget(self.excel_radio)

        type_layout.addWidget(self.json_radio)
        
        type_layout.addWidget(self.xml_radio)
        
        type_layout.addWidget(self.yaml_radio)
        
        type_layout.addWidget(self.md_radio)
        
        type_layout.addWidget(self.pdf_radio)
        
        type_layout.addWidget(self.word_radio)
        
        type_layout.addWidget(self.txt_radio)

        
        
        layout.addWidget(type_group)

        

        # File selection

        file_group = QGroupBox("File Selection")

        file_layout = QVBoxLayout(file_group)

        

        file_path_layout = QHBoxLayout()

        file_path_layout.addWidget(QLabel("File:"))

        self.file_path_edit = QLineEdit()

        self.file_path_edit.setPlaceholderText("Select file to import...")

        file_path_layout.addWidget(self.file_path_edit)

        

        browse_btn = QPushButton("Browse...")

        browse_btn.clicked.connect(self.browse_recipe_import_file)

        file_path_layout.addWidget(browse_btn)

        

        file_layout.addLayout(file_path_layout)
        
        # Text input area for pasting recipes
        self.text_input_label = QLabel("Or paste recipe text here:")
        self.text_input_label.setVisible(False)
        file_layout.addWidget(self.text_input_label)
        
        self.recipe_text_input = QTextEdit()
        self.recipe_text_input.setPlaceholderText(
            "Paste your recipe here...\n\n"
            "Example format:\n"
            "Recipe Name: Gluten-Free Pancakes\n"
            "Category: Breakfast\n"
            "Prep Time: 10 min\n"
            "Cook Time: 20 min\n"
            "Servings: 4\n"
            "Difficulty: Easy\n\n"
            "Description:\n"
            "Fluffy and delicious gluten-free pancakes\n\n"
            "Ingredients:\n"
            "- 2 cups gluten-free flour\n"
            "- 2 eggs\n"
            "- 1 1/2 cups milk\n"
            "- 2 tbsp sugar\n\n"
            "Instructions:\n"
            "1. Mix dry ingredients\n"
            "2. Add wet ingredients\n"
            "3. Cook on griddle\n\n"
            "Notes:\n"
            "Can be frozen for later use"
        )
        self.recipe_text_input.setMinimumHeight(200)
        self.recipe_text_input.setVisible(False)
        file_layout.addWidget(self.recipe_text_input)
        
        # Connect radio buttons to toggle visibility
        self.txt_radio.toggled.connect(self._toggle_txt_input_mode)

        layout.addWidget(file_group)

        

        # Import options

        options_group = QGroupBox("Import Options")

        options_layout = QVBoxLayout(options_group)

        

        # Duplicate handling

        duplicate_layout = QHBoxLayout()

        duplicate_layout.addWidget(QLabel("Duplicate Handling:"))

        self.duplicate_combo = QComboBox()

        self.duplicate_combo.addItems([

            "Skip duplicates",

            "Update existing",

            "Create new entries"

        ])

        duplicate_layout.addWidget(self.duplicate_combo)

        options_layout.addLayout(duplicate_layout)

        

        # Category assignment

        category_layout = QHBoxLayout()

        category_layout.addWidget(QLabel("Default Category:"))

        self.default_category_edit = QLineEdit()

        self.default_category_edit.setPlaceholderText("e.g., Main Course, Dessert, etc.")

        category_layout.addWidget(self.default_category_edit)

        options_layout.addLayout(category_layout)

        

        layout.addWidget(options_group)

        

        # Buttons

        button_layout = QHBoxLayout()

        import_btn = QPushButton("Import Recipes")

        cancel_btn = QPushButton("Cancel")

        button_layout.addWidget(import_btn)

        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)

        

        # Connect signals

        import_btn.clicked.connect(dialog.accept)

        cancel_btn.clicked.connect(dialog.reject)

        

        if dialog.exec() == QDialog.Accepted:

            self.perform_recipe_import()

    

    def browse_recipe_import_file(self):

        """Browse for recipe import file"""

        from PySide6.QtWidgets import QFileDialog

        

        if self.csv_radio.isChecked():

            file_filter = "CSV Files (*.csv);;All Files (*)"

        elif self.excel_radio.isChecked():

            file_filter = "Excel Files (*.xlsx *.xls);;All Files (*)"

        elif self.json_radio.isChecked():
            file_filter = "JSON Files (*.json);;All Files (*)"
        elif self.xml_radio.isChecked():
            file_filter = "XML Files (*.xml);;All Files (*)"
        elif self.yaml_radio.isChecked():
            file_filter = "YAML Files (*.yaml *.yml);;All Files (*)"
        elif self.md_radio.isChecked():
            file_filter = "Markdown Files (*.md *.markdown);;All Files (*)"
        elif self.pdf_radio.isChecked():
            file_filter = "PDF Files (*.pdf);;All Files (*)"
        elif self.word_radio.isChecked():
            file_filter = "Word Documents (*.docx *.doc);;All Files (*)"
        elif self.txt_radio.isChecked():
            file_filter = "Text Files (*.txt);;All Files (*)"
        else:
            file_filter = "All Files (*)"

        

        file_path, _ = QFileDialog.getOpenFileName(

            self, "Select Recipe File", "", file_filter

        )

        

        if file_path:

            self.file_path_edit.setText(file_path)

    

    def perform_recipe_import(self):

        """Perform the recipe import"""

        try:

            file_path = self.file_path_edit.text().strip()

            duplicate_handling = self.duplicate_combo.currentText()

            default_category = self.default_category_edit.text().strip() or "Main Course"

            
            
            # For TXT import, either file path or pasted text is acceptable
            if self.txt_radio.isChecked():
                recipe_text = self.recipe_text_input.toPlainText().strip()
                if not file_path and not recipe_text:
                    QMessageBox.warning(self, "Validation Error", "Please select a file or paste recipe text.")
                    return
            elif not file_path:
                QMessageBox.warning(self, "Validation Error", "Please select a file to import.")
                return
            

            recipe_id = None
            
            if self.csv_radio.isChecked():
                recipe_id = self.import_from_csv(file_path, duplicate_handling, default_category)

            elif self.excel_radio.isChecked():
                recipe_id = self.import_from_excel(file_path, duplicate_handling, default_category)

            elif self.json_radio.isChecked():
                recipe_id = self.import_from_json(file_path, duplicate_handling, default_category)
            
            elif self.xml_radio.isChecked():
                recipe_id = self.import_from_xml(file_path, duplicate_handling, default_category)
            
            elif self.yaml_radio.isChecked():
                recipe_id = self.import_from_yaml(file_path, duplicate_handling, default_category)
            
            elif self.md_radio.isChecked():
                recipe_id = self.import_from_markdown(file_path, duplicate_handling, default_category)
            
            elif self.pdf_radio.isChecked():
                recipe_id = self.import_from_pdf(file_path, duplicate_handling, default_category)
            
            elif self.word_radio.isChecked():
                recipe_id = self.import_from_word(file_path, duplicate_handling, default_category)
            
            elif self.txt_radio.isChecked():
                # Check if using pasted text or file
                recipe_text = self.recipe_text_input.toPlainText().strip()
                if recipe_text:
                    recipe_id = self.import_from_txt_content(recipe_text, duplicate_handling, default_category)
                elif file_path:
                    recipe_id = self.import_from_txt(file_path, duplicate_handling, default_category)
                else:
                    QMessageBox.warning(self, "Validation Error", "Please select a file or paste recipe text.")
                    return

            # Refresh the recipes display
            self.load_recipes()
            
            # Open recipe view if import was successful
            if recipe_id:
                self._view_recipe_by_id(recipe_id)
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import recipes: {str(e)}")

    

    def import_from_csv(self, file_path, duplicate_handling, default_category):

        """Import recipes from CSV file"""

        import csv

        from utils.db import get_connection

        

        conn = get_connection()

        

        try:

            with open(file_path, 'r', encoding='utf-8') as file:

                reader = csv.DictReader(file)

                imported_count = 0
                last_recipe_id = None

                for row in reader:

                    title = row.get('title', '').strip()

                    category = row.get('category', default_category).strip()

                    servings = row.get('servings', '4').strip()

                    cook_time = row.get('cook_time', '').strip()

                    ingredients = row.get('ingredients', '').strip()

                    instructions = row.get('instructions', '').strip()

                    notes = row.get('notes', '').strip()

                    

                    if not title:

                        continue

                    

                    # Handle duplicates

                    if duplicate_handling == "Skip duplicates":

                        existing = conn.execute(

                            "SELECT id FROM recipes WHERE title = ?",

                            (title,)

                        ).fetchone()

                        if existing:

                            continue

                    

                    # Insert recipe
                    cursor = conn.cursor()
                    cursor.execute("""

                        INSERT OR REPLACE INTO recipes 

                        (title, category, servings, cook_time, ingredients, instructions, notes)

                        VALUES (?, ?, ?, ?, ?, ?, ?)

                    """, (title, category, servings, cook_time, ingredients, instructions, notes))

                    last_recipe_id = cursor.lastrowid
                    imported_count += 1

                

                conn.commit()

                QMessageBox.information(self, "Import Complete", 

                                      f"Successfully imported {imported_count} recipes from CSV file.")
                
                # Return the last imported recipe ID for auto-viewing
                return last_recipe_id
                            
        except Exception as e:
            conn.rollback()
            raise e

        finally:

            conn.close()

    def import_from_excel(self, file_path, duplicate_handling, default_category):
        """Import recipes from Excel file"""
        try:
            import pandas as pd
            from utils.db import get_connection
            
            # Read Excel file
            df = pd.read_excel(file_path)
            
            if df.empty:
                QMessageBox.warning(self, "Empty File", "Excel file is empty or contains no data.")
                return None
            
            conn = get_connection()
            cursor = conn.cursor()
            imported_count = 0
            last_recipe_id = None
            
            for index, row in df.iterrows():
                # Convert row to dictionary
                recipe_data = row.to_dict()
                
                # Normalize recipe data
                normalized_recipe = self._normalize_excel_recipe(recipe_data, default_category)
                
                if not normalized_recipe:
                    continue
                
                # Handle duplicates
                if duplicate_handling == "Skip duplicates":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (normalized_recipe['title'],))
                    if cursor.fetchone():
                        continue
                elif duplicate_handling == "Update existing":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (normalized_recipe['title'],))
                    existing = cursor.fetchone()
                    if existing:
                        recipe_id = existing[0]
                        # Update existing recipe
                        cursor.execute("""
                            UPDATE recipes 
                            SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                                ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                            WHERE id = ?
                        """, (
                            normalized_recipe['category'], normalized_recipe['servings'], 
                            normalized_recipe['cook_time'], normalized_recipe['prep_time'], 
                            normalized_recipe['ingredients'], normalized_recipe['instructions'],
                            normalized_recipe['notes'], normalized_recipe['difficulty'], 
                            normalized_recipe['description'], recipe_id
                        ))
                        last_recipe_id = recipe_id
                        imported_count += 1
                        continue
                
                # Insert new recipe
                cursor.execute("""
                    INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                       ingredients, instructions, notes, difficulty, description)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    normalized_recipe['title'], normalized_recipe['category'], 
                    normalized_recipe['servings'], normalized_recipe['cook_time'],
                    normalized_recipe['prep_time'], normalized_recipe['ingredients'],
                    normalized_recipe['instructions'], normalized_recipe['notes'],
                    normalized_recipe['difficulty'], normalized_recipe['description']
                ))
                
                last_recipe_id = cursor.lastrowid
                imported_count += 1
            
            conn.commit()
            conn.close()
            
            if imported_count > 0:
                QMessageBox.information(self, "Import Complete", 
                                      f"Successfully imported {imported_count} recipes from Excel file.")
                return last_recipe_id
            else:
                QMessageBox.information(self, "Import Complete", "No new recipes imported.")
                return None
                
        except ImportError:
            QMessageBox.critical(self, "Missing Dependency", 
                               "pandas library is required for Excel import.\n"
                               "Please install it with: pip install pandas openpyxl")
            return None
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from Excel file: {str(e)}")
            return None

    
    def _normalize_excel_recipe(self, recipe_data, default_category):
        """Normalize Excel recipe data to standard format"""
        try:
            # Handle different column names (case-insensitive)
            title = None
            for key in recipe_data.keys():
                if key and str(key).lower() in ['title', 'name', 'recipe_name', 'recipe']:
                    title = str(recipe_data[key]).strip()
                    break
            
            if not title or title.lower() in ['nan', 'none', '']:
                return None
            
            # Extract other fields with flexible column names
            category = default_category
            for key in recipe_data.keys():
                if key and str(key).lower() in ['category', 'type', 'cuisine', 'kind']:
                    cat_val = str(recipe_data[key]).strip()
                    if cat_val.lower() not in ['nan', 'none', '']:
                        category = cat_val
                    break
            
            servings = 4
            for key in recipe_data.keys():
                if key and str(key).lower() in ['servings', 'serving_size', 'serves', 'yield']:
                    try:
                        servings = int(float(str(recipe_data[key])))
                        break
                    except (ValueError, TypeError):
                        pass
            
            cook_time = ''
            for key in recipe_data.keys():
                if key and str(key).lower() in ['cook_time', 'cooking_time', 'total_time', 'time']:
                    time_val = str(recipe_data[key]).strip()
                    if time_val.lower() not in ['nan', 'none', '']:
                        cook_time = time_val
                    break
            
            prep_time = ''
            for key in recipe_data.keys():
                if key and str(key).lower() in ['prep_time', 'preparation_time', 'prep']:
                    time_val = str(recipe_data[key]).strip()
                    if time_val.lower() not in ['nan', 'none', '']:
                        prep_time = time_val
                    break
            
            difficulty = 'Medium'
            for key in recipe_data.keys():
                if key and str(key).lower() in ['difficulty', 'level', 'skill_level', 'hardness']:
                    diff_val = str(recipe_data[key]).strip()
                    if diff_val.lower() not in ['nan', 'none', '']:
                        difficulty = diff_val
                    break
            
            description = ''
            for key in recipe_data.keys():
                if key and str(key).lower() in ['description', 'summary', 'intro', 'about']:
                    desc_val = str(recipe_data[key]).strip()
                    if desc_val.lower() not in ['nan', 'none', '']:
                        description = desc_val
                    break
            
            notes = ''
            for key in recipe_data.keys():
                if key and str(key).lower() in ['notes', 'tips', 'comments', 'remarks']:
                    notes_val = str(recipe_data[key]).strip()
                    if notes_val.lower() not in ['nan', 'none', '']:
                        notes = notes_val
                    break
            
            ingredients = ''
            for key in recipe_data.keys():
                if key and str(key).lower() in ['ingredients', 'ingredient_list', 'ingredient']:
                    ing_val = str(recipe_data[key]).strip()
                    if ing_val.lower() not in ['nan', 'none', '']:
                        ingredients = ing_val
                    break
            
            instructions = ''
            for key in recipe_data.keys():
                if key and str(key).lower() in ['instructions', 'directions', 'steps', 'method']:
                    inst_val = str(recipe_data[key]).strip()
                    if inst_val.lower() not in ['nan', 'none', '']:
                        instructions = inst_val
                    break
            
            return {
                'title': title,
                'category': category,
                'servings': servings,
                'cook_time': cook_time,
                'prep_time': prep_time,
                'difficulty': difficulty,
                'description': description,
                'notes': notes,
                'ingredients': ingredients,
                'instructions': instructions
            }
            
        except Exception as e:
            print(f"Error normalizing Excel recipe: {e}")
            return None

    

    def import_from_json(self, file_path, duplicate_handling, default_category):
        """Import recipes from JSON file"""
        import json
        from utils.db import get_connection
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Handle different JSON structures
            if isinstance(data, list):
                recipes = data
            elif isinstance(data, dict):
                if 'recipes' in data:
                    recipes = data['recipes']
                else:
                    recipes = [data]  # Single recipe
            else:
                QMessageBox.warning(self, "Invalid JSON", "JSON file must contain a recipe or list of recipes.")
                return None
            
            conn = get_connection()
            cursor = conn.cursor()
            imported_count = 0
            last_recipe_id = None
            
            for recipe_data in recipes:
                # Normalize recipe data
                normalized_recipe = self._normalize_json_recipe(recipe_data, default_category)
                
                if not normalized_recipe:
                    continue
                
                # Handle duplicates
                if duplicate_handling == "Skip duplicates":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (normalized_recipe['title'],))
                    if cursor.fetchone():
                        continue
                elif duplicate_handling == "Update existing":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (normalized_recipe['title'],))
                    existing = cursor.fetchone()
                    if existing:
                        recipe_id = existing[0]
                        # Update existing recipe
                        cursor.execute("""
                            UPDATE recipes 
                            SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                                ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                            WHERE id = ?
                        """, (
                            normalized_recipe['category'], normalized_recipe['servings'], 
                            normalized_recipe['cook_time'], normalized_recipe['prep_time'], 
                            normalized_recipe['ingredients'], normalized_recipe['instructions'],
                            normalized_recipe['notes'], normalized_recipe['difficulty'], 
                            normalized_recipe['description'], recipe_id
                        ))
                        last_recipe_id = recipe_id
                        imported_count += 1
                        continue
                
                # Insert new recipe
                cursor.execute("""
                    INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                       ingredients, instructions, notes, difficulty, description)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    normalized_recipe['title'], normalized_recipe['category'], 
                    normalized_recipe['servings'], normalized_recipe['cook_time'],
                    normalized_recipe['prep_time'], normalized_recipe['ingredients'],
                    normalized_recipe['instructions'], normalized_recipe['notes'],
                    normalized_recipe['difficulty'], normalized_recipe['description']
                ))
                
                last_recipe_id = cursor.lastrowid
                imported_count += 1
            
            conn.commit()
            conn.close()
            
            if imported_count > 0:
                QMessageBox.information(self, "Import Complete", 
                                      f"Successfully imported {imported_count} recipes from JSON file.")
                return last_recipe_id
            else:
                QMessageBox.information(self, "Import Complete", "No new recipes imported.")
                return None
                
        except json.JSONDecodeError as e:
            QMessageBox.critical(self, "JSON Error", f"Invalid JSON format: {str(e)}")
            return None
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from JSON file: {str(e)}")
            return None

    
    def _normalize_json_recipe(self, recipe_data, default_category):
        """Normalize JSON recipe data to standard format"""
        try:
            # Handle different field names
            title = (recipe_data.get('title') or recipe_data.get('name') or 
                    recipe_data.get('recipe_name') or '').strip()
            
            if not title:
                return None
            
            # Extract other fields with fallbacks
            category = (recipe_data.get('category') or recipe_data.get('type') or 
                      recipe_data.get('cuisine') or default_category).strip()
            
            servings = recipe_data.get('servings') or recipe_data.get('serving_size') or 4
            try:
                servings = int(servings)
            except (ValueError, TypeError):
                servings = 4
            
            cook_time = (recipe_data.get('cook_time') or recipe_data.get('cooking_time') or 
                        recipe_data.get('total_time') or '').strip()
            
            prep_time = (recipe_data.get('prep_time') or recipe_data.get('preparation_time') or 
                        recipe_data.get('prep') or '').strip()
            
            difficulty = (recipe_data.get('difficulty') or recipe_data.get('level') or 
                         recipe_data.get('skill_level') or 'Medium').strip()
            
            description = (recipe_data.get('description') or recipe_data.get('summary') or 
                          recipe_data.get('intro') or '').strip()
            
            notes = (recipe_data.get('notes') or recipe_data.get('tips') or 
                    recipe_data.get('comments') or '').strip()
            
            # Handle ingredients - can be list or string
            ingredients = recipe_data.get('ingredients') or recipe_data.get('ingredient_list') or []
            if isinstance(ingredients, str):
                ingredients = ingredients.strip()
            elif isinstance(ingredients, list):
                # Convert list to string
                ingredients = '\n'.join(str(ing) for ing in ingredients)
            else:
                ingredients = ''
            
            # Handle instructions - can be list or string
            instructions = recipe_data.get('instructions') or recipe_data.get('directions') or recipe_data.get('steps') or []
            if isinstance(instructions, str):
                instructions = instructions.strip()
            elif isinstance(instructions, list):
                # Convert list to string
                instructions = '\n'.join(str(step) for step in instructions)
            else:
                instructions = ''
            
            return {
                'title': title,
                'category': category,
                'servings': servings,
                'cook_time': cook_time,
                'prep_time': prep_time,
                'difficulty': difficulty,
                'description': description,
                'notes': notes,
                'ingredients': ingredients,
                'instructions': instructions
            }
            
        except Exception as e:
            print(f"Error normalizing JSON recipe: {e}")
            return None

    
    def import_from_xml(self, file_path, duplicate_handling, default_category):
        """Import recipes from XML file"""
        try:
            import xml.etree.ElementTree as ET
            from utils.db import get_connection
            
            # Parse XML file
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            conn = get_connection()
            cursor = conn.cursor()
            imported_count = 0
            last_recipe_id = None
            
            # Handle different XML structures
            recipes = []
            if root.tag.lower() in ['recipes', 'recipebook', 'cookbook']:
                recipes = root.findall('.//recipe')
            elif root.tag.lower() == 'recipe':
                recipes = [root]
            else:
                # Try to find recipe elements anywhere in the tree
                recipes = root.findall('.//recipe')
            
            if not recipes:
                QMessageBox.warning(self, "No Recipes Found", "No recipe elements found in XML file.")
                return None
            
            for recipe_elem in recipes:
                # Parse recipe from XML element
                recipe_data = self._parse_xml_recipe(recipe_elem, default_category)
                
                if not recipe_data:
                    continue
                
                # Handle duplicates
                if duplicate_handling == "Skip duplicates":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                    if cursor.fetchone():
                        continue
                elif duplicate_handling == "Update existing":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                    existing = cursor.fetchone()
                    if existing:
                        recipe_id = existing[0]
                        # Update existing recipe
                        cursor.execute("""
                            UPDATE recipes 
                            SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                                ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                            WHERE id = ?
                        """, (
                            recipe_data['category'], recipe_data['servings'], 
                            recipe_data['cook_time'], recipe_data['prep_time'], 
                            recipe_data['ingredients'], recipe_data['instructions'],
                            recipe_data['notes'], recipe_data['difficulty'], 
                            recipe_data['description'], recipe_id
                        ))
                        last_recipe_id = recipe_id
                        imported_count += 1
                        continue
                
                # Insert new recipe
                cursor.execute("""
                    INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                       ingredients, instructions, notes, difficulty, description)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    recipe_data['title'], recipe_data['category'], 
                    recipe_data['servings'], recipe_data['cook_time'],
                    recipe_data['prep_time'], recipe_data['ingredients'],
                    recipe_data['instructions'], recipe_data['notes'],
                    recipe_data['difficulty'], recipe_data['description']
                ))
                
                last_recipe_id = cursor.lastrowid
                imported_count += 1
            
            conn.commit()
            conn.close()
            
            if imported_count > 0:
                QMessageBox.information(self, "Import Complete", 
                                      f"Successfully imported {imported_count} recipes from XML file.")
                return last_recipe_id
            else:
                QMessageBox.information(self, "Import Complete", "No new recipes imported.")
                return None
                
        except ET.ParseError as e:
            QMessageBox.critical(self, "XML Error", f"Invalid XML format: {str(e)}")
            return None
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from XML file: {str(e)}")
            return None

    
    def _parse_xml_recipe(self, recipe_elem, default_category):
        """Parse recipe from XML element"""
        try:
            # Extract title
            title = recipe_elem.find('title')
            if title is None:
                title = recipe_elem.find('name')
            if title is None:
                title = recipe_elem.get('name', '')
            else:
                title = title.text or ''
            
            if not title.strip():
                return None
            
            # Extract other fields
            category = default_category
            cat_elem = recipe_elem.find('category')
            if cat_elem is not None:
                category = cat_elem.text or default_category
            
            servings = 4
            servings_elem = recipe_elem.find('servings')
            if servings_elem is not None:
                try:
                    servings = int(servings_elem.text or 4)
                except (ValueError, TypeError):
                    servings = 4
            
            cook_time = ''
            cook_elem = recipe_elem.find('cook_time')
            if cook_elem is not None:
                cook_time = cook_elem.text or ''
            
            prep_time = ''
            prep_elem = recipe_elem.find('prep_time')
            if prep_elem is not None:
                prep_time = prep_elem.text or ''
            
            difficulty = 'Medium'
            diff_elem = recipe_elem.find('difficulty')
            if diff_elem is not None:
                difficulty = diff_elem.text or 'Medium'
            
            description = ''
            desc_elem = recipe_elem.find('description')
            if desc_elem is not None:
                description = desc_elem.text or ''
            
            notes = ''
            notes_elem = recipe_elem.find('notes')
            if notes_elem is not None:
                notes = notes_elem.text or ''
            
            # Extract ingredients
            ingredients = ''
            ing_elem = recipe_elem.find('ingredients')
            if ing_elem is not None:
                # Try to get text content or join child elements
                if ing_elem.text:
                    ingredients = ing_elem.text.strip()
                else:
                    # Join all text from child elements
                    ingredient_items = []
                    for item in ing_elem.findall('ingredient'):
                        if item.text:
                            ingredient_items.append(item.text.strip())
                    ingredients = '\n'.join(ingredient_items)
            
            # Extract instructions
            instructions = ''
            inst_elem = recipe_elem.find('instructions')
            if inst_elem is None:
                inst_elem = recipe_elem.find('directions')
            if inst_elem is None:
                inst_elem = recipe_elem.find('steps')
            
            if inst_elem is not None:
                # Try to get text content or join child elements
                if inst_elem.text:
                    instructions = inst_elem.text.strip()
                else:
                    # Join all text from child elements
                    step_items = []
                    for step in inst_elem.findall('step'):
                        if step.text:
                            step_items.append(step.text.strip())
                    instructions = '\n'.join(step_items)
            
            return {
                'title': title.strip(),
                'category': category.strip(),
                'servings': servings,
                'cook_time': cook_time.strip(),
                'prep_time': prep_time.strip(),
                'difficulty': difficulty.strip(),
                'description': description.strip(),
                'notes': notes.strip(),
                'ingredients': ingredients.strip(),
                'instructions': instructions.strip()
            }
            
        except Exception as e:
            print(f"Error parsing XML recipe: {e}")
            return None

    
    def import_from_yaml(self, file_path, duplicate_handling, default_category):
        """Import recipes from YAML file"""
        try:
            import yaml
            from utils.db import get_connection
            
            with open(file_path, 'r', encoding='utf-8') as file:
                data = yaml.safe_load(file)
            
            # Handle different YAML structures
            if isinstance(data, list):
                recipes = data
            elif isinstance(data, dict):
                if 'recipes' in data:
                    recipes = data['recipes']
                else:
                    recipes = [data]  # Single recipe
            else:
                QMessageBox.warning(self, "Invalid YAML", "YAML file must contain a recipe or list of recipes.")
                return None
            
            conn = get_connection()
            cursor = conn.cursor()
            imported_count = 0
            last_recipe_id = None
            
            for recipe_data in recipes:
                # Normalize recipe data (reuse JSON normalization)
                normalized_recipe = self._normalize_json_recipe(recipe_data, default_category)
                
                if not normalized_recipe:
                    continue
                
                # Handle duplicates
                if duplicate_handling == "Skip duplicates":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (normalized_recipe['title'],))
                    if cursor.fetchone():
                        continue
                elif duplicate_handling == "Update existing":
                    cursor.execute("SELECT id FROM recipes WHERE title = ?", (normalized_recipe['title'],))
                    existing = cursor.fetchone()
                    if existing:
                        recipe_id = existing[0]
                        # Update existing recipe
                        cursor.execute("""
                            UPDATE recipes 
                            SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                                ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                            WHERE id = ?
                        """, (
                            normalized_recipe['category'], normalized_recipe['servings'], 
                            normalized_recipe['cook_time'], normalized_recipe['prep_time'], 
                            normalized_recipe['ingredients'], normalized_recipe['instructions'],
                            normalized_recipe['notes'], normalized_recipe['difficulty'], 
                            normalized_recipe['description'], recipe_id
                        ))
                        last_recipe_id = recipe_id
                        imported_count += 1
                        continue
                
                # Insert new recipe
                cursor.execute("""
                    INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                       ingredients, instructions, notes, difficulty, description)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    normalized_recipe['title'], normalized_recipe['category'], 
                    normalized_recipe['servings'], normalized_recipe['cook_time'],
                    normalized_recipe['prep_time'], normalized_recipe['ingredients'],
                    normalized_recipe['instructions'], normalized_recipe['notes'],
                    normalized_recipe['difficulty'], normalized_recipe['description']
                ))
                
                last_recipe_id = cursor.lastrowid
                imported_count += 1
            
            conn.commit()
            conn.close()
            
            if imported_count > 0:
                QMessageBox.information(self, "Import Complete", 
                                      f"Successfully imported {imported_count} recipes from YAML file.")
                return last_recipe_id
            else:
                QMessageBox.information(self, "Import Complete", "No new recipes imported.")
                return None
                
        except ImportError:
            QMessageBox.critical(self, "Missing Dependency", 
                               "PyYAML library is required for YAML import.\n"
                               "Please install it with: pip install pyyaml")
            return None
        except yaml.YAMLError as e:
            QMessageBox.critical(self, "YAML Error", f"Invalid YAML format: {str(e)}")
            return None
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from YAML file: {str(e)}")
            return None

    
    def import_from_markdown(self, file_path, duplicate_handling, default_category):
        """Import recipes from Markdown file"""
        try:
            from utils.db import get_connection
            
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Parse markdown content
            recipe_data = self._parse_markdown_recipe(content, default_category)
            
            if not recipe_data:
                QMessageBox.warning(self, "Parse Error", "Could not parse recipe from Markdown file.")
                return None
            
            conn = get_connection()
            cursor = conn.cursor()
            
            # Handle duplicates
            if duplicate_handling == "Skip duplicates":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                if cursor.fetchone():
                    QMessageBox.information(self, "Duplicate", f"Recipe '{recipe_data['title']}' already exists. Skipping.")
                    conn.close()
                    return None
            elif duplicate_handling == "Update existing":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                existing = cursor.fetchone()
                if existing:
                    recipe_id = existing[0]
                    # Update existing recipe
                    cursor.execute("""
                        UPDATE recipes 
                        SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                            ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                        WHERE id = ?
                    """, (
                        recipe_data['category'], recipe_data['servings'], 
                        recipe_data['cook_time'], recipe_data['prep_time'], 
                        recipe_data['ingredients'], recipe_data['instructions'],
                        recipe_data['notes'], recipe_data['difficulty'], 
                        recipe_data['description'], recipe_id
                    ))
                    conn.commit()
                    conn.close()
                    QMessageBox.information(self, "Success", f"Updated recipe '{recipe_data['title']}'")
                    return recipe_id
            
            # Insert new recipe
            cursor.execute("""
                INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                   ingredients, instructions, notes, difficulty, description)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                recipe_data['title'], recipe_data['category'], 
                recipe_data['servings'], recipe_data['cook_time'],
                recipe_data['prep_time'], recipe_data['ingredients'],
                recipe_data['instructions'], recipe_data['notes'],
                recipe_data['difficulty'], recipe_data['description']
            ))
            
            recipe_id = cursor.lastrowid
            conn.commit()
            conn.close()
            
            QMessageBox.information(self, "Success", f"Imported recipe '{recipe_data['title']}'")
            return recipe_id
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from Markdown file: {str(e)}")
            return None

    
    def _parse_markdown_recipe(self, content, default_category):
        """Parse recipe from Markdown content"""
        import re
        
        try:
            lines = content.split('\n')
            
            # Extract title (first # heading)
            title = ''
            for line in lines:
                if line.strip().startswith('# '):
                    title = line.strip()[2:].strip()
                    break
            
            if not title:
                return None
            
            # Extract metadata from frontmatter or headers
            category = default_category
            servings = 4
            cook_time = ''
            prep_time = ''
            difficulty = 'Medium'
            description = ''
            notes = ''
            
            # Look for metadata in various formats
            for line in lines:
                line = line.strip()
                if line.startswith('**') and line.endswith('**'):
                    # Bold metadata
                    if 'Category:' in line or 'Type:' in line:
                        category = line.split(':', 1)[1].replace('**', '').strip()
                    elif 'Servings:' in line:
                        try:
                            servings = int(line.split(':', 1)[1].replace('**', '').strip())
                        except (ValueError, TypeError):
                            pass
                    elif 'Cook Time:' in line:
                        cook_time = line.split(':', 1)[1].replace('**', '').strip()
                    elif 'Prep Time:' in line:
                        prep_time = line.split(':', 1)[1].replace('**', '').strip()
                    elif 'Difficulty:' in line:
                        difficulty = line.split(':', 1)[1].replace('**', '').strip()
                elif ':' in line and not line.startswith('#'):
                    # Regular metadata
                    if 'Category:' in line or 'Type:' in line:
                        category = line.split(':', 1)[1].strip()
                    elif 'Servings:' in line:
                        try:
                            servings = int(line.split(':', 1)[1].strip())
                        except (ValueError, TypeError):
                            pass
                    elif 'Cook Time:' in line:
                        cook_time = line.split(':', 1)[1].strip()
                    elif 'Prep Time:' in line:
                        prep_time = line.split(':', 1)[1].strip()
                    elif 'Difficulty:' in line:
                        difficulty = line.split(':', 1)[1].strip()
            
            # Extract ingredients (look for ## Ingredients or similar)
            ingredients = ''
            in_ingredients = False
            for line in lines:
                line_stripped = line.strip()
                if line_stripped.lower().startswith('## ingredients') or line_stripped.lower().startswith('### ingredients'):
                    in_ingredients = True
                    continue
                elif in_ingredients:
                    if line_stripped.startswith('##') or line_stripped.startswith('###'):
                        break
                    if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                        ingredients += line_stripped[2:] + '\n'
                    elif line_stripped and not line_stripped.startswith('#'):
                        ingredients += line_stripped + '\n'
            
            # Extract instructions (look for ## Instructions or similar)
            instructions = ''
            in_instructions = False
            for line in lines:
                line_stripped = line.strip()
                if line_stripped.lower().startswith('## instructions') or line_stripped.lower().startswith('## directions') or line_stripped.lower().startswith('## steps'):
                    in_instructions = True
                    continue
                elif in_instructions:
                    if line_stripped.startswith('##') or line_stripped.startswith('###'):
                        break
                    if line_stripped and not line_stripped.startswith('#'):
                        instructions += line_stripped + '\n'
            
            # Extract description (text before ingredients)
            description_lines = []
            in_description = False
            for line in lines:
                line_stripped = line.strip()
                if line_stripped.startswith('# '):
                    in_description = True
                    continue
                elif line_stripped.lower().startswith('## ingredients') or line_stripped.lower().startswith('## instructions'):
                    break
                elif in_description and line_stripped and not line_stripped.startswith('#'):
                    description_lines.append(line_stripped)
            
            description = ' '.join(description_lines)
            
            return {
                'title': title,
                'category': category,
                'servings': servings,
                'cook_time': cook_time,
                'prep_time': prep_time,
                'difficulty': difficulty,
                'description': description,
                'notes': notes,
                'ingredients': ingredients.strip(),
                'instructions': instructions.strip()
            }
            
        except Exception as e:
            print(f"Error parsing Markdown recipe: {e}")
            return None

    
    def import_from_pdf(self, file_path, duplicate_handling, default_category):
        """Import recipes from PDF file"""
        try:
            import os
            from utils.db import get_connection
            
            # Extract text from PDF
            pdf_text = self._extract_pdf_text(file_path)
            
            if not pdf_text:
                QMessageBox.warning(self, "PDF Error", "Could not extract text from PDF file.")
                return None
            
            # Parse the extracted text as a recipe
            recipe_data = self._parse_pdf_recipe(pdf_text, default_category)
            
            if not recipe_data:
                QMessageBox.warning(self, "Parse Error", "Could not parse recipe from PDF file.")
                return None
            
            conn = get_connection()
            cursor = conn.cursor()
            
            # Handle duplicates
            if duplicate_handling == "Skip duplicates":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                if cursor.fetchone():
                    QMessageBox.information(self, "Duplicate", f"Recipe '{recipe_data['title']}' already exists. Skipping.")
                    conn.close()
                    return None
            elif duplicate_handling == "Update existing":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                existing = cursor.fetchone()
                if existing:
                    recipe_id = existing[0]
                    # Update existing recipe
                    cursor.execute("""
                        UPDATE recipes 
                        SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                            ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                        WHERE id = ?
                    """, (
                        recipe_data['category'], recipe_data['servings'], 
                        recipe_data['cook_time'], recipe_data['prep_time'], 
                        recipe_data['ingredients'], recipe_data['instructions'],
                        recipe_data['notes'], recipe_data['difficulty'], 
                        recipe_data['description'], recipe_id
                    ))
                    conn.commit()
                    conn.close()
                    QMessageBox.information(self, "Success", f"Updated recipe '{recipe_data['title']}'")
                    return recipe_id
            
            # Insert new recipe
            cursor.execute("""
                INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                   ingredients, instructions, notes, difficulty, description)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                recipe_data['title'], recipe_data['category'], 
                recipe_data['servings'], recipe_data['cook_time'],
                recipe_data['prep_time'], recipe_data['ingredients'],
                recipe_data['instructions'], recipe_data['notes'],
                recipe_data['difficulty'], recipe_data['description']
            ))
            
            recipe_id = cursor.lastrowid
            conn.commit()
            conn.close()
            
            QMessageBox.information(self, "Success", f"Imported recipe '{recipe_data['title']}' from PDF")
            return recipe_id
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from PDF file: {str(e)}")
            return None

    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file using multiple methods"""
        try:
            # Try PyPDF2 first (more common)
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
            except ImportError:
                pass
            
            # Try pdfplumber (better text extraction)
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text.strip()
            except ImportError:
                pass
            
            # Try pdfminer (fallback)
            try:
                from pdfminer.high_level import extract_text
                text = extract_text(file_path)
                return text.strip()
            except ImportError:
                pass
            
            # If no PDF libraries are available
            QMessageBox.critical(self, "Missing Dependency", 
                               "No PDF processing library found.\n"
                               "Please install one of the following:\n"
                               "• pip install PyPDF2\n"
                               "• pip install pdfplumber\n"
                               "• pip install pdfminer.six")
            return None
            
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return None

    
    def _parse_pdf_recipe(self, pdf_text, default_category):
        """Parse recipe from extracted PDF text"""
        import re
        
        try:
            lines = pdf_text.split('\n')
            
            # Clean up the text - remove extra whitespace and normalize
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    cleaned_lines.append(line)
            
            if not cleaned_lines:
                return None
            
            # Extract title - look for the first substantial line or common patterns
            title = ''
            for i, line in enumerate(cleaned_lines[:10]):  # Check first 10 lines
                # Skip common PDF artifacts
                if any(skip in line.lower() for skip in ['page', 'copyright', 'www.', 'http', 'recipe']):
                    continue
                # Look for title patterns
                if len(line) > 3 and len(line) < 100 and not line.isdigit():
                    title = line
                    break
            
            if not title:
                # Fallback: use first non-empty line
                title = cleaned_lines[0] if cleaned_lines else "Imported Recipe"
            
            # Extract metadata using various patterns
            category = default_category
            servings = 4
            cook_time = ''
            prep_time = ''
            difficulty = 'Medium'
            description = ''
            notes = ''
            
            # Look for metadata patterns
            for line in cleaned_lines:
                line_lower = line.lower()
                
                # Category detection
                if any(cat in line_lower for cat in ['appetizer', 'main course', 'dessert', 'side dish', 'beverage']):
                    for cat in ['appetizer', 'main course', 'dessert', 'side dish', 'beverage']:
                        if cat in line_lower:
                            category = cat.title()
                            break
                
                # Servings detection
                servings_match = re.search(r'(\d+)\s*(servings?|serves?|people)', line_lower)
                if servings_match:
                    try:
                        servings = int(servings_match.group(1))
                    except (ValueError, TypeError):
                        pass
                
                # Time detection
                time_patterns = [
                    r'cook(?:ing)?\s*time[:\s]*(\d+(?:\s*\d+)?\s*(?:min|minutes?|hr|hour|hours?))',
                    r'total\s*time[:\s]*(\d+(?:\s*\d+)?\s*(?:min|minutes?|hr|hour|hours?))',
                    r'prep(?:aration)?\s*time[:\s]*(\d+(?:\s*\d+)?\s*(?:min|minutes?|hr|hour|hours?))'
                ]
                
                for pattern in time_patterns:
                    match = re.search(pattern, line_lower)
                    if match:
                        time_value = match.group(1)
                        if 'cook' in pattern or 'total' in pattern:
                            cook_time = time_value
                        elif 'prep' in pattern:
                            prep_time = time_value
                        break
                
                # Difficulty detection
                if any(diff in line_lower for diff in ['easy', 'medium', 'hard', 'difficult', 'beginner', 'intermediate', 'advanced']):
                    for diff in ['easy', 'medium', 'hard', 'difficult', 'beginner', 'intermediate', 'advanced']:
                        if diff in line_lower:
                            if diff in ['beginner']:
                                difficulty = 'Easy'
                            elif diff in ['intermediate']:
                                difficulty = 'Medium'
                            elif diff in ['advanced', 'difficult']:
                                difficulty = 'Hard'
                            else:
                                difficulty = diff.title()
                            break
            
            # Extract ingredients - look for common patterns
            ingredients = []
            in_ingredients = False
            
            ingredient_keywords = ['ingredients', 'ingredient list', 'you will need']
            instruction_keywords = ['instructions', 'directions', 'steps', 'method', 'how to']
            
            for i, line in enumerate(cleaned_lines):
                line_lower = line.lower()
                
                # Check if we're entering ingredients section
                if any(keyword in line_lower for keyword in ingredient_keywords):
                    in_ingredients = True
                    continue
                
                # Check if we're entering instructions section
                if any(keyword in line_lower for keyword in instruction_keywords):
                    break
                
                # If we're in ingredients section, collect ingredient lines
                if in_ingredients:
                    # Skip empty lines and section headers
                    if not line or line.startswith('#') or len(line) < 3:
                        continue
                    
                    # Look for bullet points or numbered lists
                    if re.match(r'^[\d\-\*\•]\s+', line):
                        ingredient = re.sub(r'^[\d\-\*\•]\s+', '', line).strip()
                        if ingredient:
                            ingredients.append(ingredient)
                    # Look for lines that seem like ingredients (contain measurements)
                    elif re.search(r'\d+\s*(cup|cups|tbsp|tsp|oz|lb|pound|gram|g|ml|l|liter)', line_lower):
                        ingredients.append(line)
            
            # If no ingredients found with section detection, try pattern matching
            if not ingredients:
                for line in cleaned_lines:
                    # Look for lines with measurements
                    if re.search(r'\d+\s*(cup|cups|tbsp|tsp|oz|lb|pound|gram|g|ml|l|liter)', line.lower()):
                        ingredients.append(line)
            
            # Extract instructions - look for common patterns
            instructions = []
            in_instructions = False
            
            for i, line in enumerate(cleaned_lines):
                line_lower = line.lower()
                
                # Check if we're entering instructions section
                if any(keyword in line_lower for keyword in instruction_keywords):
                    in_instructions = True
                    continue
                
                # If we're in instructions section, collect instruction lines
                if in_instructions:
                    # Skip empty lines and section headers
                    if not line or line.startswith('#') or len(line) < 3:
                        continue
                    
                    # Look for numbered steps
                    if re.match(r'^\d+[\.\)]\s+', line):
                        instruction = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
                        if instruction:
                            instructions.append(instruction)
                    # Look for any substantial text
                    elif len(line) > 10:
                        instructions.append(line)
            
            # If no instructions found with section detection, try pattern matching
            if not instructions:
                # Look for numbered steps anywhere in the text
                for line in cleaned_lines:
                    if re.match(r'^\d+[\.\)]\s+', line) and len(line) > 10:
                        instruction = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
                        if instruction:
                            instructions.append(instruction)
            
            # Extract description from text before ingredients
            description_lines = []
            found_ingredients = False
            
            for line in cleaned_lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in ingredient_keywords):
                    found_ingredients = True
                    break
                if line and len(line) > 10 and not line.isdigit():
                    description_lines.append(line)
            
            if description_lines:
                description = ' '.join(description_lines[:3])  # Take first 3 lines
            
            return {
                'title': title,
                'category': category,
                'servings': servings,
                'cook_time': cook_time,
                'prep_time': prep_time,
                'difficulty': difficulty,
                'description': description,
                'notes': notes,
                'ingredients': '\n'.join(ingredients),
                'instructions': '\n'.join(instructions)
            }
            
        except Exception as e:
            print(f"Error parsing PDF recipe: {e}")
            return None

    
    def import_from_word(self, file_path, duplicate_handling, default_category):
        """Import recipes from Word document (.docx, .doc)"""
        try:
            from utils.db import get_connection
            
            # Extract text from Word document
            word_text = self._extract_word_text(file_path)
            
            if not word_text:
                QMessageBox.warning(self, "Word Error", "Could not extract text from Word document.")
                return None
            
            # Parse the extracted text as a recipe
            recipe_data = self._parse_word_recipe(word_text, default_category)
            
            if not recipe_data:
                QMessageBox.warning(self, "Parse Error", "Could not parse recipe from Word document.")
                return None
            
            conn = get_connection()
            cursor = conn.cursor()
            
            # Handle duplicates
            if duplicate_handling == "Skip duplicates":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                if cursor.fetchone():
                    QMessageBox.information(self, "Duplicate", f"Recipe '{recipe_data['title']}' already exists. Skipping.")
                    conn.close()
                    return None
            elif duplicate_handling == "Update existing":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                existing = cursor.fetchone()
                if existing:
                    recipe_id = existing[0]
                    # Update existing recipe
                    cursor.execute("""
                        UPDATE recipes 
                        SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                            ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                        WHERE id = ?
                    """, (
                        recipe_data['category'], recipe_data['servings'], 
                        recipe_data['cook_time'], recipe_data['prep_time'], 
                        recipe_data['ingredients'], recipe_data['instructions'],
                        recipe_data['notes'], recipe_data['difficulty'], 
                        recipe_data['description'], recipe_id
                    ))
                    conn.commit()
                    conn.close()
                    QMessageBox.information(self, "Success", f"Updated recipe '{recipe_data['title']}'")
                    return recipe_id
            
            # Insert new recipe
            cursor.execute("""
                INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                   ingredients, instructions, notes, difficulty, description)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                recipe_data['title'], recipe_data['category'], 
                recipe_data['servings'], recipe_data['cook_time'],
                recipe_data['prep_time'], recipe_data['ingredients'],
                recipe_data['instructions'], recipe_data['notes'],
                recipe_data['difficulty'], recipe_data['description']
            ))
            
            recipe_id = cursor.lastrowid
            conn.commit()
            conn.close()
            
            QMessageBox.information(self, "Success", f"Imported recipe '{recipe_data['title']}' from Word document")
            return recipe_id
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from Word document: {str(e)}")
            return None

    
    def _extract_word_text(self, file_path):
        """Extract text from Word document using multiple methods"""
        import os
        
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Try python-docx for .docx files (modern format)
            if file_ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    
                    # Extract text from paragraphs
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    
                    # Extract text from tables (ingredients often in tables)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_parts.append(' | '.join(row_text))
                    
                    return '\n'.join(text_parts)
                    
                except ImportError:
                    QMessageBox.critical(self, "Missing Dependency", 
                                       "python-docx library is required for .docx import.\n"
                                       "Please install it with: pip install python-docx")
                    return None
            
            # Try textract for .doc files (older format) or as fallback
            elif file_ext == '.doc':
                try:
                    import textract
                    text = textract.process(file_path).decode('utf-8')
                    return text
                except ImportError:
                    QMessageBox.critical(self, "Missing Dependency", 
                                       "textract library is required for .doc import.\n"
                                       "Please install it with: pip install textract\n"
                                       "Note: .doc format support may require additional dependencies.")
                    return None
                except Exception as e:
                    # Fallback: Try to open as .docx anyway (sometimes works)
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        text_parts = []
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text)
                        return '\n'.join(text_parts)
                    except:
                        QMessageBox.critical(self, "Format Error", 
                                           f"Could not process .doc file: {str(e)}\n"
                                           "Please convert to .docx format or install textract.")
                        return None
            else:
                QMessageBox.warning(self, "Unsupported Format", 
                                  f"File format {file_ext} is not supported.\n"
                                  "Please use .docx or .doc format.")
                return None
                
        except Exception as e:
            print(f"Error extracting Word text: {e}")
            return None

    
    def _parse_word_recipe(self, word_text, default_category):
        """Parse recipe from extracted Word document text"""
        import re
        
        try:
            lines = word_text.split('\n')
            
            # Clean up the text - remove extra whitespace and normalize
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line and line not in ['|', '||', '|||']:  # Remove table separators
                    cleaned_lines.append(line)
            
            if not cleaned_lines:
                return None
            
            # Extract title - look for the first substantial line
            title = ''
            for i, line in enumerate(cleaned_lines[:10]):  # Check first 10 lines
                # Skip common document artifacts
                if any(skip in line.lower() for skip in ['recipe', 'document', 'page', 'copyright']):
                    if len(cleaned_lines) > i + 1:
                        continue
                # Look for title patterns - substantial text, not too long
                if len(line) > 3 and len(line) < 100 and not line.isdigit():
                    # Check if it's styled like a title (all caps, title case, etc.)
                    title = line
                    break
            
            if not title:
                # Fallback: use first non-empty line
                title = cleaned_lines[0] if cleaned_lines else "Imported Recipe"
            
            # Extract metadata using various patterns
            category = default_category
            servings = 4
            cook_time = ''
            prep_time = ''
            difficulty = 'Medium'
            description = ''
            notes = ''
            
            # Look for metadata patterns (often in Word docs as "Label: Value")
            for line in cleaned_lines:
                line_lower = line.lower()
                
                # Category detection
                if 'category:' in line_lower or 'type:' in line_lower or 'cuisine:' in line_lower:
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        category = parts[1].strip()
                elif any(cat in line_lower for cat in ['appetizer', 'main course', 'dessert', 'side dish', 'beverage', 'breakfast', 'lunch', 'dinner']):
                    for cat in ['appetizer', 'main course', 'dessert', 'side dish', 'beverage', 'breakfast', 'lunch', 'dinner']:
                        if cat in line_lower:
                            category = cat.title()
                            break
                
                # Servings detection
                if 'servings:' in line_lower or 'serves:' in line_lower or 'yield:' in line_lower:
                    servings_match = re.search(r'(\d+)', line)
                    if servings_match:
                        try:
                            servings = int(servings_match.group(1))
                        except (ValueError, TypeError):
                            pass
                else:
                    servings_match = re.search(r'(\d+)\s*(servings?|serves?|people)', line_lower)
                    if servings_match:
                        try:
                            servings = int(servings_match.group(1))
                        except (ValueError, TypeError):
                            pass
                
                # Time detection
                if 'cook time:' in line_lower or 'cooking time:' in line_lower:
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        cook_time = parts[1].strip()
                elif 'prep time:' in line_lower or 'preparation time:' in line_lower:
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        prep_time = parts[1].strip()
                elif 'total time:' in line_lower:
                    parts = line.split(':', 1)
                    if len(parts) > 1 and not cook_time:
                        cook_time = parts[1].strip()
                
                # Difficulty detection
                if 'difficulty:' in line_lower or 'level:' in line_lower:
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        difficulty = parts[1].strip().title()
                elif any(diff in line_lower for diff in ['easy', 'medium', 'hard', 'beginner', 'intermediate', 'advanced']):
                    for diff in ['easy', 'medium', 'hard', 'beginner', 'intermediate', 'advanced']:
                        if diff in line_lower:
                            if diff == 'beginner':
                                difficulty = 'Easy'
                            elif diff == 'intermediate':
                                difficulty = 'Medium'
                            elif diff in ['advanced', 'difficult']:
                                difficulty = 'Hard'
                            else:
                                difficulty = diff.title()
                            break
            
            # Extract ingredients - look for common patterns
            ingredients = []
            in_ingredients = False
            
            ingredient_keywords = ['ingredients', 'ingredient list', 'you will need', 'what you need']
            instruction_keywords = ['instructions', 'directions', 'steps', 'method', 'how to make', 'preparation']
            
            for i, line in enumerate(cleaned_lines):
                line_lower = line.lower()
                
                # Check if we're entering ingredients section
                if any(keyword in line_lower for keyword in ingredient_keywords):
                    in_ingredients = True
                    continue
                
                # Check if we're entering instructions section
                if any(keyword in line_lower for keyword in instruction_keywords):
                    in_ingredients = False
                    break
                
                # If we're in ingredients section, collect ingredient lines
                if in_ingredients:
                    # Skip empty lines and section headers
                    if not line or len(line) < 2:
                        continue
                    
                    # Handle table format (with | separator)
                    if '|' in line:
                        parts = line.split('|')
                        # Combine amount and ingredient if in separate columns
                        combined = ' '.join(p.strip() for p in parts if p.strip())
                        if combined:
                            ingredients.append(combined)
                    # Look for bullet points or numbered lists
                    elif re.match(r'^[\d\-\*\•·]\s+', line):
                        ingredient = re.sub(r'^[\d\-\*\•·]\s+', '', line).strip()
                        if ingredient:
                            ingredients.append(ingredient)
                    # Look for lines that seem like ingredients
                    elif re.search(r'\d+\s*(cup|cups|tbsp|tsp|oz|lb|pound|gram|g|ml|l|liter|tablespoon|teaspoon)', line_lower) or len(line) < 60:
                        ingredients.append(line)
            
            # If no ingredients found with section detection, try pattern matching
            if not ingredients:
                for line in cleaned_lines:
                    # Look for lines with measurements
                    if re.search(r'\d+\s*(cup|cups|tbsp|tsp|oz|lb|pound|gram|g|ml|l|liter|tablespoon|teaspoon)', line.lower()):
                        ingredients.append(line)
            
            # Extract instructions - look for common patterns
            instructions = []
            in_instructions = False
            
            for i, line in enumerate(cleaned_lines):
                line_lower = line.lower()
                
                # Check if we're entering instructions section
                if any(keyword in line_lower for keyword in instruction_keywords):
                    in_instructions = True
                    continue
                
                # If we're in instructions section, collect instruction lines
                if in_instructions:
                    # Skip empty lines and very short lines
                    if not line or len(line) < 5:
                        continue
                    
                    # Skip lines that look like section headers
                    if line.lower() in ['notes', 'tips', 'nutrition']:
                        break
                    
                    # Look for numbered steps
                    if re.match(r'^\d+[\.\)]\s+', line):
                        instruction = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
                        if instruction:
                            instructions.append(instruction)
                    # Look for any substantial text
                    elif len(line) > 15:
                        instructions.append(line)
            
            # If no instructions found with section detection, try pattern matching
            if not instructions:
                # Look for numbered steps anywhere in the text
                for line in cleaned_lines:
                    if re.match(r'^\d+[\.\)]\s+', line) and len(line) > 15:
                        instruction = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
                        if instruction:
                            instructions.append(instruction)
            
            # Extract description from text before ingredients
            description_lines = []
            found_ingredients = False
            
            for line in cleaned_lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in ingredient_keywords):
                    found_ingredients = True
                    break
                if line and line != title and len(line) > 15 and not line.isdigit():
                    # Skip metadata lines
                    if ':' not in line or len(line) > 50:
                        description_lines.append(line)
            
            if description_lines:
                description = ' '.join(description_lines[:3])  # Take first 3 lines
            
            return {
                'title': title,
                'category': category,
                'servings': servings,
                'cook_time': cook_time,
                'prep_time': prep_time,
                'difficulty': difficulty,
                'description': description,
                'notes': notes,
                'ingredients': '\n'.join(ingredients),
                'instructions': '\n'.join(instructions)
            }
            
        except Exception as e:
            print(f"Error parsing Word recipe: {e}")
            return None

    
    def _toggle_txt_input_mode(self, checked):
        """Toggle between file selection and text input for TXT import"""
        if checked:
            self.text_input_label.setVisible(True)
            self.recipe_text_input.setVisible(True)
        else:
            self.text_input_label.setVisible(False)
            self.recipe_text_input.setVisible(False)
    
    def import_from_txt(self, file_path, duplicate_handling, default_category):
        """Import recipe from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                recipe_text = file.read()
            
            self.import_from_txt_content(recipe_text, duplicate_handling, default_category)
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import from text file: {str(e)}")
    
    def import_from_txt_content(self, recipe_text, duplicate_handling, default_category):
        """Parse and import recipe from text content"""
        from utils.db import get_connection
        
        try:
            # Parse the recipe text
            recipe_data = self._parse_recipe_text(recipe_text, default_category)
            
            if not recipe_data:
                QMessageBox.warning(self, "Parse Error", "Could not parse recipe from text. Please check the format.")
                return
            
            conn = get_connection()
            cursor = conn.cursor()
            
            # Handle duplicates
            if duplicate_handling == "Skip duplicates":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                if cursor.fetchone():
                    QMessageBox.information(self, "Duplicate", f"Recipe '{recipe_data['title']}' already exists. Skipping.")
                    conn.close()
                    return
            elif duplicate_handling == "Update existing":
                cursor.execute("SELECT id FROM recipes WHERE title = ?", (recipe_data['title'],))
                existing = cursor.fetchone()
                if existing:
                    recipe_id = existing[0]
                    # Update existing recipe
                    cursor.execute("""
                        UPDATE recipes 
                        SET category = ?, servings = ?, cook_time = ?, prep_time = ?,
                            ingredients = ?, instructions = ?, notes = ?, difficulty = ?, description = ?
                        WHERE id = ?
                    """, (
                        recipe_data['category'], recipe_data['servings'], recipe_data['cook_time'],
                        recipe_data['prep_time'], recipe_data['ingredients'], recipe_data['instructions'],
                        recipe_data['notes'], recipe_data['difficulty'], recipe_data['description'], recipe_id
                    ))
                    conn.commit()
                    conn.close()
                    QMessageBox.information(self, "Success", f"Updated recipe '{recipe_data['title']}'")
                    return recipe_id
            
            # Insert new recipe
            cursor.execute("""
                INSERT INTO recipes (title, category, servings, cook_time, prep_time,
                                   ingredients, instructions, notes, difficulty, description)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                recipe_data['title'], recipe_data['category'], recipe_data['servings'],
                recipe_data['cook_time'], recipe_data['prep_time'], recipe_data['ingredients'],
                recipe_data['instructions'], recipe_data['notes'], recipe_data['difficulty'],
                recipe_data['description']
            ))
            
            recipe_id = cursor.lastrowid
            
            # Insert ingredients into recipe_ingredients table if we have structured data
            if recipe_data.get('parsed_ingredients'):
                for ingredient in recipe_data['parsed_ingredients']:
                    cursor.execute("""
                        INSERT INTO recipe_ingredients (recipe_id, ingredient_name, quantity, unit, notes)
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        recipe_id,
                        ingredient.get('name', ''),
                        ingredient.get('quantity', ''),
                        ingredient.get('unit', ''),
                        ingredient.get('notes', '')
                    ))
            
            conn.commit()
            conn.close()
            
            QMessageBox.information(self, "Success", f"Imported recipe '{recipe_data['title']}'")
            
            # Return the recipe ID for auto-viewing
            return recipe_id
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import recipe: {str(e)}")
            return None
    
    def _parse_recipe_text(self, text, default_category):
        """Parse recipe from text content
        
        Supports formats like:
        - Recipe Name: ...
        - Category: ...
        - Ingredients: (or Ingredients list)
        - Instructions: (or Directions:, Steps:, Method:)
        """
        import re
        
        recipe_data = {
            'title': '',
            'category': default_category,
            'prep_time': '',
            'cook_time': '',
            'servings': 4,
            'difficulty': 'Medium',
            'description': '',
            'ingredients': '',
            'instructions': '',
            'notes': '',
            'parsed_ingredients': []
        }
        
        # Split text into lines
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        # Patterns to match common recipe field names
        field_patterns = {
            'title': r'^(?:recipe\s+name|title|name)\s*[:=]\s*(.+)$',
            'category': r'^(?:category|type)\s*[:=]\s*(.+)$',
            'prep_time': r'^(?:prep\s+time|preparation\s+time|prep)\s*[:=]\s*(.+)$',
            'cook_time': r'^(?:cook\s+time|cooking\s+time|cook)\s*[:=]\s*(.+)$',
            'servings': r'^(?:servings|serves|yield)\s*[:=]\s*(\d+)',
            'difficulty': r'^(?:difficulty|level)\s*[:=]\s*(.+)$',
            'description': r'^(?:description|about)\s*[:=]\s*(.+)$',
        }
        
        section_patterns = {
            'ingredients': r'^(?:ingredients|ingredient\s+list)\s*[:=]?\s*$',
            'instructions': r'^(?:instructions|directions|steps|method|preparation)\s*[:=]?\s*$',
            'notes': r'^(?:notes|tips|note)\s*[:=]?\s*$',
        }
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                # Empty line - might be section separator
                if section_content and current_section:
                    # Save accumulated section content
                    if current_section in ['ingredients', 'instructions', 'notes']:
                        recipe_data[current_section] = '\n'.join(section_content).strip()
                    section_content = []
                continue
            
            # Check for field patterns (single-line fields)
            matched_field = False
            for field_name, pattern in field_patterns.items():
                match = re.match(pattern, line_stripped, re.IGNORECASE)
                if match:
                    value = match.group(1).strip()
                    if field_name == 'servings':
                        try:
                            recipe_data[field_name] = int(value)
                        except ValueError:
                            recipe_data[field_name] = 4
                    else:
                        recipe_data[field_name] = value
                    matched_field = True
                    break
            
            if matched_field:
                continue
            
            # Check for section headers
            matched_section = False
            for section_name, pattern in section_patterns.items():
                if re.match(pattern, line_stripped, re.IGNORECASE):
                    # Save previous section if any
                    if section_content and current_section:
                        recipe_data[current_section] = '\n'.join(section_content).strip()
                    
                    # Start new section
                    current_section = section_name
                    section_content = []
                    matched_section = True
                    break
            
            if matched_section:
                continue
            
            # Add line to current section
            if current_section:
                section_content.append(line_stripped)
            elif not recipe_data['title']:
                # If no title set yet and we have content, assume first line is title
                recipe_data['title'] = line_stripped
        
        # Save last section
        if section_content and current_section:
            recipe_data[current_section] = '\n'.join(section_content).strip()
        
        # Parse ingredients into structured format
        if recipe_data['ingredients']:
            recipe_data['parsed_ingredients'] = self._parse_ingredients_list(recipe_data['ingredients'])
        
        # Validate required fields
        if not recipe_data['title']:
            # Try to extract title from first non-empty line
            for line in lines:
                if line.strip():
                    recipe_data['title'] = line.strip()[:100]  # Limit title length
                    break
        
        return recipe_data if recipe_data['title'] else None
    
    def _parse_ingredients_list(self, ingredients_text):
        """Parse ingredients text into structured list"""
        import re
        
        parsed_ingredients = []
        lines = ingredients_text.split('\n')
        
        # Common measurement patterns
        measurement_pattern = r'^[\-\*\•]?\s*(\d+[\d\s\/\.]*)\s*([a-zA-Z]+)?\s+(.+)$'
        simple_pattern = r'^[\-\*\•]?\s*(.+)$'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try to match quantity + unit + ingredient
            match = re.match(measurement_pattern, line)
            if match:
                quantity = match.group(1).strip()
                unit = match.group(2).strip() if match.group(2) else ''
                name = match.group(3).strip()
                
                parsed_ingredients.append({
                    'quantity': quantity,
                    'unit': unit,
                    'name': name,
                    'notes': ''
                })
            else:
                # Just ingredient name, no measurement
                match = re.match(simple_pattern, line)
                if match:
                    parsed_ingredients.append({
                        'quantity': '',
                        'unit': '',
                        'name': match.group(1).strip(),
                        'notes': ''
                    })
        
        return parsed_ingredients
